#!/usr/bin/env python3
"""
tools/generate_invoice_template.py — JR Anchored owner use only

Generates two templates in this folder:
  - proforma_invoice_template.docx  (sent with the zip after PayPal payment)
  - receipt_template.docx           (confirms payment received)

Fill in your real details once and save personal copies outside the repo.

Usage:
    python3 tools/generate_invoice_template.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))

BLUE  = RGBColor(0x2E, 0x5B, 0xBA)
GREY  = RGBColor(0x70, 0x80, 0xA0)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def add_run(para, text, bold=False, size=10, color=BLACK, italic=False):
    run = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    return run


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    return doc


def add_sender_header(doc, label):
    """Brand left, document type label right."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment    = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(10)
    tbl.columns[1].width = Cm(6)

    left = tbl.cell(0, 0)
    for text, bold, size, color in [
        ("JR Anchored",         True,  14, BLUE),
        ("[YOUR TRADING NAME]", False, 9,  GREY),
        ("[ADDRESS LINE 1]",    False, 9,  GREY),
        ("[POSTCODE  CITY]",    False, 9,  GREY),
        ("[EMAIL ADDRESS]",     False, 9,  GREY),
        ("KvK: [KVK NUMBER]",  False, 9,  GREY),
    ]:
        p = left.add_paragraph()
        no_space(p)
        add_run(p, text, bold=bold, size=size, color=color)
    left.paragraphs[0]._element.getparent().remove(left.paragraphs[0]._element)

    right = tbl.cell(0, 1)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    no_space(p)
    add_run(p, label, bold=True, size=16, color=BLUE)


def add_divider(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    no_space(p)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xD4, 0xEE)
    run.font.size = Pt(7)


def add_meta_and_bill_to(doc, meta_rows):
    """Two-column table: meta fields left, Bill to right."""
    meta = doc.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.columns[0].width = Cm(8)
    meta.columns[1].width = Cm(8)

    ml = meta.cell(0, 0)
    mr = meta.cell(0, 1)

    for label, value in meta_rows:
        p = ml.add_paragraph()
        no_space(p)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{label}: ", bold=True, size=9, color=BLACK)
        add_run(p, value, size=9, color=GREY)
    ml.paragraphs[0]._element.getparent().remove(ml.paragraphs[0]._element)

    p = mr.add_paragraph()
    no_space(p)
    add_run(p, "Bill to", bold=True, size=9, color=BLACK)
    mr.paragraphs[0]._element.getparent().remove(mr.paragraphs[0]._element)
    for line in ["[BUYER NAME / COMPANY]", "[ADDRESS]", "[CITY, COUNTRY]", "[EMAIL]"]:
        p = mr.add_paragraph()
        no_space(p)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, line, size=9, color=GREY)


def add_line_items(doc, description):
    doc.add_paragraph()
    items = doc.add_table(rows=2, cols=3)
    items.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths  = [Cm(11), Cm(2), Cm(3)]
    headers = ["Description", "Qty", "Amount"]

    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = items.cell(0, i)
        cell.width = w
        set_cell_bg(cell, "2E5BBA")
        p = cell.paragraphs[0]
        no_space(p)
        p.paragraph_format.space_after = Pt(4)
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, hdr, bold=True, size=9, color=WHITE)

    for i, (text, w) in enumerate(zip([description, "1", "$ 100.00"], widths)):
        cell = items.rows[1].cells[i]
        cell.width = w
        p = cell.paragraphs[0]
        no_space(p)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, text, size=9, color=BLACK)


def add_totals(doc, total_label="Total"):
    doc.add_paragraph()
    totals = doc.add_table(rows=2, cols=2)
    totals.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, (label, value) in enumerate([
        ("VAT",        "Exempt (KOR)"),
        (total_label,  "$ 100.00"),
    ]):
        lc = totals.cell(r, 0)
        vc = totals.cell(r, 1)
        lc.width = Cm(13)
        vc.width = Cm(3)
        pl = lc.paragraphs[0]
        pv = vc.paragraphs[0]
        no_space(pl); no_space(pv)
        pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        bold = (r == 1)
        add_run(pl, label, bold=bold, size=9, color=BLACK if bold else GREY)
        add_run(pv, value, bold=bold, size=9, color=BLACK if bold else GREY)


def add_kor_footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    no_space(p)
    p.paragraph_format.space_before = Pt(16)
    add_run(p, "BTW vrijgesteld op grond van de kleineondernemersregeling (KOR).",
            size=8, color=GREY, italic=True)


# ── Proforma invoice ──────────────────────────────────────────────────────────

def build_proforma_invoice():
    doc = new_doc()

    add_sender_header(doc, "PROFORMA INVOICE")
    add_divider(doc)
    add_meta_and_bill_to(doc, [
        ("Invoice number", "[INVOICE NUMBER]"),
        ("Date",           "[DD MONTH YYYY]"),
    ])
    add_line_items(doc,
        "JR Anchored Validation Pack v[VERSION] — perpetual licence, one organisation")
    add_totals(doc, total_label="Total")

    # Payment note — payment already received via PayPal
    doc.add_paragraph()
    p = doc.add_paragraph()
    no_space(p)
    p.paragraph_format.space_before = Pt(10)
    add_run(p, "Payment received via PayPal on [DD MONTH YYYY].",
            size=9, color=GREY, italic=True)

    add_kor_footer(doc)

    out = os.path.join(TOOLS_DIR, "proforma_invoice_template.docx")
    doc.save(out)
    print(f"✅  Proforma invoice template written to: {out}")


# ── Receipt ───────────────────────────────────────────────────────────────────

def build_receipt():
    doc = new_doc()

    add_sender_header(doc, "RECEIPT")
    add_divider(doc)

    # Meta: receipt number + date, issued to right
    meta = doc.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.columns[0].width = Cm(8)
    meta.columns[1].width = Cm(8)

    ml = meta.cell(0, 0)
    mr = meta.cell(0, 1)

    for label, value in [
        ("Receipt number",   "[RECEIPT NUMBER]"),
        ("Date of payment",  "[DD MONTH YYYY]"),
        ("Payment method",   "PayPal"),
        ("Transaction ID",   "[PAYPAL TRANSACTION ID]"),
    ]:
        p = ml.add_paragraph()
        no_space(p)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{label}: ", bold=True, size=9, color=BLACK)
        add_run(p, value, size=9, color=GREY)
    ml.paragraphs[0]._element.getparent().remove(ml.paragraphs[0]._element)

    p = mr.add_paragraph()
    no_space(p)
    add_run(p, "Received from", bold=True, size=9, color=BLACK)
    mr.paragraphs[0]._element.getparent().remove(mr.paragraphs[0]._element)
    for line in ["[BUYER NAME / COMPANY]", "[ADDRESS]", "[CITY, COUNTRY]", "[EMAIL]"]:
        p = mr.add_paragraph()
        no_space(p)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, line, size=9, color=GREY)

    add_line_items(doc,
        "JR Anchored Validation Pack v[VERSION] — perpetual licence, one organisation")
    add_totals(doc, total_label="Total received")
    add_kor_footer(doc)

    out = os.path.join(TOOLS_DIR, "receipt_template.docx")
    doc.save(out)
    print(f"✅  Receipt template written to:          {out}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print()
    print("JR Anchored — Generate document templates")
    print("=" * 42)
    build_proforma_invoice()
    build_receipt()
    print()
