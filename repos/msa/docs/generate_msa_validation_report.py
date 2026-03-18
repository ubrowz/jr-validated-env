"""
Generates msa_validation_report.docx — FDA-acceptable OQ Validation Report
for the JR Validated Environment MSA module v1.0.

Run from the repo root:
    python3 repos/msa/docs/generate_msa_validation_report.py
Output: repos/msa/docs/msa_validation_report.docx
"""

import os
import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "msa_validation_report.docx")

DARK  = "1A1A2E"
GREY  = "F2F2F2"
WHITE = "FFFFFF"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def header_row(table, *texts):
    row = table.rows[0]
    for cell, text in zip(row.cells, texts):
        set_cell_shading(cell, DARK)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)


def data_row(table, *texts, alt=True):
    row = table.add_row()
    fill = GREY if alt else WHITE
    for cell, text in zip(row.cells, texts):
        set_cell_shading(cell, fill)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text))
        run.font.size = Pt(9)
    return row


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def para(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = docx.Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)

for sn in ("Normal", "Heading 1", "Heading 2", "List Bullet", "List Number"):
    try:
        s = doc.styles[sn]
        s.font.name = "Calibri"
        if sn == "Normal":
            s.font.size = Pt(10)
        elif sn in ("Heading 1", "Heading 2"):
            s.font.bold = True
            s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            s.font.size = Pt(14 if sn == "Heading 1" else 13)
    except KeyError:
        pass

# ===========================================================================
# COVER PAGE
# ===========================================================================

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("VALIDATION REPORT\nOPERATIONAL QUALIFICATION\nMSA MODULE")
r.bold = True
r.font.size = Pt(18)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("JR Validated Environment — MSA Module v1.0")
r2.bold = True
r2.font.size = Pt(13)
r2.font.name = "Calibri"

doc.add_paragraph()

tbl_cover = doc.add_table(rows=1, cols=2)
tbl_cover.style = "Table Grid"
header_row(tbl_cover, "Field", "Value")
cover_rows = [
    ("Document Number",   "JR-VR-MSA-001"),
    ("Title",             "Validation Report — Operational Qualification, MSA Module"),
    ("Validation Plan",   "JR-VP-MSA-001 v1.0"),
    ("System",            "JR Validated Environment — MSA Module"),
    ("Module Version",    "1.0"),
    ("Document Version",  "1.0"),
    ("Execution Date",    "2026-03-18"),
    ("OQ Result",         "PASS — 53 / 53 tests passed, 0 failures, 0 deviations"),
    ("Author",            "Joep Rous"),
    ("Reviewer",          "[Name]"),
    ("Approver",          "[Name]"),
]
for i, (f, v) in enumerate(cover_rows):
    data_row(tbl_cover, f, v, alt=(i % 2 == 0))
set_col_widths(tbl_cover, [5, 11])

doc.add_page_break()

# ===========================================================================
# VERSION HISTORY
# ===========================================================================

h1(doc, "Version History")
tbl_ver = doc.add_table(rows=1, cols=4)
tbl_ver.style = "Table Grid"
header_row(tbl_ver, "Version", "Date", "Author", "Description")
data_row(tbl_ver, "1.0", "2026-03-18", "Joep Rous",
         "Initial OQ execution report. All 53 test cases passed. No deviations.", alt=True)
set_col_widths(tbl_ver, [2, 3, 4, 8])

doc.add_paragraph()

# ===========================================================================
# APPROVAL SIGNATURES
# ===========================================================================

h1(doc, "Approval Signatures")
tbl_appr = doc.add_table(rows=1, cols=4)
tbl_appr.style = "Table Grid"
header_row(tbl_appr, "Role", "Name", "Signature", "Date")
data_row(tbl_appr, "Author",        "Joep Rous", "", "2026-03-18", alt=True)
data_row(tbl_appr, "Reviewer",      "[Name]",    "", "",            alt=False)
data_row(tbl_appr, "Approver (QA)", "[Name]",    "", "",            alt=True)
set_col_widths(tbl_appr, [4, 4, 5, 4])

doc.add_page_break()

# ===========================================================================
# 1. PURPOSE
# ===========================================================================

h1(doc, "1.  Purpose")
para(doc,
    "This document reports the execution and outcome of the Operational Qualification "
    "(OQ) test suite for the JR Validated Environment MSA Module, version 1.0. It "
    "provides objective evidence that all MSA scripts perform as specified in the OQ "
    "Validation Plan JR-VP-MSA-001 v1.0 under the defined test conditions.")
para(doc,
    "This report records the test environment, execution results for all 53 test cases, "
    "and the formal OQ conclusion.")

# ===========================================================================
# 2. SCOPE
# ===========================================================================

h1(doc, "2.  Scope")
h2(doc, "2.1  In Scope")
for item in [
    "All 5 R scripts in repos/msa/R/ (jrc_msa_gauge_rr, jrc_msa_nested_grr, "
    "jrc_msa_linearity_bias, jrc_msa_type1, jrc_msa_attribute)",
    "Correct computation of outputs for valid inputs",
    "Correct rejection of invalid inputs with informative error messages",
    "Bypass-protection verification (RENV_PATHS_ROOT check)",
    "PNG output to ~/Downloads/",
]:
    bullet(doc, item)

h2(doc, "2.2  Out of Scope")
for item in [
    "Demo scripts jrc_msa_R_hello.R and jrc_msa_py_hello.py",
    "Infrastructure scripts in bin/ and admin/ (covered by JR-IQ-001)",
    "Performance qualification (PQ) with real production data",
    "Validation of R interpreter and third-party packages (ggplot2, grid)",
]:
    bullet(doc, item)

# ===========================================================================
# 3. REFERENCE DOCUMENTS
# ===========================================================================

h1(doc, "3.  Reference Documents")
doc.add_paragraph()
tbl_ref = doc.add_table(rows=1, cols=3)
tbl_ref.style = "Table Grid"
header_row(tbl_ref, "Document ID", "Title", "Location")
refs = [
    ("JR-VP-MSA-001 v1.0",
     "Validation Plan — OQ, MSA Module",
     "repos/msa/docs/msa_validation_plan.docx"),
    ("JR-VP-002 v1.0",
     "Validation Plan — OQ, Community Script Suite v1.1.0",
     "docs/ignore/oq_validation_plan.docx"),
    ("JR-IQ-001",
     "IQ Execution Evidence",
     "docs/IQ_validation_20260311_205146.txt"),
    ("OQ Evidence",
     "MSA OQ Execution Evidence (pytest output + header)",
     "~/.jrscript/MyProject/validation/msa_oq_execution_20260318T150625.txt"),
]
for i, row in enumerate(refs):
    data_row(tbl_ref, *row, alt=(i % 2 == 0))
set_col_widths(tbl_ref, [4, 7, 6])

# ===========================================================================
# 4. TEST ENVIRONMENT
# ===========================================================================

h1(doc, "4.  Test Environment")
doc.add_paragraph()
tbl_env = doc.add_table(rows=1, cols=2)
tbl_env.style = "Table Grid"
header_row(tbl_env, "Item", "Value")
env_rows = [
    ("Execution date / time",       "2026-03-18T15:06:25"),
    ("Completion date / time",      "2026-03-18T15:06:47"),
    ("Host",                        "Joeps-MacBook-Air.local"),
    ("Operating system",            "macOS 15.7.3 (Darwin 24.6.0, aarch64)"),
    ("R version",                   "4.5.0"),
    ("Python version",              "3.11.9"),
    ("pytest version",              "8.3.4"),
    ("OQ virtual environment",      "~/.venvs/MyProject_oq/  (reused from core OQ suite)"),
    ("Script entry point",          "bin/jrrun  (sets RENV_PATHS_ROOT, loads renv library)"),
    ("Test runner",                 "repos/msa/admin_msa_oq"),
    ("Test directory",              "repos/msa/oq/"),
    ("Test data directory",         "repos/msa/oq/data/  (22 committed synthetic CSV files)"),
    ("Evidence file",               "~/.jrscript/MyProject/validation/"
                                    "msa_oq_execution_20260318T150625.txt"),
]
for i, row in enumerate(env_rows):
    data_row(tbl_env, *row, alt=(i % 2 == 0))
set_col_widths(tbl_env, [6, 11])

# ===========================================================================
# 5. TEST DATA
# ===========================================================================

h1(doc, "5.  Test Data")
para(doc,
    "All test data files are synthetic CSV files committed to repos/msa/oq/data/ "
    "and included in the project integrity check (admin/project_integrity.sha256). "
    "No real patient or production data is used.")
doc.add_paragraph()
tbl_data = doc.add_table(rows=1, cols=3)
tbl_data.style = "Table Grid"
header_row(tbl_data, "File", "Content", "Used by")
data_files = [
    ("gauge_rr_balanced.csv",
     "10 parts × 3 operators × 3 reps, low noise",
     "TC-MSA-GRR-001..004"),
    ("gauge_rr_missing_col.csv",
     "Missing 'operator' column",
     "TC-MSA-GRR-007"),
    ("gauge_rr_unbalanced.csv",
     "Unequal replicates per cell",
     "TC-MSA-GRR-008"),
    ("gauge_rr_one_operator.csv",
     "Single operator only",
     "TC-MSA-GRR-009"),
    ("nested_grr_good.csv",
     "5 parts × 3 operators × 2 reps, low noise, nested design",
     "TC-MSA-NGR-001..005"),
    ("nested_grr_poor.csv",
     "High within-operator replicate variation (large EV)",
     "TC-MSA-NGR-003"),
    ("nested_grr_missing_col.csv",
     "Missing 'replicate' column",
     "TC-MSA-NGR-008"),
    ("nested_grr_one_operator.csv",
     "Single operator only",
     "TC-MSA-NGR-009"),
    ("nested_grr_unbalanced.csv",
     "Operator A has 3 parts, operator B has 2 parts",
     "TC-MSA-NGR-010"),
    ("linearity_bias_good.csv",
     "5 reference levels × 3 reps, slope ≈ 0.05",
     "TC-MSA-LB-001..004"),
    ("linearity_bias_missing_col.csv",
     "Missing 'reference' column",
     "TC-MSA-LB-007"),
    ("linearity_bias_one_part.csv",
     "Single reference level (cannot fit regression)",
     "TC-MSA-LB-008"),
    ("linearity_bias_inconsistent_ref.csv",
     "Same part ID mapped to two different reference values",
     "TC-MSA-LB-009"),
    ("type1_good.csv",
     "25 measurements near reference (50.0), low bias",
     "TC-MSA-T1-001..004"),
    ("type1_biased.csv",
     "25 measurements with systematic positive bias ≈ 0.15",
     "TC-MSA-T1-003"),
    ("type1_missing_value_col.csv",
     "Missing 'value' column",
     "TC-MSA-T1-009"),
    ("type1_too_few.csv",
     "Only 5 measurements (< 10 required)",
     "TC-MSA-T1-010"),
    ("attribute_with_ref.csv",
     "3 appraisers × 10 parts × 2 trials, with reference column; "
     "appraiser A is perfect vs reference",
     "TC-MSA-ATT-001, 003..005"),
    ("attribute_no_ref.csv",
     "Same design without reference column",
     "TC-MSA-ATT-002"),
    ("attribute_missing_col.csv",
     "Missing 'trial' column",
     "TC-MSA-ATT-008"),
    ("attribute_one_appraiser.csv",
     "Single appraiser only",
     "TC-MSA-ATT-009"),
    ("attribute_unbalanced.csv",
     "Unequal trials per appraiser-part cell",
     "TC-MSA-ATT-010"),
]
for i, row in enumerate(data_files):
    data_row(tbl_data, *row, alt=(i % 2 == 0))
set_col_widths(tbl_data, [5.5, 6, 5.5])

# ===========================================================================
# 6. TEST EXECUTION RESULTS
# ===========================================================================

h1(doc, "6.  Test Execution Results")
h2(doc, "6.1  Summary")
doc.add_paragraph()
tbl_sum = doc.add_table(rows=1, cols=2)
tbl_sum.style = "Table Grid"
header_row(tbl_sum, "Item", "Value")
summary = [
    ("Total test cases",    "53"),
    ("Passed",              "53"),
    ("Failed",              "0"),
    ("Errors",              "0"),
    ("Duration",            "22.45 s"),
    ("OQ result",           "PASS"),
]
for i, row in enumerate(summary):
    data_row(tbl_sum, *row, alt=(i % 2 == 0))
set_col_widths(tbl_sum, [6, 11])

doc.add_paragraph()
h2(doc, "6.2  Results by Test File")
doc.add_paragraph()
tbl_byfile = doc.add_table(rows=1, cols=4)
tbl_byfile.style = "Table Grid"
header_row(tbl_byfile, "Test file", "Script(s) covered", "Tests", "Result")
byfile = [
    ("test_msa_attribute.py",    "jrc_msa_attribute",      "11", "11 / 11  PASS"),
    ("test_msa_gauge_rr.py",     "jrc_msa_gauge_rr",       "10", "10 / 10  PASS"),
    ("test_msa_linearity_bias.py","jrc_msa_linearity_bias","10", "10 / 10  PASS"),
    ("test_msa_nested_grr.py",   "jrc_msa_nested_grr",     "11", "11 / 11  PASS"),
    ("test_msa_type1.py",        "jrc_msa_type1",          "11", "11 / 11  PASS"),
]
for i, row in enumerate(byfile):
    data_row(tbl_byfile, *row, alt=(i % 2 == 0))
set_col_widths(tbl_byfile, [5.5, 5, 2, 4.5])

doc.add_paragraph()
h2(doc, "6.3  Individual Test Case Results")
para(doc,
    "All 53 test cases passed. The full pytest output, including individual test IDs, "
    "pass/fail status, and timing, is recorded in the evidence file:")
para(doc,
    "    ~/.jrscript/MyProject/validation/msa_oq_execution_20260318T150625.txt",
    bold=True)
para(doc,
    "That file also contains the evidence header (host, R version, Python version, "
    "pytest version, timestamp) and the OQ RESULT: PASS footer. The per-test "
    "results are reproduced in the table below for traceability.")

doc.add_paragraph()
tbl_tc = doc.add_table(rows=1, cols=4)
tbl_tc.style = "Table Grid"
header_row(tbl_tc, "Test Case ID", "Description", "Expected", "Result")
tc_rows = [
    # --- Gauge R&R ---
    ("TC-MSA-GRR-001", "Happy path — exit 0, sections present",             "Exit 0, sections present",        "PASS"),
    ("TC-MSA-GRR-002", "Known data — %GRR < 10%, verdict ACCEPTABLE",       "%GRR < 10%, ACCEPTABLE",          "PASS"),
    ("TC-MSA-GRR-003", "--tolerance flag",                                   "Exit 0, tolerance output",        "PASS"),
    ("TC-MSA-GRR-004", "PNG written to ~/Downloads/",                        "PNG present, recent mtime",       "PASS"),
    ("TC-MSA-GRR-005", "No arguments → usage message",                       "Exit ≠ 0, 'Usage' in output",    "PASS"),
    ("TC-MSA-GRR-006", "File not found",                                     "Exit ≠ 0, 'not found' in output","PASS"),
    ("TC-MSA-GRR-007", "Missing column → column named",                      "Exit ≠ 0, column name in output","PASS"),
    ("TC-MSA-GRR-008", "Unbalanced design",                                  "Exit ≠ 0, 'unbalanced' in output","PASS"),
    ("TC-MSA-GRR-009", "Single operator",                                    "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-GRR-010", "Bypass protection — direct Rscript call",            "Exit ≠ 0, 'RENV_PATHS_ROOT'",    "PASS"),
    # --- Nested GRR ---
    ("TC-MSA-NGR-001", "Happy path — exit 0, sections present",             "Exit 0, sections present",        "PASS"),
    ("TC-MSA-NGR-002", "Known good data — %GRR < 30%, ACCEPTABLE/MARGINAL", "%GRR < 30%",                      "PASS"),
    ("TC-MSA-NGR-003", "Known poor data — verdict UNACCEPTABLE",            "UNACCEPTABLE in output",          "PASS"),
    ("TC-MSA-NGR-004", "--tolerance flag",                                   "Exit 0, tolerance output",        "PASS"),
    ("TC-MSA-NGR-005", "PNG written to ~/Downloads/",                        "PNG present, recent mtime",       "PASS"),
    ("TC-MSA-NGR-006", "No arguments → usage message",                       "Exit ≠ 0, 'Usage' in output",    "PASS"),
    ("TC-MSA-NGR-007", "File not found",                                     "Exit ≠ 0, 'not found' in output","PASS"),
    ("TC-MSA-NGR-008", "Missing 'replicate' column",                         "Exit ≠ 0, 'replicate' in output","PASS"),
    ("TC-MSA-NGR-009", "Single operator",                                    "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-NGR-010", "Unbalanced design",                                  "Exit ≠ 0, 'unbalanced' in output","PASS"),
    ("TC-MSA-NGR-011", "Bypass protection — direct Rscript call",            "Exit ≠ 0, 'RENV_PATHS_ROOT'",    "PASS"),
    # --- Linearity & Bias ---
    ("TC-MSA-LB-001",  "Happy path — exit 0, sections present",             "Exit 0, sections present",        "PASS"),
    ("TC-MSA-LB-002",  "Known data — slope significant",                    "Slope in [0.01, 0.15]",           "PASS"),
    ("TC-MSA-LB-003",  "--tolerance flag",                                   "Exit 0, tolerance output",        "PASS"),
    ("TC-MSA-LB-004",  "PNG written to ~/Downloads/",                        "PNG present, recent mtime",       "PASS"),
    ("TC-MSA-LB-005",  "No arguments → usage message",                       "Exit ≠ 0, 'Usage' in output",    "PASS"),
    ("TC-MSA-LB-006",  "File not found",                                     "Exit ≠ 0, 'not found' in output","PASS"),
    ("TC-MSA-LB-007",  "Missing 'reference' column",                         "Exit ≠ 0, column named",          "PASS"),
    ("TC-MSA-LB-008",  "Single reference level",                             "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-LB-009",  "Inconsistent reference per part",                   "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-LB-010",  "Bypass protection — direct Rscript call",            "Exit ≠ 0, 'RENV_PATHS_ROOT'",    "PASS"),
    # --- Type 1 ---
    ("TC-MSA-T1-001",  "Happy path — exit 0, Cg/Cgk present",              "Exit 0, sections present",        "PASS"),
    ("TC-MSA-T1-002",  "Known good data — Cg ≥ 1.33, CAPABLE",             "Cg ≥ 1.33, Cgk ≥ 1.33, CAPABLE", "PASS"),
    ("TC-MSA-T1-003",  "Biased data — Cgk < Cg",                           "Cgk < Cg",                        "PASS"),
    ("TC-MSA-T1-004",  "PNG written to ~/Downloads/",                        "PNG present, recent mtime",       "PASS"),
    ("TC-MSA-T1-005",  "No arguments → usage message",                       "Exit ≠ 0, 'Usage' in output",    "PASS"),
    ("TC-MSA-T1-006",  "--reference missing",                                "Exit ≠ 0, 'reference' in output","PASS"),
    ("TC-MSA-T1-007",  "--tolerance missing",                                "Exit ≠ 0, 'tolerance' in output","PASS"),
    ("TC-MSA-T1-008",  "File not found",                                     "Exit ≠ 0, 'not found' in output","PASS"),
    ("TC-MSA-T1-009",  "Missing 'value' column",                             "Exit ≠ 0, 'value' in output",    "PASS"),
    ("TC-MSA-T1-010",  "Too few measurements (< 10)",                        "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-T1-011",  "Bypass protection — direct Rscript call",            "Exit ≠ 0, 'RENV_PATHS_ROOT'",    "PASS"),
    # --- Attribute ---
    ("TC-MSA-ATT-001", "Dataset with reference — exit 0, all sections",     "Exit 0, all sections present",    "PASS"),
    ("TC-MSA-ATT-002", "Dataset without reference — no vs-ref section",     "Exit 0, 'Vs Reference' absent",   "PASS"),
    ("TC-MSA-ATT-003", "Known data — Fleiss Kappa in [0.7, 1.0], verdict",  "Kappa ∈ [0.7,1.0], verdict",     "PASS"),
    ("TC-MSA-ATT-004", "Perfect appraiser A vs reference — Kappa = 1.0",    "Kappa = 1.0000 for appraiser A", "PASS"),
    ("TC-MSA-ATT-005", "PNG written to ~/Downloads/",                        "PNG present, recent mtime",       "PASS"),
    ("TC-MSA-ATT-006", "No arguments → usage message",                       "Exit ≠ 0, 'Usage' in output",    "PASS"),
    ("TC-MSA-ATT-007", "File not found",                                     "Exit ≠ 0, 'not found' in output","PASS"),
    ("TC-MSA-ATT-008", "Missing 'trial' column",                             "Exit ≠ 0, 'trial' in output",    "PASS"),
    ("TC-MSA-ATT-009", "Only one appraiser",                                 "Exit ≠ 0",                        "PASS"),
    ("TC-MSA-ATT-010", "Unbalanced design",                                  "Exit ≠ 0, 'unbalanced' in output","PASS"),
    ("TC-MSA-ATT-011", "Bypass protection — direct Rscript call",            "Exit ≠ 0, 'RENV_PATHS_ROOT'",    "PASS"),
]
for i, row in enumerate(tc_rows):
    data_row(tbl_tc, *row, alt=(i % 2 == 0))
set_col_widths(tbl_tc, [3.5, 6, 4.5, 3])

# ===========================================================================
# 7. DEVIATIONS
# ===========================================================================

h1(doc, "7.  Deviations from OQ Plan (JR-VP-MSA-001 v1.0)")
para(doc, "No deviations from the approved OQ plan were identified during this OQ execution.")
doc.add_paragraph()
tbl_dev = doc.add_table(rows=1, cols=4)
tbl_dev.style = "Table Grid"
header_row(tbl_dev, "ID", "Description", "Resolution", "Status")
data_row(tbl_dev, "—", "None", "—", "N/A", alt=True)
set_col_widths(tbl_dev, [2, 7, 6, 2])

# ===========================================================================
# 8. REQUIREMENTS TRACEABILITY
# ===========================================================================

h1(doc, "8.  Requirements Traceability Matrix")
para(doc,
    "The table below maps each User Requirement from JR-VP-MSA-001 to the "
    "test cases executed and their result.")
doc.add_paragraph()
tbl_rtm = doc.add_table(rows=1, cols=4)
tbl_rtm.style = "Table Grid"
header_row(tbl_rtm, "UR", "Requirement (summary)", "Test Cases", "Result")
rtm = [
    ("UR-MSA-001", "jrc_msa_gauge_rr — ANOVA, variance components, %GRR",
     "TC-MSA-GRR-001..003", "3 / 3  PASS"),
    ("UR-MSA-002", "Gauge R&R classification (ACCEPTABLE/MARGINAL/UNACCEPTABLE)",
     "TC-MSA-GRR-002", "1 / 1  PASS"),
    ("UR-MSA-003", "Gauge R&R four-panel PNG to ~/Downloads/",
     "TC-MSA-GRR-004", "1 / 1  PASS"),
    ("UR-MSA-004", "Gauge R&R — reject invalid inputs",
     "TC-MSA-GRR-007..009", "3 / 3  PASS"),
    ("UR-MSA-005", "jrc_msa_nested_grr — nested ANOVA, %GRR",
     "TC-MSA-NGR-001, 002, 004", "3 / 3  PASS"),
    ("UR-MSA-006", "Nested GRR classification",
     "TC-MSA-NGR-002..003", "2 / 2  PASS"),
    ("UR-MSA-007", "Nested GRR two-panel PNG to ~/Downloads/",
     "TC-MSA-NGR-005", "1 / 1  PASS"),
    ("UR-MSA-008", "Nested GRR — reject invalid inputs",
     "TC-MSA-NGR-008..010", "3 / 3  PASS"),
    ("UR-MSA-009", "jrc_msa_linearity_bias — regression, slope, %Linearity",
     "TC-MSA-LB-001..003", "3 / 3  PASS"),
    ("UR-MSA-010", "Linearity/bias verdict",
     "TC-MSA-LB-002, 004", "2 / 2  PASS"),
    ("UR-MSA-011", "Linearity/bias two-panel PNG to ~/Downloads/",
     "TC-MSA-LB-004", "1 / 1  PASS"),
    ("UR-MSA-012", "Linearity/bias — reject invalid inputs",
     "TC-MSA-LB-007..009", "3 / 3  PASS"),
    ("UR-MSA-013", "jrc_msa_type1 — Cg, Cgk, bias t-test",
     "TC-MSA-T1-001..003", "3 / 3  PASS"),
    ("UR-MSA-014", "Type 1 verdict; Cgk < Cg when bias present",
     "TC-MSA-T1-002..003", "2 / 2  PASS"),
    ("UR-MSA-015", "Type 1 two-panel PNG to ~/Downloads/",
     "TC-MSA-T1-004", "1 / 1  PASS"),
    ("UR-MSA-016", "Type 1 — mandatory args; reject < 10 measurements",
     "TC-MSA-T1-006..010", "5 / 5  PASS"),
    ("UR-MSA-017", "jrc_msa_attribute — within/between Kappa",
     "TC-MSA-ATT-001..003", "3 / 3  PASS"),
    ("UR-MSA-018", "Attribute — vs-reference Kappa when column present",
     "TC-MSA-ATT-001, 004", "2 / 2  PASS"),
    ("UR-MSA-019", "Attribute Kappa classification and verdict",
     "TC-MSA-ATT-003", "1 / 1  PASS"),
    ("UR-MSA-020", "Attribute two-panel PNG to ~/Downloads/",
     "TC-MSA-ATT-005", "1 / 1  PASS"),
    ("UR-MSA-021", "Attribute — reject invalid inputs",
     "TC-MSA-ATT-008..010", "3 / 3  PASS"),
    ("UR-MSA-022", "All scripts — RENV_PATHS_ROOT bypass protection",
     "TC-MSA-GRR-010, NGR-011, LB-010, T1-011, ATT-011", "5 / 5  PASS"),
    ("UR-MSA-023", "All scripts — no-arguments usage message",
     "TC-MSA-GRR-005, NGR-006, LB-005, T1-005, ATT-006", "5 / 5  PASS"),
    ("UR-MSA-024", "All scripts — file not found error",
     "TC-MSA-GRR-006, NGR-007, LB-006, T1-008, ATT-007", "5 / 5  PASS"),
]
for i, row in enumerate(rtm):
    data_row(tbl_rtm, *row, alt=(i % 2 == 0))
set_col_widths(tbl_rtm, [3, 7, 5, 2])

# ===========================================================================
# 9. OQ CONCLUSION
# ===========================================================================

h1(doc, "9.  OQ Conclusion")
para(doc,
    "The Operational Qualification of the JR Validated Environment MSA Module "
    "v1.0 is complete. All 53 test cases defined in JR-VP-MSA-001 were executed "
    "on 2026-03-18 and passed with no failures, no errors, and no deviations.")
doc.add_paragraph()
para(doc,
    "All 24 user requirements (UR-MSA-001 through UR-MSA-024) are covered by at "
    "least one passing test case. The OQ acceptance criteria stated in "
    "JR-VP-MSA-001 Section 10 are fully met.")
doc.add_paragraph()
para(doc,
    "The MSA module is hereby declared OPERATIONALLY QUALIFIED and released for "
    "use in design verification activities in accordance with the JR Validated "
    "Environment quality system.", bold=True)
doc.add_paragraph()
para(doc,
    "OQ Evidence file: "
    "~/.jrscript/MyProject/validation/msa_oq_execution_20260318T150625.txt",
    italic=True)

# ===========================================================================
# Save
# ===========================================================================

doc.save(OUT)
print(f"Saved: {OUT}")
