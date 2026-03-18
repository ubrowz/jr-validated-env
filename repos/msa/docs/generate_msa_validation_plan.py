"""
Generates msa_validation_plan.docx — FDA-acceptable OQ Validation Plan
for the JR Validated Environment MSA module v1.0.

Run from the repo root:
    python3 repos/msa/docs/generate_msa_validation_plan.py
Output: repos/msa/docs/msa_validation_plan.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "msa_validation_plan.docx")

# ---------------------------------------------------------------------------
# Helpers (same style conventions as the core OQ plan generator)
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text="", bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def shade_cell(cell, hex_fill, font_color=None):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)
    if font_color:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = font_color


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, "1A1A2E", RGBColor(0xFF, 0xFF, 0xFF))
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            shade_cell(cell, bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_tc_block(doc, tc_id, tests_ur, command, rationale, pass_criterion):
    """Add a formatted test case block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tc_id)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x21, 0x66, 0xAC)

    def labeled_line(label, value):
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Cm(0.8)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        r1 = lp.add_run(label + ": ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = lp.add_run(value)
        r2.font.size = Pt(9)
        if label == "Command":
            r2.font.name = "Courier New"

    labeled_line("Tests", tests_ur)
    labeled_line("Command", command)
    if rationale:
        labeled_line("Rationale", rationale)
    labeled_line("Pass criterion", pass_criterion)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ---------------------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------------------

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("VALIDATION PLAN\nOPERATIONAL QUALIFICATION\nMSA MODULE")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("JR Validated Environment — MSA Module v1.0")
r2.font.size = Pt(13)
r2.bold = True

doc.add_paragraph()

add_table(doc,
    ["Field", "Value"],
    [
        ("Document Number",  "JR-VP-MSA-001"),
        ("Title",            "Validation Plan — Operational Qualification, MSA Module"),
        ("System",           "JR Validated Environment — MSA Module"),
        ("Module Version",   "1.0"),
        ("Document Version", "1.0"),
        ("Status",           "Draft"),
        ("Effective Date",   "2026-03-18"),
        ("Author",           "Joep Rous"),
        ("Reviewer",         "[Name]"),
        ("Approver",         "[Name]"),
    ],
    col_widths=[5, 11]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# VERSION HISTORY
# ---------------------------------------------------------------------------

add_heading(doc, "Version History", level=1)
add_table(doc,
    ["Version", "Date", "Author", "Description"],
    [("1.0", "2026-03-18", "Joep Rous",
      "Initial release. Covers all 5 MSA scripts in module v1.0.")],
    col_widths=[2, 3, 4, 8]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# APPROVAL SIGNATURES
# ---------------------------------------------------------------------------

add_heading(doc, "Approval Signatures", level=1)
add_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ("Author",        "", "", ""),
        ("Reviewer",      "", "", ""),
        ("Approver (QA)", "", "", ""),
    ],
    col_widths=[4, 5, 5, 3]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. PURPOSE
# ---------------------------------------------------------------------------

add_heading(doc, "1.  Purpose", level=1)
add_para(doc,
    "This document defines the Validation Plan for the Operational Qualification (OQ) "
    "of the MSA (Measurement System Analysis) module included in version 1.0 of the "
    "JR Validated Environment. It specifies the user requirements, test cases, "
    "acceptance criteria, and traceability between requirements and tests needed to "
    "demonstrate that each MSA script performs correctly for its intended use.")
add_para(doc,
    "Execution of the tests defined in this plan, together with satisfactory results "
    "and documented evidence, constitutes the OQ for the MSA module.")

# ---------------------------------------------------------------------------
# 2. SCOPE
# ---------------------------------------------------------------------------

add_heading(doc, "2.  Scope", level=1)
add_heading(doc, "2.1  In Scope", level=2)

in_scope = [
    "All 5 R scripts in the repos/msa/R/ directory",
    "Correct computation of outputs for valid inputs (Kappa, variance components, %GRR, Cg/Cgk, bias/linearity)",
    "Correct rejection of invalid inputs with informative error messages",
    "Correct file I/O behaviour (CSV input, PNG output to ~/Downloads/)",
    "Verification that all R scripts enforce execution within the jrrun validated "
    "environment by checking RENV_PATHS_ROOT at startup",
    "Balanced-design validation for crossed and nested Gauge R&R",
]
for item in in_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

add_heading(doc, "2.2  Out of Scope", level=2)
out_scope = [
    "Infrastructure scripts in bin/ and admin/ (covered by IQ; see JR-VP-001)",
    "Demo scripts jrc_msa_R_hello.R and jrc_msa_py_hello.py",
    "Performance qualification (PQ) — end-to-end studies using production data "
    "are conducted separately under individual design verification protocols",
    "User acceptance testing (UAT)",
    "Validation of R language interpreter and third-party packages (ggplot2, grid) — "
    "package integrity is confirmed at installation via SHA256 (JR-IQ-001)",
    "Statistical validation of the underlying methods (Kappa, ANOVA, Weibull) — "
    "method references are cited in Section 3",
]
for item in out_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 3. REGULATORY BASIS
# ---------------------------------------------------------------------------

add_heading(doc, "3.  Regulatory Basis and References", level=1)
add_heading(doc, "3.1  Regulations", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("21 CFR Part 820, §820.70(i)", "Quality System Regulation — Automated Data Processing"),
        ("21 CFR Part 11",              "Electronic Records; Electronic Signatures"),
    ],
    col_widths=[5, 12]
)

doc.add_paragraph()
add_heading(doc, "3.2  FDA Guidance", level=2)
add_table(doc,
    ["Reference", "Title", "Date"],
    [
        ("FDA GPSV",
         "General Principles of Software Validation; Final Guidance for Industry and FDA Staff",
         "January 2002"),
        ("FDA CSA",
         "Computer Software Assurance for Production and Quality System Software; Draft Guidance",
         "September 2022"),
    ],
    col_widths=[3, 11, 3]
)

doc.add_paragraph()
add_heading(doc, "3.3  Standards and Methods", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("AIAG MSA, 4th Ed.",
         "Measurement Systems Analysis Reference Manual, Automotive Industry Action Group, 2010"),
        ("ISO 22514-7:2021",
         "Statistical methods in process management — Capability and performance — Part 7: "
         "Capability of measurement processes"),
        ("ISO 13485:2016",
         "Medical devices — Quality management systems — Requirements for regulatory purposes"),
        ("Fleiss (1971)",
         "Fleiss JL. Measuring nominal scale agreement among many raters. "
         "Psychological Bulletin, 76(5), 378–382."),
        ("Cohen (1960)",
         "Cohen J. A coefficient of agreement for nominal scales. "
         "Educational and Psychological Measurement, 20(1), 37–46."),
    ],
    col_widths=[4, 13]
)

doc.add_paragraph()
add_heading(doc, "3.4  Internal Documents", level=2)
add_table(doc,
    ["Document Number", "Title"],
    [
        ("JR-VP-001", "Validation Plan — Installation Qualification, JR Validated Environment v1.0.0"),
        ("JR-IQ-001", "IQ Execution Evidence (docs/IQ_validation_20260311_205146.txt)"),
        ("JR-VP-002", "Validation Plan — OQ, Community Script Suite v1.1.0"),
    ],
    col_widths=[4, 13]
)

# ---------------------------------------------------------------------------
# 4. DEFINITIONS
# ---------------------------------------------------------------------------

add_heading(doc, "4.  Definitions and Abbreviations", level=1)
add_table(doc,
    ["Term", "Definition"],
    [
        ("MSA",     "Measurement System Analysis — a body of techniques for characterising the "
                    "ability of a measurement system to produce accurate and repeatable results"),
        ("GRR",     "Gauge Repeatability and Reproducibility — a study to quantify the variation "
                    "introduced by the measurement system relative to total process variation"),
        ("%GRR",    "Percentage of total study variation attributable to the measurement system; "
                    "< 10% Acceptable, 10–30% Marginal, > 30% Unacceptable"),
        ("EV",      "Equipment Variation (Repeatability) — variation due to the gauge itself"),
        ("AV",      "Appraiser Variation (Reproducibility) — variation due to operators"),
        ("PV",      "Part Variation — true part-to-part variation"),
        ("ndc",     "Number of Distinct Categories — how many groups the gauge can reliably "
                    "distinguish; ≥ 5 is required"),
        ("Cg",      "Gauge Capability index (Type 1 study) — ratio of tolerance to measurement "
                    "spread; ≥ 1.33 required"),
        ("Cgk",     "Gauge Capability index adjusted for bias (Type 1 study); ≥ 1.33 required"),
        ("Kappa",   "Cohen's/Fleiss' Kappa — a statistic measuring inter-rater agreement beyond "
                    "what is expected by chance; 0 = chance, 1 = perfect agreement"),
        ("IQ",      "Installation Qualification"),
        ("OQ",      "Operational Qualification"),
        ("PQ",      "Performance Qualification"),
        ("UR",      "User Requirement"),
        ("TC",      "Test Case"),
        ("jrrun",   "The validated zsh entry-point script that sets the controlled environment "
                    "before executing any community script"),
        ("CSV",     "Comma-separated values input file"),
        ("PNG",     "Portable Network Graphics output file; plots saved to ~/Downloads/"),
    ],
    col_widths=[2.5, 14.5]
)

# ---------------------------------------------------------------------------
# 5. SYSTEM DESCRIPTION
# ---------------------------------------------------------------------------

add_heading(doc, "5.  System Description", level=1)
add_heading(doc, "5.1  Module Overview", level=2)
add_para(doc,
    "The MSA module is an extension of the JR Validated Environment providing five "
    "R scripts for Measurement System Analysis. Scripts are installed under repos/msa/ "
    "and are invoked via wrapper commands (e.g., jrc_msa_gauge_rr). The module shares "
    "the core environment infrastructure (jrrun, renv library, integrity checking, "
    "logging) with the community script suite.")

add_heading(doc, "5.2  Scripts Under Test", level=2)
add_table(doc,
    ["Script", "Method", "Input Columns", "Key Outputs"],
    [
        ("jrc_msa_gauge_rr",
         "Two-way ANOVA (crossed)",
         "part, operator, value",
         "%GRR, ndc, variance components, 4-panel PNG"),
        ("jrc_msa_nested_grr",
         "One-way nested ANOVA",
         "operator, part, replicate, value",
         "%GRR, variance components, 2-panel PNG"),
        ("jrc_msa_linearity_bias",
         "Linear regression",
         "part, reference, value",
         "Slope, intercept, %Linearity, per-part bias, 2-panel PNG"),
        ("jrc_msa_type1",
         "Type 1 gauge study",
         "value (optional: id)",
         "Cg, Cgk, bias t-test, 2-panel PNG"),
        ("jrc_msa_attribute",
         "Cohen's / Fleiss' Kappa",
         "part, appraiser, trial, rating (optional: reference)",
         "Within/between Kappa, % agreement, 2-panel PNG"),
    ],
    col_widths=[4, 3.5, 4.5, 5]
)

doc.add_paragraph()
add_heading(doc, "5.3  Intended Use", level=2)
add_para(doc,
    "MSA scripts are intended for use by engineers and scientists qualifying measurement "
    "systems used in medical device design verification. Intended uses include:")
uses = [
    "Quantifying the repeatability and reproducibility of a measurement gauge before "
    "using it in a design verification study",
    "Determining whether a gauge is capable of resolving real part-to-part differences",
    "Assessing gauge linearity and bias across its operating range",
    "Evaluating the capability of a gauge to measure within a specification tolerance",
    "Qualifying attribute inspection systems (pass/fail, categorical)",
    "Supporting measurement system qualification records in the Design History File (DHF)",
]
for u in uses:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(u).font.size = Pt(10)

add_heading(doc, "5.4  Hardware and Software Environment", level=2)
add_table(doc,
    ["Component", "Specification"],
    [
        ("Operating system", "macOS (arm64 / aarch64-apple-darwin)"),
        ("R version",        "4.5 (as specified in admin/r_version.txt)"),
        ("R packages",       "ggplot2, grid (versions as locked in admin/renv.lock)"),
        ("Entry point",      "bin/jrrun"),
        ("Module location",  "repos/msa/"),
        ("Test runner",      "pytest in ~/.venvs/MyProject_oq/ (reused from core OQ venv)"),
    ],
    col_widths=[5, 12]
)

# ---------------------------------------------------------------------------
# 6. RISK ASSESSMENT
# ---------------------------------------------------------------------------

add_heading(doc, "6.  Risk Assessment and Validation Level", level=1)
add_heading(doc, "6.1  Risk Classification", level=2)
add_para(doc,
    "MSA scripts are used to determine whether a measurement system is fit for purpose "
    "in a design verification study. Incorrect results could lead to accepting a "
    "measurement system with unacceptable gauge error, leading to unreliable "
    "verification data, incorrect pass/fail conclusions, and potentially unsafe "
    "medical devices reaching the market. This represents a high-consequence failure "
    "mode. Per the FDA General Principles of Software Validation (GPSV 2002), the "
    "level of validation effort is commensurate with this risk.")

add_heading(doc, "6.2  Risk Factors", level=2)
add_table(doc,
    ["Risk Factor", "Assessment", "Consequence"],
    [
        ("Incorrect variance component calculation (%GRR)",
         "Medium probability",
         "High: accept a poor gauge for use in DHF testing"),
        ("Incorrect Kappa calculation",
         "Medium probability",
         "High: accept an unreliable inspection system"),
        ("Design balance validation bypassed",
         "Low probability",
         "High: biased estimates from unbalanced data"),
        ("Silent failure (non-zero exit, no message)",
         "Low probability",
         "Medium: user unaware of failure"),
        ("Script runs outside jrrun (wrong library)",
         "Low probability",
         "High: wrong package versions, undetected results"),
    ],
    col_widths=[6, 4, 7]
)
doc.add_paragraph()
add_para(doc,
    "Strategy: Full OQ testing of all 5 scripts, covering both primary (happy-path) "
    "execution paths and boundary/error-handling paths. Numerical outputs are verified "
    "against independently determined expected ranges using known test datasets.")

# ---------------------------------------------------------------------------
# 7. ROLES AND RESPONSIBILITIES
# ---------------------------------------------------------------------------

add_heading(doc, "7.  Roles and Responsibilities", level=1)
add_table(doc,
    ["Role", "Responsibility"],
    [
        ("Author",
         "Writes and maintains this Validation Plan and the MSA module scripts"),
        ("Reviewer",
         "Reviews test specifications for completeness, accuracy, and regulatory adequacy"),
        ("QA Approver",
         "Approves this plan prior to OQ execution"),
        ("Test Executor",
         "Executes the OQ test suite (admin_msa_oq) and documents results"),
        ("QA Witness",
         "Reviews execution evidence and signs the OQ execution report"),
    ],
    col_widths=[4, 13]
)
doc.add_paragraph()
add_para(doc,
    "The Test Executor and Author may be the same person. The QA Approver must be "
    "independent of the Author for this document.", italic=True)

# ---------------------------------------------------------------------------
# 8. VALIDATION APPROACH
# ---------------------------------------------------------------------------

add_heading(doc, "8.  Validation Approach", level=1)
add_heading(doc, "8.1  Overall Strategy", level=2)
add_para(doc,
    "The MSA module is validated as an extension of the core JR Validated Environment "
    "using the same IQ/OQ/PQ model. The core environment IQ is documented in JR-IQ-001. "
    "This document covers the OQ for the MSA module only.")
add_table(doc,
    ["Phase", "Description", "Status"],
    [
        ("IQ",
         "Core environment: R version, packages, jrrun, integrity checking. "
         "MSA module scripts are included in the project_integrity.sha256 file.",
         "Completed — JR-IQ-001"),
        ("OQ (this document)",
         "Verifies correct operation of each MSA script across the range of "
         "anticipated inputs",
         "Planned"),
        ("PQ",
         "End-to-end performance with real measurement data; executed under "
         "individual design verification protocols",
         "Out of scope here"),
    ],
    col_widths=[2, 12, 3]
)

doc.add_paragraph()
add_heading(doc, "8.2  OQ Execution", level=2)
add_para(doc,
    "OQ tests are executed by the automated test runner repos/msa/admin_msa_oq, which:")
steps = [
    "Activates the OQ Python virtual environment (~/.venvs/MyProject_oq/)",
    "Invokes pytest against all test_*.py files in repos/msa/oq/",
    "For each test case, calls the MSA script as a subprocess via jrrun",
    "Captures exit code, stdout, and stderr",
    "Evaluates pass/fail against the criteria defined in Section 13",
    "Writes a timestamped qualification report to ~/.jrscript/MyProject/validation/",
]
for s in steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)
add_para(doc,
    "All test cases must pass for the OQ to be considered successful. A single failing "
    "test case constitutes an OQ failure and must be resolved before the module is "
    "released to users.")

add_heading(doc, "8.3  Sequence of Execution", level=2)
seq = [
    "Confirm test data files in repos/msa/oq/data/ are present and committed",
    "Run admin/admin_create_hash to regenerate project integrity file",
    "Obtain QA approval of this plan",
    "Execute repos/msa/admin_msa_oq on the target machine",
    "Review the qualification report",
    "Document any deviations",
    "Obtain QA sign-off on the execution evidence",
]
for s in seq:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 9. USER REQUIREMENTS
# ---------------------------------------------------------------------------

add_heading(doc, "9.  User Requirements", level=1)
add_para(doc,
    "The following user requirements define what the MSA module must do from the "
    "user's perspective. Each requirement is traceable to one or more OQ test cases "
    "in the Requirements Traceability Matrix (Section 14).")

add_heading(doc, "9.1  Gauge R&R (Crossed)", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-001",
         "The system shall perform a Gauge R&R analysis using the two-way ANOVA method "
         "(crossed design) on a balanced part/operator/value dataset, computing: "
         "ANOVA table, variance components, repeatability (EV), reproducibility (AV), "
         "%GRR vs study variation, %GRR vs tolerance (when supplied), and number of "
         "distinct categories (ndc)."),
        ("UR-MSA-002",
         "The system shall classify the measurement system as ACCEPTABLE (%GRR < 10%), "
         "MARGINAL (10% ≤ %GRR < 30%), or UNACCEPTABLE (%GRR ≥ 30%) and display the verdict."),
        ("UR-MSA-003",
         "The system shall save a four-panel PNG (components of variation, by-part, "
         "by-operator, interaction plot) to ~/Downloads/."),
        ("UR-MSA-004",
         "The system shall reject datasets that are unbalanced, have fewer than 2 operators "
         "or parts, or are missing required columns, with an informative error message."),
    ],
    col_widths=[3, 14]
)

doc.add_paragraph()
add_heading(doc, "9.2  Nested Gauge R&R (Destructive)", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-005",
         "The system shall perform a Nested Gauge R&R analysis on a balanced nested dataset "
         "(parts nested within operators), computing: ANOVA table, variance components, "
         "repeatability (EV), reproducibility (AV), part-within-operator variation, "
         "%GRR vs study variation, and %GRR vs tolerance (when supplied)."),
        ("UR-MSA-006",
         "The system shall classify the nested measurement system as ACCEPTABLE, MARGINAL, "
         "or UNACCEPTABLE using the same %GRR thresholds as UR-MSA-002."),
        ("UR-MSA-007",
         "The system shall save a two-panel PNG (components of variation, measurements by "
         "operator) to ~/Downloads/."),
        ("UR-MSA-008",
         "The system shall reject unbalanced nested datasets or datasets missing required "
         "columns, with an informative error message."),
    ],
    col_widths=[3, 14]
)

doc.add_paragraph()
add_heading(doc, "9.3  Linearity and Bias", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-009",
         "The system shall perform a linearity and bias analysis by regressing observed "
         "bias (measured value − reference) against reference value, reporting: "
         "regression slope and intercept with p-values, R², %Linearity, and per-part "
         "bias with 95% CI and significance flag."),
        ("UR-MSA-010",
         "The system shall classify the gauge as having ACCEPTABLE or UNACCEPTABLE "
         "linearity and bias based on slope significance (p < 0.05) and per-part bias "
         "significance."),
        ("UR-MSA-011",
         "The system shall save a two-panel PNG (linearity regression plot with 95% CI "
         "band, per-part bias bar chart) to ~/Downloads/."),
        ("UR-MSA-012",
         "The system shall reject datasets with fewer than 2 distinct reference values, "
         "inconsistent reference values per part, or missing required columns."),
    ],
    col_widths=[3, 14]
)

doc.add_paragraph()
add_heading(doc, "9.4  Type 1 Gauge Study", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-013",
         "The system shall perform a Type 1 gauge capability study on repeated measurements "
         "of a single reference part, computing: Cg = 0.2 × T / (6σ), "
         "Cgk = (0.1 × T − |bias|) / (3σ), bias, SD, and a t-test for bias significance."),
        ("UR-MSA-014",
         "The system shall classify the gauge as CAPABLE (Cg ≥ 1.33 and Cgk ≥ 1.33) "
         "or INCAPABLE, and indicate when Cgk < Cg (bias present)."),
        ("UR-MSA-015",
         "The system shall save a two-panel PNG (run chart with acceptance band, "
         "histogram with normal curve) to ~/Downloads/."),
        ("UR-MSA-016",
         "The system shall require both --reference and --tolerance arguments and reject "
         "datasets with fewer than 10 measurements."),
    ],
    col_widths=[3, 14]
)

doc.add_paragraph()
add_heading(doc, "9.5  Attribute Agreement Analysis", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-017",
         "The system shall perform an attribute agreement analysis on a balanced "
         "part/appraiser/trial/rating dataset, computing: within-appraiser % agreement "
         "and Cohen's Kappa for each appraiser, and between-appraiser Fleiss' Kappa."),
        ("UR-MSA-018",
         "When a reference column is present, the system shall additionally compute "
         "each appraiser's Kappa versus the reference."),
        ("UR-MSA-019",
         "The system shall classify each Kappa as ACCEPTABLE (≥ 0.9), MARGINAL "
         "(0.7–0.9), or UNACCEPTABLE (< 0.7) and display the verdict."),
        ("UR-MSA-020",
         "The system shall save a two-panel PNG (% agreement bar chart, Kappa chart) "
         "to ~/Downloads/."),
        ("UR-MSA-021",
         "The system shall reject datasets that are unbalanced, have fewer than 2 "
         "appraisers, are missing required columns, or contain only one rating category."),
    ],
    col_widths=[3, 14]
)

doc.add_paragraph()
add_heading(doc, "9.6  Common Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-MSA-022",
         "Each MSA script shall verify that RENV_PATHS_ROOT is set at startup and "
         "terminate with a non-zero exit code and descriptive error if it is not set, "
         "preventing execution outside the jrrun validated environment."),
        ("UR-MSA-023",
         "Each MSA script shall terminate with a non-zero exit code and display a "
         "usage message when invoked with no arguments."),
        ("UR-MSA-024",
         "Each MSA script shall terminate with a non-zero exit code and display an "
         "informative error message when the input file is not found."),
    ],
    col_widths=[3, 14]
)

# ---------------------------------------------------------------------------
# 10. ACCEPTANCE CRITERIA
# ---------------------------------------------------------------------------

add_heading(doc, "10.  Acceptance Criteria", level=1)
add_para(doc,
    "The OQ is considered successful when all of the following criteria are met:")
criteria = [
    "All 53 automated test cases defined in Section 13 pass with exit code 0 from pytest",
    "The qualification report generated by admin_msa_oq records 0 failures and 0 errors",
    "The qualification report is reviewed and signed by the QA Witness",
    "Any deviations are documented and dispositioned per Section 15",
]
for c in criteria:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(c).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 11. TEST DATA
# ---------------------------------------------------------------------------

add_heading(doc, "11.  Test Data", level=1)
add_para(doc,
    "All test data files are synthetic CSV files committed to repos/msa/oq/data/. "
    "They are version-controlled and included in the project integrity check. "
    "No real patient or production data is used.")

add_table(doc,
    ["File", "Used By", "Purpose"],
    [
        ("gauge_rr_balanced.csv",         "TC-MSA-GRR-001..004", "10 parts × 3 operators × 3 reps, low noise"),
        ("gauge_rr_missing_col.csv",       "TC-MSA-GRR-007",      "Missing 'operator' column"),
        ("gauge_rr_unbalanced.csv",        "TC-MSA-GRR-008",      "Unequal replicates per cell"),
        ("gauge_rr_one_operator.csv",      "TC-MSA-GRR-009",      "Single operator only"),
        ("nested_grr_good.csv",           "TC-MSA-NGR-001..005", "5 parts × 3 operators × 2 reps, low noise"),
        ("nested_grr_poor.csv",           "TC-MSA-NGR-003",      "High within-operator variation"),
        ("nested_grr_missing_col.csv",    "TC-MSA-NGR-008",      "Missing 'replicate' column"),
        ("nested_grr_one_operator.csv",   "TC-MSA-NGR-009",      "Single operator only"),
        ("nested_grr_unbalanced.csv",     "TC-MSA-NGR-010",      "Unequal parts per operator"),
        ("linearity_bias_good.csv",       "TC-MSA-LB-001..004",  "5 reference levels × 3 reps, linear bias"),
        ("linearity_bias_missing_col.csv","TC-MSA-LB-007",       "Missing 'reference' column"),
        ("linearity_bias_one_part.csv",   "TC-MSA-LB-008",       "Single reference level"),
        ("linearity_bias_inconsistent_ref.csv","TC-MSA-LB-009",  "Same part, different reference values"),
        ("type1_good.csv",                "TC-MSA-T1-001..004",  "25 measurements of reference part, low bias"),
        ("type1_biased.csv",              "TC-MSA-T1-003",       "25 measurements with systematic bias"),
        ("type1_missing_value_col.csv",   "TC-MSA-T1-009",       "Missing 'value' column"),
        ("type1_too_few.csv",             "TC-MSA-T1-010",       "Only 5 measurements (< 10 required)"),
        ("attribute_with_ref.csv",        "TC-MSA-ATT-001,003..005", "3 appraisers × 10 parts × 2 trials + reference"),
        ("attribute_no_ref.csv",          "TC-MSA-ATT-002",      "Same without reference column"),
        ("attribute_missing_col.csv",     "TC-MSA-ATT-008",      "Missing 'trial' column"),
        ("attribute_one_appraiser.csv",   "TC-MSA-ATT-009",      "Single appraiser only"),
        ("attribute_unbalanced.csv",      "TC-MSA-ATT-010",      "Unequal trials per appraiser-part cell"),
    ],
    col_widths=[5, 4, 8]
)

# ---------------------------------------------------------------------------
# 12. TRACEABILITY SUMMARY
# ---------------------------------------------------------------------------

add_heading(doc, "12.  Requirements Traceability Summary", level=1)
add_table(doc,
    ["Requirement", "Test Cases"],
    [
        ("UR-MSA-001", "TC-MSA-GRR-001, TC-MSA-GRR-002, TC-MSA-GRR-003"),
        ("UR-MSA-002", "TC-MSA-GRR-002"),
        ("UR-MSA-003", "TC-MSA-GRR-004"),
        ("UR-MSA-004", "TC-MSA-GRR-007, TC-MSA-GRR-008, TC-MSA-GRR-009"),
        ("UR-MSA-005", "TC-MSA-NGR-001, TC-MSA-NGR-002, TC-MSA-NGR-004"),
        ("UR-MSA-006", "TC-MSA-NGR-002, TC-MSA-NGR-003"),
        ("UR-MSA-007", "TC-MSA-NGR-005"),
        ("UR-MSA-008", "TC-MSA-NGR-008, TC-MSA-NGR-009, TC-MSA-NGR-010"),
        ("UR-MSA-009", "TC-MSA-LB-001, TC-MSA-LB-002, TC-MSA-LB-003"),
        ("UR-MSA-010", "TC-MSA-LB-002, TC-MSA-LB-004"),
        ("UR-MSA-011", "TC-MSA-LB-005"),
        ("UR-MSA-012", "TC-MSA-LB-007, TC-MSA-LB-008, TC-MSA-LB-009"),
        ("UR-MSA-013", "TC-MSA-T1-001, TC-MSA-T1-002, TC-MSA-T1-003"),
        ("UR-MSA-014", "TC-MSA-T1-002, TC-MSA-T1-003"),
        ("UR-MSA-015", "TC-MSA-T1-004"),
        ("UR-MSA-016", "TC-MSA-T1-006, TC-MSA-T1-007, TC-MSA-T1-010"),
        ("UR-MSA-017", "TC-MSA-ATT-001, TC-MSA-ATT-002, TC-MSA-ATT-003"),
        ("UR-MSA-018", "TC-MSA-ATT-001, TC-MSA-ATT-004"),
        ("UR-MSA-019", "TC-MSA-ATT-003"),
        ("UR-MSA-020", "TC-MSA-ATT-005"),
        ("UR-MSA-021", "TC-MSA-ATT-008, TC-MSA-ATT-009, TC-MSA-ATT-010"),
        ("UR-MSA-022", "TC-MSA-GRR-010, TC-MSA-NGR-011, TC-MSA-LB-010, TC-MSA-T1-011, TC-MSA-ATT-011"),
        ("UR-MSA-023", "TC-MSA-GRR-005, TC-MSA-NGR-006, TC-MSA-LB-006, TC-MSA-T1-005, TC-MSA-ATT-006"),
        ("UR-MSA-024", "TC-MSA-GRR-006, TC-MSA-NGR-007, TC-MSA-LB-007 (file), TC-MSA-T1-008, TC-MSA-ATT-007"),
    ],
    col_widths=[4, 13]
)

# ---------------------------------------------------------------------------
# 13. TEST CASES
# ---------------------------------------------------------------------------

add_heading(doc, "13.  Test Cases", level=1)
add_para(doc,
    "Each test case is executed automatically by pytest via repos/msa/admin_msa_oq. "
    "Commands shown use wrapper names; the test suite invokes scripts via jrrun.")

# --- 13.1 jrc_msa_gauge_rr ---
add_heading(doc, "13.1  jrc_msa_gauge_rr — Gauge R&R (Crossed)", level=2)

add_tc_block(doc,
    "TC-MSA-GRR-001  Happy path — exit 0, output sections present",
    "UR-MSA-001",
    "jrc_msa_gauge_rr gauge_rr_balanced.csv",
    None,
    "Exit code = 0. Output contains: 'ANOVA', 'Variance Components', 'Study Variation', "
    "'Verdict', 'Gauge R&R'."
)
add_tc_block(doc,
    "TC-MSA-GRR-002  Known data — %GRR in expected range, verdict ACCEPTABLE",
    "UR-MSA-001, UR-MSA-002",
    "jrc_msa_gauge_rr gauge_rr_balanced.csv",
    "Dataset has low gauge noise; expected %GRR < 10%.",
    "Exit code = 0. Parsed %GRR < 10%. Verdict = 'ACCEPTABLE'."
)
add_tc_block(doc,
    "TC-MSA-GRR-003  --tolerance flag",
    "UR-MSA-001",
    "jrc_msa_gauge_rr gauge_rr_balanced.csv --tolerance 0.5",
    None,
    "Exit code = 0. Output contains tolerance-referenced %GRR value."
)
add_tc_block(doc,
    "TC-MSA-GRR-004  PNG written to ~/Downloads/",
    "UR-MSA-003",
    "jrc_msa_gauge_rr gauge_rr_balanced.csv",
    None,
    "Exit code = 0. File matching *_jrc_msa_gauge_rr.png exists in ~/Downloads/ "
    "with modification time ≥ test start."
)
add_tc_block(doc,
    "TC-MSA-GRR-005  No arguments → usage message",
    "UR-MSA-023",
    "jrc_msa_gauge_rr",
    None,
    "Exit code ≠ 0. Output contains 'Usage' or 'usage'."
)
add_tc_block(doc,
    "TC-MSA-GRR-006  File not found",
    "UR-MSA-024",
    "jrc_msa_gauge_rr /tmp/no_such.csv",
    None,
    "Exit code ≠ 0. Output contains 'not found' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-GRR-007  Missing column → column named in error",
    "UR-MSA-004",
    "jrc_msa_gauge_rr gauge_rr_missing_col.csv",
    None,
    "Exit code ≠ 0. Output contains name of missing column."
)
add_tc_block(doc,
    "TC-MSA-GRR-008  Unbalanced design",
    "UR-MSA-004",
    "jrc_msa_gauge_rr gauge_rr_unbalanced.csv",
    None,
    "Exit code ≠ 0. Output contains 'unbalanced' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-GRR-009  Single operator",
    "UR-MSA-004",
    "jrc_msa_gauge_rr gauge_rr_one_operator.csv",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-GRR-010  Bypass protection — direct Rscript call",
    "UR-MSA-022",
    "Rscript repos/msa/R/jrc_msa_gauge_rr.R gauge_rr_balanced.csv  (RENV_PATHS_ROOT unset)",
    "Ensures script cannot be run outside jrrun.",
    "Exit code ≠ 0. Output contains 'RENV_PATHS_ROOT'."
)

# --- 13.2 jrc_msa_nested_grr ---
add_heading(doc, "13.2  jrc_msa_nested_grr — Nested Gauge R&R", level=2)

add_tc_block(doc,
    "TC-MSA-NGR-001  Happy path — exit 0, output sections present",
    "UR-MSA-005",
    "jrc_msa_nested_grr nested_grr_good.csv",
    None,
    "Exit code = 0. Output contains: 'ANOVA', 'Variance Components', 'Study Variation', "
    "'Verdict', 'Gauge R&R', 'Nested'."
)
add_tc_block(doc,
    "TC-MSA-NGR-002  Known good data — %GRR < 30%, verdict ACCEPTABLE or MARGINAL",
    "UR-MSA-005, UR-MSA-006",
    "jrc_msa_nested_grr nested_grr_good.csv",
    "Dataset has low gauge noise relative to part variation.",
    "Exit code = 0. Parsed %GRR < 30%. Verdict = 'ACCEPTABLE' or 'MARGINAL'."
)
add_tc_block(doc,
    "TC-MSA-NGR-003  Known poor data — verdict UNACCEPTABLE",
    "UR-MSA-006",
    "jrc_msa_nested_grr nested_grr_poor.csv",
    "Dataset has high within-operator replicate variation.",
    "Exit code = 0. Verdict = 'UNACCEPTABLE'."
)
add_tc_block(doc,
    "TC-MSA-NGR-004  --tolerance flag",
    "UR-MSA-005",
    "jrc_msa_nested_grr nested_grr_good.csv --tolerance 10.0",
    None,
    "Exit code = 0. Output contains tolerance-referenced output and a percentage value."
)
add_tc_block(doc,
    "TC-MSA-NGR-005  PNG written to ~/Downloads/",
    "UR-MSA-007",
    "jrc_msa_nested_grr nested_grr_good.csv",
    None,
    "Exit code = 0. File matching *_jrc_msa_nested_grr.png exists in ~/Downloads/ "
    "with modification time ≥ test start."
)
add_tc_block(doc,
    "TC-MSA-NGR-006  No arguments → usage message",
    "UR-MSA-023",
    "jrc_msa_nested_grr",
    None,
    "Exit code ≠ 0. Output contains 'Usage' or 'usage'."
)
add_tc_block(doc,
    "TC-MSA-NGR-007  File not found",
    "UR-MSA-024",
    "jrc_msa_nested_grr /tmp/no_such_nested.csv",
    None,
    "Exit code ≠ 0. Output contains 'not found' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-NGR-008  Missing 'replicate' column",
    "UR-MSA-008",
    "jrc_msa_nested_grr nested_grr_missing_col.csv",
    None,
    "Exit code ≠ 0. Output contains 'replicate' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-NGR-009  Single operator",
    "UR-MSA-008",
    "jrc_msa_nested_grr nested_grr_one_operator.csv",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-NGR-010  Unbalanced design — unequal parts per operator",
    "UR-MSA-008",
    "jrc_msa_nested_grr nested_grr_unbalanced.csv",
    None,
    "Exit code ≠ 0. Output contains 'unbalanced', 'parts', or 'equal' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-NGR-011  Bypass protection — direct Rscript call",
    "UR-MSA-022",
    "Rscript repos/msa/R/jrc_msa_nested_grr.R nested_grr_good.csv  (RENV_PATHS_ROOT unset)",
    None,
    "Exit code ≠ 0. Output contains 'RENV_PATHS_ROOT'."
)

# --- 13.3 jrc_msa_linearity_bias ---
add_heading(doc, "13.3  jrc_msa_linearity_bias — Linearity and Bias", level=2)

add_tc_block(doc,
    "TC-MSA-LB-001  Happy path — exit 0, output sections present",
    "UR-MSA-009",
    "jrc_msa_linearity_bias linearity_bias_good.csv",
    None,
    "Exit code = 0. Output contains: 'Linearity', 'Bias', 'Slope', 'Intercept', 'Verdict'."
)
add_tc_block(doc,
    "TC-MSA-LB-002  Known data — slope and %Linearity in expected range",
    "UR-MSA-009, UR-MSA-010",
    "jrc_msa_linearity_bias linearity_bias_good.csv",
    "Dataset has known linear bias; slope ≈ 0.05.",
    "Exit code = 0. Parsed slope in [0.01, 0.15]. %Linearity > 0."
)
add_tc_block(doc,
    "TC-MSA-LB-003  --tolerance flag",
    "UR-MSA-009",
    "jrc_msa_linearity_bias linearity_bias_good.csv --tolerance 5.0",
    None,
    "Exit code = 0. Output references tolerance."
)
add_tc_block(doc,
    "TC-MSA-LB-004  Verdict present",
    "UR-MSA-010",
    "jrc_msa_linearity_bias linearity_bias_good.csv",
    None,
    "Exit code = 0. Output contains 'ACCEPTABLE' or 'UNACCEPTABLE'."
)
add_tc_block(doc,
    "TC-MSA-LB-005  PNG written to ~/Downloads/",
    "UR-MSA-011",
    "jrc_msa_linearity_bias linearity_bias_good.csv",
    None,
    "File matching *_jrc_msa_linearity_bias.png exists in ~/Downloads/ "
    "with modification time ≥ test start."
)
add_tc_block(doc,
    "TC-MSA-LB-006  No arguments → usage message",
    "UR-MSA-023",
    "jrc_msa_linearity_bias",
    None,
    "Exit code ≠ 0. Output contains 'Usage' or 'usage'."
)
add_tc_block(doc,
    "TC-MSA-LB-007  File not found",
    "UR-MSA-024",
    "jrc_msa_linearity_bias /tmp/no_such_lb.csv",
    None,
    "Exit code ≠ 0. Output contains 'not found' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-LB-008  Missing column",
    "UR-MSA-012",
    "jrc_msa_linearity_bias linearity_bias_missing_col.csv",
    None,
    "Exit code ≠ 0. Output names missing column."
)
add_tc_block(doc,
    "TC-MSA-LB-009  Single reference level",
    "UR-MSA-012",
    "jrc_msa_linearity_bias linearity_bias_one_part.csv",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-LB-010  Inconsistent reference per part",
    "UR-MSA-012",
    "jrc_msa_linearity_bias linearity_bias_inconsistent_ref.csv",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-LB-011  Bypass protection — direct Rscript call",
    "UR-MSA-022",
    "Rscript repos/msa/R/jrc_msa_linearity_bias.R linearity_bias_good.csv  (RENV_PATHS_ROOT unset)",
    None,
    "Exit code ≠ 0. Output contains 'RENV_PATHS_ROOT'."
)

# Note: LB tests are 10 in current suite; TC-MSA-LB-011 = bypass; renumbered above

# --- 13.4 jrc_msa_type1 ---
add_heading(doc, "13.4  jrc_msa_type1 — Type 1 Gauge Study", level=2)

add_tc_block(doc,
    "TC-MSA-T1-001  Happy path — exit 0, output sections present",
    "UR-MSA-013",
    "jrc_msa_type1 type1_good.csv --reference 50.0 --tolerance 1.0",
    None,
    "Exit code = 0. Output contains: 'Cg', 'Cgk', 'Bias', 'Verdict'."
)
add_tc_block(doc,
    "TC-MSA-T1-002  Known good data — Cg ≥ 1.33, Cgk ≥ 1.33, verdict CAPABLE",
    "UR-MSA-013, UR-MSA-014",
    "jrc_msa_type1 type1_good.csv --reference 50.0 --tolerance 1.0",
    "Dataset: 25 measurements near reference, low noise.",
    "Exit code = 0. Parsed Cg ≥ 1.33. Parsed Cgk ≥ 1.33. Verdict = 'CAPABLE'."
)
add_tc_block(doc,
    "TC-MSA-T1-003  Biased data — Cgk < Cg",
    "UR-MSA-014",
    "jrc_msa_type1 type1_biased.csv --reference 50.0 --tolerance 1.0",
    "Dataset has systematic positive bias.",
    "Exit code = 0. Parsed Cgk < Cg."
)
add_tc_block(doc,
    "TC-MSA-T1-004  PNG written to ~/Downloads/",
    "UR-MSA-015",
    "jrc_msa_type1 type1_good.csv --reference 50.0 --tolerance 1.0",
    None,
    "File matching *_jrc_msa_type1.png exists in ~/Downloads/ "
    "with modification time ≥ test start."
)
add_tc_block(doc,
    "TC-MSA-T1-005  No arguments → usage message",
    "UR-MSA-023",
    "jrc_msa_type1",
    None,
    "Exit code ≠ 0. Output contains 'Usage' or 'usage'."
)
add_tc_block(doc,
    "TC-MSA-T1-006  --reference missing",
    "UR-MSA-016",
    "jrc_msa_type1 type1_good.csv --tolerance 1.0",
    None,
    "Exit code ≠ 0. Output contains 'reference' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-T1-007  --tolerance missing",
    "UR-MSA-016",
    "jrc_msa_type1 type1_good.csv --reference 50.0",
    None,
    "Exit code ≠ 0. Output contains 'tolerance' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-T1-008  File not found",
    "UR-MSA-024",
    "jrc_msa_type1 /tmp/no_such_t1.csv --reference 50.0 --tolerance 1.0",
    None,
    "Exit code ≠ 0. Output contains 'not found' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-T1-009  Missing 'value' column",
    "UR-MSA-016",
    "jrc_msa_type1 type1_missing_value_col.csv --reference 50.0 --tolerance 1.0",
    None,
    "Exit code ≠ 0. Output contains 'value' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-T1-010  Too few measurements (< 10)",
    "UR-MSA-016",
    "jrc_msa_type1 type1_too_few.csv --reference 50.0 --tolerance 1.0",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-T1-011  Bypass protection — direct Rscript call",
    "UR-MSA-022",
    "Rscript repos/msa/R/jrc_msa_type1.R type1_good.csv --reference 50.0 --tolerance 1.0  (RENV_PATHS_ROOT unset)",
    None,
    "Exit code ≠ 0. Output contains 'RENV_PATHS_ROOT'."
)

# --- 13.5 jrc_msa_attribute ---
add_heading(doc, "13.5  jrc_msa_attribute — Attribute Agreement Analysis", level=2)

add_tc_block(doc,
    "TC-MSA-ATT-001  Dataset with reference → exit 0, all sections present",
    "UR-MSA-017, UR-MSA-018",
    "jrc_msa_attribute attribute_with_ref.csv",
    None,
    "Exit code = 0. Output contains: 'Within-Appraiser', 'Between-Appraiser', "
    "'vs Reference' (or 'Appraiser vs'), 'Verdict', 'Fleiss'."
)
add_tc_block(doc,
    "TC-MSA-ATT-002  Dataset without reference → exit 0, vs-reference section absent",
    "UR-MSA-017",
    "jrc_msa_attribute attribute_no_ref.csv",
    None,
    "Exit code = 0. Output contains 'Within-Appraiser' and 'Between-Appraiser'. "
    "Output does NOT contain 'Vs Reference'."
)
add_tc_block(doc,
    "TC-MSA-ATT-003  Known data — Fleiss' Kappa in expected range, verdict present",
    "UR-MSA-017, UR-MSA-019",
    "jrc_msa_attribute attribute_with_ref.csv",
    "Known dataset; Fleiss' Kappa expected in [0.7, 1.0].",
    "Exit code = 0. Parsed Kappa ∈ [0.7, 1.0]. Verdict = 'MARGINAL' or 'ACCEPTABLE'."
)
add_tc_block(doc,
    "TC-MSA-ATT-004  Perfect appraiser vs reference — Kappa = 1.0",
    "UR-MSA-018",
    "jrc_msa_attribute attribute_with_ref.csv",
    "Appraiser A agrees with reference on all trials in the test dataset.",
    "Exit code = 0. Parsed Kappa for appraiser A vs reference = 1.0000."
)
add_tc_block(doc,
    "TC-MSA-ATT-005  PNG written to ~/Downloads/",
    "UR-MSA-020",
    "jrc_msa_attribute attribute_with_ref.csv",
    None,
    "File matching *_jrc_msa_attribute.png exists in ~/Downloads/ "
    "with modification time ≥ test start."
)
add_tc_block(doc,
    "TC-MSA-ATT-006  No arguments → usage message",
    "UR-MSA-023",
    "jrc_msa_attribute",
    None,
    "Exit code ≠ 0. Output contains 'Usage' or 'usage'."
)
add_tc_block(doc,
    "TC-MSA-ATT-007  File not found",
    "UR-MSA-024",
    "jrc_msa_attribute /tmp/no_such_attr.csv",
    None,
    "Exit code ≠ 0. Output contains 'not found' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-ATT-008  Missing 'trial' column",
    "UR-MSA-021",
    "jrc_msa_attribute attribute_missing_col.csv",
    None,
    "Exit code ≠ 0. Output contains 'trial' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-ATT-009  Only one appraiser",
    "UR-MSA-021",
    "jrc_msa_attribute attribute_one_appraiser.csv",
    None,
    "Exit code ≠ 0."
)
add_tc_block(doc,
    "TC-MSA-ATT-010  Unbalanced design",
    "UR-MSA-021",
    "jrc_msa_attribute attribute_unbalanced.csv",
    None,
    "Exit code ≠ 0. Output contains 'unbalanced', 'trials', or 'equal' (case-insensitive)."
)
add_tc_block(doc,
    "TC-MSA-ATT-011  Bypass protection — direct Rscript call",
    "UR-MSA-022",
    "Rscript repos/msa/R/jrc_msa_attribute.R attribute_with_ref.csv  (RENV_PATHS_ROOT unset)",
    None,
    "Exit code ≠ 0. Output contains 'RENV_PATHS_ROOT'."
)

# ---------------------------------------------------------------------------
# 14. DEVIATION PROCEDURE
# ---------------------------------------------------------------------------

add_heading(doc, "14.  Deviation Procedure", level=1)
add_para(doc,
    "Any test case that does not meet its pass criterion constitutes a deviation. "
    "Deviations shall be handled as follows:")
dev_steps = [
    "Record the deviation in the OQ execution evidence, noting: test case ID, "
    "actual result, and expected result",
    "Investigate the root cause (software defect, test data error, environmental issue)",
    "If a software defect is identified: correct the script, regenerate the integrity "
    "file, and re-execute the full test suite",
    "If a test data error is identified: correct the test data, document the change, "
    "and re-execute the affected test cases",
    "Document the resolution and obtain QA approval before releasing the module",
]
for s in dev_steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)
add_para(doc,
    "The OQ shall not be considered successful until all deviations are resolved "
    "and all 53 test cases pass.", italic=True)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT)
print(f"Saved: {OUT}")
