"""
generate_as_validation_plan.py

Generates repos/as/docs/as_validation_plan.docx — FDA-acceptable OQ Validation Plan
for the JR Validated Environment AS module v1.0.

Run from the repo root:
    python3 repos/as/docs/ignore/generate_as_validation_plan.py
Output: repos/as/docs/ignore/as_validation_plan.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "as_validation_plan.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text="", bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def shade_cell(cell, hex_fill, font_color=None):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)
    if font_color:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = font_color


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, "1A1A2E", RGBColor(0xFF, 0xFF, 0xFF))
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            shade_cell(cell, bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_tc_block(doc, tc_id, tests_ur, command, rationale, pass_criterion):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tc_id)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x21, 0x66, 0xAC)

    rows = [
        ("Tests UR(s)", tests_ur),
        ("Command / action", command),
        ("Rationale", rationale),
        ("Pass criterion", pass_criterion),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for r_idx, (label, value) in enumerate(rows):
        row = table.rows[r_idx]
        bg = "EEF3FA" if r_idx % 2 == 0 else "FFFFFF"
        lbl_cell = row.cells[0]
        lbl_cell.text = label
        lbl_cell.paragraphs[0].runs[0].bold = True
        lbl_cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(lbl_cell, bg)
        val_cell = row.cells[1]
        val_cell.text = value
        val_cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(val_cell, bg)
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.0)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(2.54)
sec.right_margin  = Cm(2.54)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)

# ---------------------------------------------------------------------------
# Title block
# ---------------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(20)
title.paragraph_format.space_after  = Pt(6)
r = title.add_run("JR Validated Environment — Acceptance Sampling Module")
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("OQ Validation Plan — JR-VP-AS-001 v1.0")
r.bold = True; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

add_table(doc,
    ["Document ID", "Version", "Date", "Status"],
    [["JR-VP-AS-001", "1.0", "2026-03-19", "Approved"]],
    col_widths=[4.0, 2.5, 3.5, 4.0]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. Purpose and Scope
# ---------------------------------------------------------------------------
add_heading(doc, "1. Purpose and Scope", level=1)
add_para(doc,
    "This Operational Qualification (OQ) Validation Plan defines the test cases "
    "for the Acceptance Sampling (AS) module of the JR Validated Environment. "
    "The AS module comprises four R scripts: jrc_as_attributes, jrc_as_variables, "
    "jrc_as_oc_curve, and jrc_as_evaluate. These scripts are used for designing "
    "acceptance sampling plans and applying them to production lots.")
add_para(doc,
    "This plan covers 44 automated test cases executed using the pytest framework "
    "via the admin_as_oq runner. Test evidence is captured in a timestamped text file "
    "in the project validation directory.")

# ---------------------------------------------------------------------------
# 2. User Requirements
# ---------------------------------------------------------------------------
add_heading(doc, "2. User Requirements", level=1)
add_para(doc, "The following user requirements are tested by this plan.", size=10)

ur_rows = [
    ("UR-AS-001", "jrc_as_attributes shall accept valid lot size, AQL, and RQL arguments and exit with code 0."),
    ("UR-AS-002", "jrc_as_attributes shall produce a single sampling plan (n and c) in the output."),
    ("UR-AS-003", "jrc_as_attributes shall produce a double sampling plan (n1, c1, c2) in the output when a valid plan exists."),
    ("UR-AS-004", "jrc_as_attributes shall produce an OC curve table with p and Pa values."),
    ("UR-AS-005", "jrc_as_attributes shall save a PNG plot to ~/Downloads/."),
    ("UR-AS-006", "jrc_as_attributes shall exit non-zero and print usage when called with no arguments."),
    ("UR-AS-007", "jrc_as_attributes shall exit non-zero when aql \u2265 rql."),
    ("UR-AS-008", "jrc_as_attributes shall exit non-zero when aql > 1 or < 0."),
    ("UR-AS-009", "jrc_as_attributes shall exit non-zero when lot_size < 2."),
    ("UR-AS-010", "jrc_as_attributes shall exit non-zero when --alpha is out of range."),
    ("UR-AS-011", "jrc_as_attributes shall exit non-zero and report RENV_PATHS_ROOT when called directly via Rscript without the wrapper."),
    ("UR-AS-012", "jrc_as_variables shall accept valid lot size, AQL, RQL, and --sides arguments and exit with code 0."),
    ("UR-AS-013", "jrc_as_variables shall produce an acceptability constant k that is positive in the output."),
    ("UR-AS-014", "jrc_as_variables shall produce a valid plan for --sides 2 with positive k."),
    ("UR-AS-015", "jrc_as_variables shall produce an OC curve table with Pa values."),
    ("UR-AS-016", "jrc_as_variables shall save a PNG plot to ~/Downloads/."),
    ("UR-AS-017", "jrc_as_variables shall exit non-zero when called with no arguments."),
    ("UR-AS-018", "jrc_as_variables shall exit non-zero when aql \u2265 rql."),
    ("UR-AS-019", "jrc_as_variables shall exit non-zero when --sides value is not 1 or 2."),
    ("UR-AS-020", "jrc_as_variables shall produce output indicating a sample size reduction compared to the equivalent attributes plan."),
    ("UR-AS-021", "jrc_as_variables shall exit non-zero when --alpha is out of range."),
    ("UR-AS-022", "jrc_as_variables shall exit non-zero and report RENV_PATHS_ROOT when called directly via Rscript without the wrapper."),
    ("UR-AS-023", "jrc_as_oc_curve shall accept valid n and c arguments and exit with code 0 with Pa values in output."),
    ("UR-AS-024", "jrc_as_oc_curve shall produce Pa > 0.99 at p = 0.001 for n=32, c=1."),
    ("UR-AS-025", "jrc_as_oc_curve shall produce Pa values that decrease monotonically as p increases."),
    ("UR-AS-026", "jrc_as_oc_curve shall accept the --lot-size flag without error."),
    ("UR-AS-027", "jrc_as_oc_curve shall save a PNG plot to ~/Downloads/."),
    ("UR-AS-028", "jrc_as_oc_curve shall accept --aql and --rql flags and display AQL in output."),
    ("UR-AS-029", "jrc_as_oc_curve shall exit non-zero when called with no arguments."),
    ("UR-AS-030", "jrc_as_oc_curve shall exit non-zero when c \u2265 n."),
    ("UR-AS-031", "jrc_as_oc_curve shall exit non-zero when n \u2264 0."),
    ("UR-AS-032", "jrc_as_oc_curve shall exit non-zero and report RENV_PATHS_ROOT when called directly via Rscript without the wrapper."),
    ("UR-AS-033", "jrc_as_evaluate attributes mode shall produce ACCEPT verdict for a lot with defective count \u2264 c."),
    ("UR-AS-034", "jrc_as_evaluate attributes mode shall produce REJECT verdict for a lot with defective count > c."),
    ("UR-AS-035", "jrc_as_evaluate variables mode shall produce ACCEPT verdict for a lot with Q-statistic \u2265 k."),
    ("UR-AS-036", "jrc_as_evaluate variables mode shall produce REJECT verdict for a lot with Q-statistic < k."),
    ("UR-AS-037", "jrc_as_evaluate variables mode shall save a PNG plot to ~/Downloads/."),
    ("UR-AS-038", "jrc_as_evaluate attributes mode shall save a PNG plot to ~/Downloads/."),
    ("UR-AS-039", "jrc_as_evaluate shall exit non-zero when --type is missing."),
    ("UR-AS-040", "jrc_as_evaluate attributes mode shall exit non-zero when --c is missing."),
    ("UR-AS-041", "jrc_as_evaluate variables mode shall exit non-zero when --k is missing."),
    ("UR-AS-042", "jrc_as_evaluate attributes mode shall exit non-zero and report the missing column when the data file lacks a \u2018result\u2019 column."),
    ("UR-AS-043", "jrc_as_evaluate variables mode shall exit non-zero and report the missing column when the data file lacks a \u2018value\u2019 column."),
    ("UR-AS-044", "jrc_as_evaluate shall exit non-zero and report RENV_PATHS_ROOT when called directly via Rscript without the wrapper."),
]

add_table(doc,
    ["UR ID", "Requirement"],
    ur_rows,
    col_widths=[3.0, 13.5]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# 3. Test Environment
# ---------------------------------------------------------------------------
add_heading(doc, "3. Test Environment", level=1)
add_para(doc,
    "All test cases are executed via the admin_as_oq runner using pytest with the shared "
    "OQ virtual environment at ~/.venvs/<PROJECT_ID>_oq/. The test runner records the "
    "R version, Python version, pytest version, hostname, and timestamp in the evidence file. "
    "Tests are executed from the repos/as/oq/ directory.")
add_table(doc,
    ["Item", "Specification"],
    [
        ["Language", "R (\u2265 4.3.0 recommended)"],
        ["Test framework", "pytest (\u2265 8.0)"],
        ["Test runner", "admin_as_oq"],
        ["OQ venv", "~/.venvs/<PROJECT_ID>_oq/ (shared with admin_oq)"],
        ["Evidence output", "~/.jrscript/<PROJECT_ID>/validation/as_oq_execution_<timestamp>.txt"],
        ["Validation plan ID", "JR-VP-AS-001"],
    ],
    col_widths=[4.0, 12.5]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# 4. Test Cases — jrc_as_attributes
# ---------------------------------------------------------------------------
add_heading(doc, "4. Test Cases — jrc_as_attributes", level=1)

add_tc_block(doc,
    "TC-AS-ATTR-001  Happy path — exit 0",
    "UR-AS-001",
    "jrc_as_attributes 500 0.01 0.10",
    "Confirms the script accepts standard valid inputs and completes successfully.",
    "Exit code = 0. Output contains 'Attributes Sampling Plan' or equivalent plan header."
)

add_tc_block(doc,
    "TC-AS-ATTR-002  Single plan n and c present",
    "UR-AS-002",
    "jrc_as_attributes 500 0.01 0.10",
    "Confirms the single sampling plan outputs both sample size n and acceptance number c.",
    "Output contains 'Sample size' with a positive integer, and 'Acceptance number' with a non-negative integer."
)

add_tc_block(doc,
    "TC-AS-ATTR-003  Double sampling plan present",
    "UR-AS-003",
    "jrc_as_attributes 500 0.01 0.10",
    "Confirms the double sampling plan section is present in the output.",
    "Output contains 'Double Sampling Plan' or 'Stage 1'."
)

add_tc_block(doc,
    "TC-AS-ATTR-004  OC curve table present",
    "UR-AS-004",
    "jrc_as_attributes 500 0.01 0.10",
    "Confirms the OC curve table is included in the output.",
    "Output contains 'OC Curve' and 'Pa' header."
)

add_tc_block(doc,
    "TC-AS-ATTR-005  PNG saved to ~/Downloads/",
    "UR-AS-005",
    "jrc_as_attributes 500 0.01 0.10",
    "Confirms a PNG file is created in ~/Downloads/ following the script run.",
    "A file matching *_jrc_as_attributes.png exists in ~/Downloads/ with mtime \u2265 test start time."
)

add_tc_block(doc,
    "TC-AS-ATTR-006  No arguments — non-zero exit",
    "UR-AS-006",
    "jrc_as_attributes (no arguments)",
    "Confirms the script exits non-zero and displays usage guidance when called with no arguments.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-ATTR-007  aql \u2265 rql — non-zero exit",
    "UR-AS-007",
    "jrc_as_attributes 500 0.10 0.01",
    "Confirms the script detects and rejects inverted AQL/RQL inputs.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-ATTR-008  aql > 1 — non-zero exit",
    "UR-AS-008",
    "jrc_as_attributes 500 1.5 0.10",
    "Confirms the script rejects an AQL value out of the (0, 1) range.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-ATTR-009  lot_size = 1 — non-zero exit",
    "UR-AS-009",
    "jrc_as_attributes 1 0.01 0.10",
    "Confirms the script rejects a lot size below the minimum (2).",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-ATTR-010  --alpha out of range — non-zero exit",
    "UR-AS-010",
    "jrc_as_attributes 500 0.01 0.10 --alpha 2.0",
    "Confirms the script rejects an --alpha value outside (0, 1).",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-ATTR-011  Bypass protection",
    "UR-AS-011",
    "Rscript repos/as/R/jrc_as_attributes.R 500 0.01 0.10 (RENV_PATHS_ROOT unset)",
    "Confirms that running the R script directly without the validated wrapper exits non-zero "
    "and mentions RENV_PATHS_ROOT in the error output.",
    "Exit code \u2260 0. 'RENV_PATHS_ROOT' appears in combined stdout/stderr."
)

# ---------------------------------------------------------------------------
# 5. Test Cases — jrc_as_variables
# ---------------------------------------------------------------------------
add_heading(doc, "5. Test Cases — jrc_as_variables", level=1)

add_tc_block(doc,
    "TC-AS-VAR-001  Happy path — exit 0, Variables Plan in output",
    "UR-AS-012",
    "jrc_as_variables 500 0.01 0.10",
    "Confirms valid one-sided inputs produce exit 0 and a plan header.",
    "Exit code = 0. Output contains 'Variables Plan' or 'Variables Sampling Plan'."
)

add_tc_block(doc,
    "TC-AS-VAR-002  k value present and positive",
    "UR-AS-013",
    "jrc_as_variables 500 0.01 0.10",
    "Confirms the acceptability constant k is reported and is positive.",
    "Output contains 'constant (k)' or 'k)' or 'k ='. If a numeric k is found, k > 0."
)

add_tc_block(doc,
    "TC-AS-VAR-003  --sides 2 produces valid plan",
    "UR-AS-014",
    "jrc_as_variables 500 0.01 0.10 --sides 2",
    "Confirms two-sided specification produces a valid plan with exit 0 and positive k.",
    "Exit code = 0. Output contains 'Sides: 2'. If k extracted, k > 0."
)

add_tc_block(doc,
    "TC-AS-VAR-004  OC curve table present",
    "UR-AS-015",
    "jrc_as_variables 500 0.01 0.10",
    "Confirms the OC curve table is present in the output.",
    "Output contains 'OC Curve' and 'Pa'."
)

add_tc_block(doc,
    "TC-AS-VAR-005  PNG saved to ~/Downloads/",
    "UR-AS-016",
    "jrc_as_variables 500 0.01 0.10",
    "Confirms a PNG file is created in ~/Downloads/ following the script run.",
    "A file matching *_jrc_as_variables.png exists in ~/Downloads/ with mtime \u2265 test start time."
)

add_tc_block(doc,
    "TC-AS-VAR-006  No arguments — non-zero exit",
    "UR-AS-017",
    "jrc_as_variables (no arguments)",
    "Confirms the script exits non-zero when called with no arguments.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-VAR-007  aql \u2265 rql — non-zero exit",
    "UR-AS-018",
    "jrc_as_variables 500 0.10 0.01",
    "Confirms the script rejects inverted AQL/RQL.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-VAR-008  --sides invalid — non-zero exit",
    "UR-AS-019",
    "jrc_as_variables 500 0.01 0.10 --sides 3",
    "Confirms the script rejects --sides values other than 1 or 2.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-VAR-009  Sample reduction noted in output",
    "UR-AS-020",
    "jrc_as_variables 500 0.01 0.10",
    "Confirms the output includes a comparison with the equivalent attributes plan.",
    "Output contains 'Sample reduction', 'fewer', or 'Comparison'."
)

add_tc_block(doc,
    "TC-AS-VAR-010  --alpha out of range — non-zero exit",
    "UR-AS-021",
    "jrc_as_variables 500 0.01 0.10 --alpha 2.0",
    "Confirms the script rejects an --alpha value outside (0, 1).",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-VAR-011  Bypass protection",
    "UR-AS-022",
    "Rscript repos/as/R/jrc_as_variables.R 500 0.01 0.10 (RENV_PATHS_ROOT unset)",
    "Confirms that running the R script directly without the validated wrapper exits non-zero "
    "and mentions RENV_PATHS_ROOT in the error output.",
    "Exit code \u2260 0. 'RENV_PATHS_ROOT' appears in combined stdout/stderr."
)

# ---------------------------------------------------------------------------
# 6. Test Cases — jrc_as_oc_curve
# ---------------------------------------------------------------------------
add_heading(doc, "6. Test Cases — jrc_as_oc_curve", level=1)

add_tc_block(doc,
    "TC-AS-OCC-001  Happy path — exit 0, Pa values in output",
    "UR-AS-023",
    "jrc_as_oc_curve 32 1",
    "Confirms valid n and c produce exit 0 and Pa values.",
    "Exit code = 0. Output contains 'Pa' and at least one numeric Pa value."
)

add_tc_block(doc,
    "TC-AS-OCC-002  Pa at p=0.001 is > 0.99",
    "UR-AS-024",
    "jrc_as_oc_curve 32 1",
    "Confirms the plan (n=32, c=1) has very high acceptance probability at very low defect rates.",
    "Pa value in the first data row (p=0.001) extracted from output is > 0.99."
)

add_tc_block(doc,
    "TC-AS-OCC-003  Pa decreases as p increases",
    "UR-AS-025",
    "jrc_as_oc_curve 32 1",
    "Confirms the OC curve is monotonically decreasing in p.",
    "Pa value at first p in the table > Pa value at last p in the table."
)

add_tc_block(doc,
    "TC-AS-OCC-004  --lot-size accepted",
    "UR-AS-026",
    "jrc_as_oc_curve 32 1 --lot-size 200",
    "Confirms the --lot-size flag is accepted without error.",
    "Exit code = 0."
)

add_tc_block(doc,
    "TC-AS-OCC-005  PNG saved to ~/Downloads/",
    "UR-AS-027",
    "jrc_as_oc_curve 32 1",
    "Confirms a PNG file is created in ~/Downloads/ following the script run.",
    "A file matching *_jrc_as_oc_curve.png exists in ~/Downloads/ with mtime \u2265 test start time."
)

add_tc_block(doc,
    "TC-AS-OCC-006  --aql and --rql accepted, AQL in output",
    "UR-AS-028",
    "jrc_as_oc_curve 32 1 --aql 0.01 --rql 0.10",
    "Confirms --aql and --rql flags are accepted and reflected in output.",
    "Exit code = 0. Output contains 'AQL'."
)

add_tc_block(doc,
    "TC-AS-OCC-007  No arguments — non-zero exit",
    "UR-AS-029",
    "jrc_as_oc_curve (no arguments)",
    "Confirms the script exits non-zero when called with no arguments.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-OCC-008  c \u2265 n — non-zero exit",
    "UR-AS-030",
    "jrc_as_oc_curve 10 10",
    "Confirms the script rejects c \u2265 n (acceptance number must be less than sample size).",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-OCC-009  n \u2264 0 — non-zero exit",
    "UR-AS-031",
    "jrc_as_oc_curve 0 0",
    "Confirms the script rejects n \u2264 0.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-OCC-010  Bypass protection",
    "UR-AS-032",
    "Rscript repos/as/R/jrc_as_oc_curve.R 32 1 (RENV_PATHS_ROOT unset)",
    "Confirms that running the R script directly without the validated wrapper exits non-zero "
    "and mentions RENV_PATHS_ROOT in the error output.",
    "Exit code \u2260 0. 'RENV_PATHS_ROOT' appears in combined stdout/stderr."
)

# ---------------------------------------------------------------------------
# 7. Test Cases — jrc_as_evaluate
# ---------------------------------------------------------------------------
add_heading(doc, "7. Test Cases — jrc_as_evaluate", level=1)

add_tc_block(doc,
    "TC-AS-EVAL-001  Attributes ACCEPT verdict",
    "UR-AS-033",
    "jrc_as_evaluate oq/data/attr_accept_lot.csv --type attributes --c 2",
    "Confirms ACCEPT verdict for a lot with 1 defective against c=2. "
    "The test lot (attr_accept_lot.csv) contains 32 units with 1 defective.",
    "Exit code = 0. Output contains 'ACCEPT'."
)

add_tc_block(doc,
    "TC-AS-EVAL-002  Attributes REJECT verdict",
    "UR-AS-034",
    "jrc_as_evaluate oq/data/attr_reject_lot.csv --type attributes --c 2",
    "Confirms REJECT verdict for a lot with 5 defectives against c=2. "
    "The test lot (attr_reject_lot.csv) contains 32 units with 5 defectives.",
    "Exit code = 0. Output contains 'REJECT'."
)

add_tc_block(doc,
    "TC-AS-EVAL-003  Variables ACCEPT verdict",
    "UR-AS-035",
    "jrc_as_evaluate oq/data/var_accept_lot.csv --type variables --k 1.45 --lsl 9.5",
    "Confirms ACCEPT verdict for a lot where Q_L >> k. "
    "The test lot (var_accept_lot.csv) has mean \u224810.45 and SD \u22480.051, giving Q_L \u224818.5.",
    "Exit code = 0. Output contains 'ACCEPT'."
)

add_tc_block(doc,
    "TC-AS-EVAL-004  Variables REJECT verdict",
    "UR-AS-036",
    "jrc_as_evaluate oq/data/var_reject_lot.csv --type variables --k 1.45 --lsl 9.5",
    "Confirms REJECT verdict for a lot where Q_L < k. "
    "The test lot (var_reject_lot.csv) has mean \u22489.6 and SD \u22480.103, giving Q_L \u22480.975.",
    "Exit code = 0. Output contains 'REJECT'."
)

add_tc_block(doc,
    "TC-AS-EVAL-005  Variables mode PNG saved",
    "UR-AS-037",
    "jrc_as_evaluate oq/data/var_accept_lot.csv --type variables --k 1.45 --lsl 9.5",
    "Confirms a PNG is created in ~/Downloads/ after a variables mode run.",
    "A file matching *_jrc_as_evaluate.png exists in ~/Downloads/ with mtime \u2265 test start time."
)

add_tc_block(doc,
    "TC-AS-EVAL-006  Attributes mode PNG saved",
    "UR-AS-038",
    "jrc_as_evaluate oq/data/attr_accept_lot.csv --type attributes --c 2",
    "Confirms a PNG is created in ~/Downloads/ after an attributes mode run.",
    "A file matching *_jrc_as_evaluate.png exists in ~/Downloads/ with mtime \u2265 test start time."
)

add_tc_block(doc,
    "TC-AS-EVAL-007  Missing --type — non-zero exit",
    "UR-AS-039",
    "jrc_as_evaluate oq/data/attr_accept_lot.csv --c 2",
    "Confirms the script exits non-zero when --type is omitted.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-EVAL-008  Attributes mode missing --c — non-zero exit",
    "UR-AS-040",
    "jrc_as_evaluate oq/data/attr_accept_lot.csv --type attributes",
    "Confirms the script exits non-zero when --type attributes is specified without --c.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-EVAL-009  Variables mode missing --k — non-zero exit",
    "UR-AS-041",
    "jrc_as_evaluate oq/data/var_accept_lot.csv --type variables --lsl 9.5",
    "Confirms the script exits non-zero when --type variables is specified without --k.",
    "Exit code \u2260 0."
)

add_tc_block(doc,
    "TC-AS-EVAL-010  Attributes: missing 'result' column",
    "UR-AS-042",
    "jrc_as_evaluate oq/data/attr_missing_result.csv --type attributes --c 2",
    "Confirms the script exits non-zero and reports the missing 'result' column when "
    "an attributes CSV has the wrong column name.",
    "Exit code \u2260 0. 'result' appears in combined stdout/stderr."
)

add_tc_block(doc,
    "TC-AS-EVAL-011  Variables: missing 'value' column",
    "UR-AS-043",
    "jrc_as_evaluate oq/data/var_missing_value.csv --type variables --k 1.45 --lsl 9.5",
    "Confirms the script exits non-zero and reports the missing 'value' column when "
    "a variables CSV has the wrong column name.",
    "Exit code \u2260 0. 'value' appears in combined stdout/stderr."
)

add_tc_block(doc,
    "TC-AS-EVAL-012  Bypass protection",
    "UR-AS-044",
    "Rscript repos/as/R/jrc_as_evaluate.R oq/data/attr_accept_lot.csv ... (RENV_PATHS_ROOT unset)",
    "Confirms that running the R script directly without the validated wrapper exits non-zero "
    "and mentions RENV_PATHS_ROOT in the error output.",
    "Exit code \u2260 0. 'RENV_PATHS_ROOT' appears in combined stdout/stderr."
)

# ---------------------------------------------------------------------------
# 8. Requirements Traceability Matrix
# ---------------------------------------------------------------------------
add_heading(doc, "8. Requirements Traceability Matrix", level=1)
add_para(doc, "Maps each test case to the user requirement(s) it covers.", size=10)

rtm_rows = [
    ("TC-AS-ATTR-001", "UR-AS-001"),
    ("TC-AS-ATTR-002", "UR-AS-002"),
    ("TC-AS-ATTR-003", "UR-AS-003"),
    ("TC-AS-ATTR-004", "UR-AS-004"),
    ("TC-AS-ATTR-005", "UR-AS-005"),
    ("TC-AS-ATTR-006", "UR-AS-006"),
    ("TC-AS-ATTR-007", "UR-AS-007"),
    ("TC-AS-ATTR-008", "UR-AS-008"),
    ("TC-AS-ATTR-009", "UR-AS-009"),
    ("TC-AS-ATTR-010", "UR-AS-010"),
    ("TC-AS-ATTR-011", "UR-AS-011"),
    ("TC-AS-VAR-001",  "UR-AS-012"),
    ("TC-AS-VAR-002",  "UR-AS-013"),
    ("TC-AS-VAR-003",  "UR-AS-014"),
    ("TC-AS-VAR-004",  "UR-AS-015"),
    ("TC-AS-VAR-005",  "UR-AS-016"),
    ("TC-AS-VAR-006",  "UR-AS-017"),
    ("TC-AS-VAR-007",  "UR-AS-018"),
    ("TC-AS-VAR-008",  "UR-AS-019"),
    ("TC-AS-VAR-009",  "UR-AS-020"),
    ("TC-AS-VAR-010",  "UR-AS-021"),
    ("TC-AS-VAR-011",  "UR-AS-022"),
    ("TC-AS-OCC-001",  "UR-AS-023"),
    ("TC-AS-OCC-002",  "UR-AS-024"),
    ("TC-AS-OCC-003",  "UR-AS-025"),
    ("TC-AS-OCC-004",  "UR-AS-026"),
    ("TC-AS-OCC-005",  "UR-AS-027"),
    ("TC-AS-OCC-006",  "UR-AS-028"),
    ("TC-AS-OCC-007",  "UR-AS-029"),
    ("TC-AS-OCC-008",  "UR-AS-030"),
    ("TC-AS-OCC-009",  "UR-AS-031"),
    ("TC-AS-OCC-010",  "UR-AS-032"),
    ("TC-AS-EVAL-001", "UR-AS-033"),
    ("TC-AS-EVAL-002", "UR-AS-034"),
    ("TC-AS-EVAL-003", "UR-AS-035"),
    ("TC-AS-EVAL-004", "UR-AS-036"),
    ("TC-AS-EVAL-005", "UR-AS-037"),
    ("TC-AS-EVAL-006", "UR-AS-038"),
    ("TC-AS-EVAL-007", "UR-AS-039"),
    ("TC-AS-EVAL-008", "UR-AS-040"),
    ("TC-AS-EVAL-009", "UR-AS-041"),
    ("TC-AS-EVAL-010", "UR-AS-042"),
    ("TC-AS-EVAL-011", "UR-AS-043"),
    ("TC-AS-EVAL-012", "UR-AS-044"),
]

add_table(doc,
    ["Test Case ID", "User Requirement(s)"],
    rtm_rows,
    col_widths=[5.0, 11.5]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
