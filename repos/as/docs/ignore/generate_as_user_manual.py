"""
generate_as_user_manual.py

Generates repos/as/docs/as_user_manual.docx — the engineer-facing user guide
for the JR Validated Environment AS Module.

Run with:
    python3 repos/as/docs/ignore/generate_as_user_manual.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "as_user_manual.docx")

NAVY  = "1A1A2E"
GRAY  = "F2F2F2"
WHITE = "FFFFFF"
INFO  = "E8F0F7"
WARN  = "FFF3CD"
GREEN = "E8F5E9"

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def header_row(table, cols, widths=None):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        set_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def data_row(table, row_idx, values, shade=True, bold_first=False):
    row = table.rows[row_idx]
    fill = GRAY if (shade and row_idx % 2 == 1) else WHITE
    for i, text in enumerate(values):
        cell = row.cells[i]
        set_shading(cell, fill)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if bold_first and i == 0:
            run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def para(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = "Courier New"


def box(doc, fill, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    p.clear()
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.name = "Calibri"
    doc.add_paragraph()


def info(doc, label, text):  box(doc, INFO, label, text)
def warn(doc, label, text):  box(doc, WARN, label, text)
def tip(doc, label, text):   box(doc, GREEN, label, text)


def script_section(doc, number, script_name, when, args_table, example_cmds, output_note):
    h2(doc, f"{number}  {script_name}")
    h3(doc, "When do I use this?")
    para(doc, when)
    h3(doc, "What do I need?")
    tbl = doc.add_table(rows=len(args_table) + 1, cols=2)
    tbl.style = "Table Grid"
    header_row(tbl, ["Argument", "What to provide"], [1.8, 5.1])
    for i, (arg, desc) in enumerate(args_table, start=1):
        data_row(tbl, i, [arg, desc], bold_first=True)
    doc.add_paragraph()
    h3(doc, "Example")
    for cmd in example_cmds:
        code(doc, cmd)
    h3(doc, "What to look for in the output")
    para(doc, output_note)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(0.984)
section.top_margin    = Inches(0.984)
section.bottom_margin = Inches(0.984)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
cover.paragraph_format.space_after  = Pt(6)
r = cover.add_run("JR Validated Environment \u2014 AS Module")
r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("User Guide for Acceptance Sampling")
r.bold = True; r.font.size = Pt(13); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  |  Date: 2026-03-19  |  Audience: Design and Manufacturing Engineers")
r.font.size = Pt(10); r.font.name = "Calibri"
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 1 — How to Use This Guide
# ---------------------------------------------------------------------------
h1(doc, "1. How to Use This Guide")
para(doc,
    "This guide describes the four acceptance sampling scripts available in the JR Validated "
    "Environment. Acceptance sampling is used to decide whether to accept or reject a production "
    "lot based on the inspection of a sample. The workflow is:")

bullet(doc, "Step 1 \u2014 Design a plan: use jrc_as_attributes or jrc_as_variables to compute sample size and acceptance criteria.")
bullet(doc, "Step 2 \u2014 Visualise the plan: use jrc_as_oc_curve to inspect the OC curve before committing to production sampling.")
bullet(doc, "Step 3 \u2014 Evaluate the lot: use jrc_as_evaluate to apply the plan to real inspection data and obtain a formal ACCEPT or REJECT verdict.")

para(doc,
    "The two plan types are: attributes (pass/fail counts, simpler but less efficient) and "
    "variables (continuous measurements, requires fewer samples but the data must follow a "
    "normal distribution). See Section 2 for the quick reference chart.")

# ---------------------------------------------------------------------------
# Section 2 — Quick Reference
# ---------------------------------------------------------------------------
h1(doc, "2. Quick Reference")

qr_rows = [
    ("I need to design a sampling plan \u2014 my inspection data will be pass/fail (0/1)",
     "jrc_as_attributes", "4.1"),
    ("I need to design a sampling plan \u2014 my inspection data will be continuous measurements",
     "jrc_as_variables", "4.2"),
    ("I want to plot the OC curve for a plan I already have (n, c)",
     "jrc_as_oc_curve", "4.3"),
    ("I have collected lot inspection data and need an ACCEPT/REJECT verdict",
     "jrc_as_evaluate", "4.4"),
]

tbl = doc.add_table(rows=len(qr_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["My situation is\u2026", "Script", "Section"], [4.0, 2.0, 0.9])
for i, row_data in enumerate(qr_rows, start=1):
    data_row(tbl, i, list(row_data))
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 3 — Understanding Acceptance Sampling
# ---------------------------------------------------------------------------
h1(doc, "3. Understanding Acceptance Sampling")

h2(doc, "3.1  AQL and RQL")
para(doc,
    "Two quality levels define the plan:")
bullet(doc,
    "AQL (Acceptable Quality Level) \u2014 the fraction nonconforming at which the lot is "
    "considered acceptable. The plan must accept lots of this quality with high probability "
    "(probability of acceptance \u2265 1 \u2212 \u03b1, where \u03b1 is the producer\u2019s risk).")
bullet(doc,
    "RQL (Rejectable Quality Level, also called LTPD or LQ) \u2014 the fraction nonconforming "
    "at which the lot should be rejected. The plan must reject lots of this quality with high "
    "probability (\u2265 1 \u2212 \u03b2, where \u03b2 is the consumer\u2019s risk).")

h2(doc, "3.2  Producer\u2019s Risk (\u03b1) and Consumer\u2019s Risk (\u03b2)")
para(doc,
    "The producer\u2019s risk \u03b1 is the probability of rejecting a good lot (one at or better than AQL). "
    "The consumer\u2019s risk \u03b2 is the probability of accepting a bad lot (one at or worse than RQL). "
    "Both default to industry-standard values: \u03b1 = 0.05, \u03b2 = 0.10.")

h2(doc, "3.3  Attributes vs Variables")
para(doc,
    "An attributes plan inspects units as pass/fail and counts the number of defectives. "
    "A variables plan records a continuous measurement for each unit and compares the sample "
    "mean and standard deviation against a critical value k. Variables plans typically require "
    "significantly fewer samples because each measurement carries more information than a "
    "pass/fail result. However, variables plans require that measurements follow a normal "
    "distribution and that a specification limit is known.")

h2(doc, "3.4  OC Curve")
para(doc,
    "The Operating Characteristic (OC) curve plots Pa (probability of accepting the lot) "
    "against p (fraction nonconforming in the lot). A good plan has Pa close to 1.0 at AQL "
    "and Pa close to 0 at RQL. The steeper and more S-shaped the OC curve, the better the "
    "plan discriminates between acceptable and unacceptable lots.")

info(doc, "INFO:",
    "The scripts use the hypergeometric distribution when the sample fraction n/N exceeds 10% "
    "of the lot (large sample relative to lot size). Otherwise the binomial distribution is used. "
    "The hypergeometric calculation is more accurate for small finite lots.")

# ---------------------------------------------------------------------------
# Section 4 — Per-Script Reference
# ---------------------------------------------------------------------------
h1(doc, "4. Per-Script Reference")

# 4.1 jrc_as_attributes
script_section(doc,
    "4.1", "jrc_as_attributes \u2014 Attributes Sampling Plan (Single and Double)",
    "Use when your lot inspection will produce pass/fail results (each unit is either "
    "conforming or nonconforming). The script finds the minimum sample size n and acceptance "
    "number c for a single sampling plan, and n1, c1, c2 for a double sampling plan. "
    "Double sampling reduces the expected number of units inspected (ASN) at the cost of "
    "a more complex procedure.",
    [
        ("lot_size",
         "Total number of units in the lot (positive integer \u2265 2). Used for hypergeometric "
         "correction when the sample fraction exceeds 10%."),
        ("aql",
         "Acceptable Quality Level as a fraction strictly between 0 and 1 "
         "(e.g. 0.01 = 1%). Must be less than rql."),
        ("rql",
         "Rejectable Quality Level as a fraction strictly between 0 and 1 "
         "(e.g. 0.10 = 10%). Must be greater than aql."),
        ("--alpha <val>",
         "Producer\u2019s risk (default 0.05). Probability of rejecting a good lot at AQL."),
        ("--beta <val>",
         "Consumer\u2019s risk (default 0.10). Probability of accepting a bad lot at RQL."),
    ],
    [
        "jrc_as_attributes 500 0.01 0.10",
        "jrc_as_attributes 1000 0.005 0.05 --alpha 0.05 --beta 0.10",
    ],
    "Check the \u2018Single Sampling Plan\u2019 section for n (sample size) and c (acceptance number). "
    "The verdict rule is: sample n units; accept if the number of defectives \u2264 c. "
    "The \u2018Double Sampling Plan\u2019 section (if found) gives n1, c1, c2 with the ASN benefit. "
    "The OC curve table shows Pa at key quality levels including AQL and RQL. "
    "The PNG saves to ~/Downloads/ with a timestamp prefix."
)

tip(doc, "TIP:",
    "For N=500, AQL=0.01, RQL=0.10: Single plan gives n=51, c=2. Double plan gives n1=29, "
    "c1=0, c2=2, reducing expected sample size to 36.2 units at AQL (vs 51 for single). "
    "Use the double plan when minimising inspection cost is important.")

# 4.2 jrc_as_variables
script_section(doc,
    "4.2", "jrc_as_variables \u2014 Variables Sampling Plan (k-method, unknown \u03c3)",
    "Use when your lot inspection will produce continuous measurements (e.g. dimensions, "
    "force, electrical resistance) and you know the specification limit(s). The k-method "
    "computes a sample size n and acceptability constant k. The lot is accepted if the "
    "Q-statistic (\u0305x \u2212 LSL)/s \u2265 k (one-sided lower) or (USL \u2212 \u0305x)/s \u2265 k (one-sided upper). "
    "Two-sided specifications use --sides 2.",
    [
        ("lot_size",
         "Total number of units in the lot (positive integer \u2265 2)."),
        ("aql",
         "Acceptable Quality Level as a fraction strictly between 0 and 1."),
        ("rql",
         "Rejectable Quality Level as a fraction strictly between 0 and 1. Must be > aql."),
        ("--alpha <val>",
         "Producer\u2019s risk (default 0.05)."),
        ("--beta <val>",
         "Consumer\u2019s risk (default 0.10)."),
        ("--sides <1|2>",
         "1 = one-sided specification (default). 2 = two-sided specification "
         "(both LSL and USL must be checked; tighter plan)."),
    ],
    [
        "jrc_as_variables 500 0.01 0.10",
        "jrc_as_variables 500 0.01 0.10 --sides 2",
        "jrc_as_variables 1000 0.005 0.05 --alpha 0.05 --beta 0.10",
    ],
    "The output shows the sample size n and the acceptability constant k. Collect n "
    "measurements from the lot, compute the sample mean \u0305x and sample standard deviation s, "
    "then compute Q = (\u0305x \u2212 LSL)/s (lower-sided) or Q = (USL \u2212 \u0305x)/s (upper-sided). "
    "Accept the lot if Q \u2265 k. The comparison section shows how many fewer units the variables "
    "plan requires compared to the equivalent attributes plan."
)

tip(doc, "TIP:",
    "For N=500, AQL=0.01, RQL=0.10 (one-sided): n=21, k=1.7608, achieving \u03b1=0.0500, "
    "\u03b2=0.0956. This is a 58.8% reduction from the equivalent attributes plan (n=51). "
    "Variables plans are most beneficial when inspection is costly per unit.")

warn(doc, "NOTE:",
    "The k-method assumes measurements follow a normal distribution within the lot. "
    "If the distribution is not normal, the risk levels \u03b1 and \u03b2 will not be achieved. "
    "Run jrc_normality on pilot data before adopting a variables plan.")

# 4.3 jrc_as_oc_curve
script_section(doc,
    "4.3", "jrc_as_oc_curve \u2014 OC Curve for Any Attributes Plan",
    "Use to plot and tabulate the Operating Characteristic curve for any attributes "
    "sampling plan defined by sample size n and acceptance number c. You can supply an "
    "optional lot size (for hypergeometric correction), AQL, and RQL (to add reference "
    "lines to the plot). This script does not design a plan \u2014 it visualises one you "
    "already have, whether from jrc_as_attributes or from a standard table.",
    [
        ("n",
         "Sample size \u2014 positive integer. The number of units to inspect."),
        ("c",
         "Acceptance number \u2014 non-negative integer. Accept the lot if defectives \u2264 c. "
         "Must be less than n."),
        ("--lot-size N",
         "Optional. Total lot size. If n/N > 0.10, hypergeometric distribution is used "
         "instead of binomial. Omit for large lots."),
        ("--aql <val>",
         "Optional. AQL as a fraction. Adds a vertical AQL marker to the plot and "
         "includes AQL in the Pa table."),
        ("--rql <val>",
         "Optional. RQL as a fraction. Adds a vertical RQL marker to the plot and "
         "includes RQL in the Pa table."),
    ],
    [
        "jrc_as_oc_curve 51 2 --lot-size 500 --aql 0.01 --rql 0.10",
        "jrc_as_oc_curve 32 1",
        "jrc_as_oc_curve 80 3 --aql 0.02 --rql 0.08",
    ],
    "The terminal shows a Pa table (p, Pa, 1\u2212Pa columns). Pa at AQL should be close to "
    "1 \u2212 \u03b1 (\u2265 0.95). Pa at RQL should be close to \u03b2 (\u2264 0.10). The PNG is saved to "
    "~/Downloads/ and shows the OC curve with AQL and RQL dashed lines if supplied."
)

# 4.4 jrc_as_evaluate
script_section(doc,
    "4.4", "jrc_as_evaluate \u2014 Apply a Plan to a Lot (ACCEPT / REJECT)",
    "Use after sampling a production lot to formally apply the sampling plan and "
    "record the lot disposition decision. The script reads a CSV of inspection results, "
    "applies the plan criteria, and outputs an ACCEPT or REJECT verdict with a timestamped "
    "PNG saved to ~/Downloads/. Supports both attributes (pass/fail) and variables (measurement) modes.",
    [
        ("data.csv",
         "Path to the inspection data CSV. "
         "Attributes mode: columns id and result (0 = conforming, 1 = defective). "
         "Variables mode: columns id and value (numeric measurements)."),
        ("--type <attributes|variables>",
         "Required. Specify the inspection type: attributes or variables."),
        ("--n <value>",
         "Optional. Expected sample size. If provided, the script checks the actual row "
         "count against n and warns if they do not match."),
        ("--c <value>",
         "Attributes mode: acceptance number. Accept if number of defectives \u2264 c."),
        ("--k <value>",
         "Variables mode: acceptability constant. Accept if Q-statistic \u2265 k."),
        ("--lsl <value>",
         "Variables mode: lower specification limit. Used to compute Q_L = (\u0305x \u2212 LSL)/s."),
        ("--usl <value>",
         "Variables mode: upper specification limit. Used to compute Q_U = (USL \u2212 \u0305x)/s."),
    ],
    [
        "jrc_as_evaluate ~/data/lot_results.csv --type attributes --n 51 --c 2",
        "jrc_as_evaluate ~/data/measurements.csv --type variables --n 21 --k 1.7608 --lsl 9.5",
        "jrc_as_evaluate ~/data/measurements.csv --type variables --k 1.8556 --lsl 9.5 --usl 11.5",
    ],
    "The terminal shows the verdict (ACCEPT or REJECT) with the key statistics: "
    "defective count vs c (attributes) or Q-statistic vs k (variables). "
    "The PNG is a visual lot disposition record saved to ~/Downloads/ with a timestamp prefix. "
    "Retain the PNG as objective evidence of the lot disposition decision."
)

warn(doc, "IMPORTANT:",
    "jrc_as_evaluate applies the plan you specify. It is your responsibility to confirm "
    "that the plan parameters (n, c or n, k) match the validated plan from jrc_as_attributes "
    "or jrc_as_variables. Mismatched parameters will produce incorrect verdicts.")

# ---------------------------------------------------------------------------
# Section 5 — Preparing Your Data
# ---------------------------------------------------------------------------
h1(doc, "5. Preparing Your Data")
para(doc,
    "Plan design scripts (jrc_as_attributes, jrc_as_variables, jrc_as_oc_curve) require "
    "no data file \u2014 all inputs are command-line arguments. The evaluation script "
    "(jrc_as_evaluate) requires a CSV file in the standard jrc two-column format.")
doc.add_paragraph()

csv_rows = [
    ("jrc_as_evaluate (attributes)", "id, result", "result: 0 = conforming, 1 = defective"),
    ("jrc_as_evaluate (variables)",  "id, value",  "value: continuous measurement"),
]

tbl = doc.add_table(rows=len(csv_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["Script / mode", "Required columns", "Notes"], [2.3, 1.7, 3.0])
for i, row_data in enumerate(csv_rows, start=1):
    data_row(tbl, i, list(row_data), bold_first=True)
doc.add_paragraph()

para(doc,
    "Column names are case-insensitive. Extra columns are ignored. The \u2018id\u2019 column may "
    "contain text labels or numbers. For attributes mode, the \u2018result\u2019 column must contain "
    "only 0 or 1. For variables mode, the \u2018value\u2019 column must be numeric. "
    "The script will exit with a clear error message if required columns are missing.")

tip(doc, "TIP:",
    "Sample data files are provided in repos/as/sample_data/: "
    "as_attributes_example.csv (50 units, 3 defectives) and "
    "as_variables_example.csv (20 measurements, mean \u2248 10.5). "
    "Use these to test the evaluate script before running it on production data.")

# ---------------------------------------------------------------------------
# Section 6 — Statistical Glossary
# ---------------------------------------------------------------------------
h1(doc, "6. Statistical Glossary")

gloss = [
    ("AQL", "Acceptable Quality Level. The fraction nonconforming at which the lot is \u2018good\u2019 and should be accepted with high probability (1\u2212\u03b1)."),
    ("RQL / LTPD", "Rejectable Quality Level (also Lot Tolerance Percent Defective). The fraction nonconforming at which the lot should be rejected with high probability (1\u2212\u03b2)."),
    ("Producer\u2019s risk (\u03b1)", "Probability of rejecting a lot at AQL quality. The probability of a false rejection. Default: 0.05 (5%)."),
    ("Consumer\u2019s risk (\u03b2)", "Probability of accepting a lot at RQL quality. The probability of a false acceptance. Default: 0.10 (10%)."),
    ("n", "Sample size. Number of units drawn from the lot for inspection."),
    ("c", "Acceptance number (attributes plan). Accept the lot if the count of defectives \u2264 c; reject if > c."),
    ("k", "Acceptability constant (variables plan). Accept the lot if the Q-statistic \u2265 k."),
    ("Q-statistic", "For a one-sided lower spec: Q_L = (\u0305x \u2212 LSL) / s. For a one-sided upper spec: Q_U = (USL \u2212 \u0305x) / s. Accept if Q \u2265 k."),
    ("OC curve", "Operating Characteristic curve. Plots Pa (probability of accepting the lot) vs p (lot fraction nonconforming). Summarises plan performance at all quality levels."),
    ("Pa", "Probability of Acceptance. The probability that the sampling plan will accept a lot with fraction nonconforming p."),
    ("ASN", "Average Sample Number. For double sampling, the expected number of units inspected per lot. ASN < n of the equivalent single plan at AQL."),
    ("Double sampling", "A two-stage plan: inspect n1 units; accept immediately if defectives \u2264 c1; reject if defectives > c2; otherwise inspect n2 more units and accept if total defectives \u2264 c2."),
    ("k-method", "Variables plan method that works with unknown process sigma. The sample standard deviation s is used in place of \u03c3. Uses the non-central t distribution for exact risk control."),
]

tbl = doc.add_table(rows=len(gloss) + 1, cols=2)
tbl.style = "Table Grid"
header_row(tbl, ["Term", "Definition"], [1.5, 5.4])
for i, (term, defn) in enumerate(gloss, start=1):
    data_row(tbl, i, [term, defn], bold_first=True)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
