"""
Generates spc_validation_report.docx — FDA-acceptable OQ Validation Report
for the JR Validated Environment SPC module v1.0.

Run from the repo root:
    python3 repos/spc/docs/ignore/generate_spc_validation_report.py
Output: repos/spc/docs/ignore/spc_validation_report.docx
"""

import os
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "spc_validation_report.docx")

DARK  = "1A1A2E"
GREY  = "F2F2F2"
WHITE = "FFFFFF"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def header_row(table, *texts):
    row = table.rows[0]
    for cell, text in zip(row.cells, texts):
        set_cell_shading(cell, DARK)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)


def data_row(table, *texts, alt=True):
    row = table.add_row()
    fill = GREY if alt else WHITE
    for cell, text in zip(row.cells, texts):
        set_cell_shading(cell, fill)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text))
        run.font.size = Pt(9)
    return row


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def para(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = docx.Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)

for sn in ("Normal", "Heading 1", "Heading 2", "List Bullet", "List Number"):
    try:
        s = doc.styles[sn]
        s.font.name = "Calibri"
        if sn == "Normal":
            s.font.size = Pt(10)
        elif sn in ("Heading 1", "Heading 2"):
            s.font.bold = True
            s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            s.font.size = Pt(14 if sn == "Heading 1" else 13)
    except KeyError:
        pass

# ===========================================================================
# COVER PAGE
# ===========================================================================

doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("VALIDATION REPORT\nOPERATIONAL QUALIFICATION\nSPC MODULE")
r.bold = True
r.font.size = Pt(18)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("JR Validated Environment — SPC Module v1.0")
r2.bold = True
r2.font.size = Pt(13)
r2.font.name = "Calibri"

doc.add_paragraph()

tbl_cover = doc.add_table(rows=1, cols=2)
tbl_cover.style = "Table Grid"
header_row(tbl_cover, "Field", "Value")
cover_rows = [
    ("Document Number",   "JR-VR-SPC-001"),
    ("Title",             "Validation Report — Operational Qualification, SPC Module"),
    ("Validation Plan",   "JR-VP-SPC-001 v1.0"),
    ("System",            "JR Validated Environment — SPC Module"),
    ("Module Version",    "1.0"),
    ("Document Version",  "1.0"),
    ("Execution Date",    "2026-03-18"),
    ("OQ Result",         "PASS — 55 / 55 tests passed, 0 failures, 0 deviations"),
    ("Author",            "Joep Rous"),
    ("Reviewer",          "[Name]"),
    ("Approver",          "[Name]"),
]
for i, (f, v) in enumerate(cover_rows):
    data_row(tbl_cover, f, v, alt=(i % 2 == 0))
set_col_widths(tbl_cover, [5, 11])

doc.add_page_break()

# ===========================================================================
# VERSION HISTORY
# ===========================================================================

h1(doc, "Version History")
tbl_ver = doc.add_table(rows=1, cols=4)
tbl_ver.style = "Table Grid"
header_row(tbl_ver, "Version", "Date", "Author", "Description")
data_row(tbl_ver, "1.0", "2026-03-18", "Joep Rous",
         "Initial OQ execution report. All 55 test cases passed. No deviations.",
         alt=True)
set_col_widths(tbl_ver, [2, 3, 4, 8])

doc.add_paragraph()

# ===========================================================================
# APPROVAL SIGNATURES
# ===========================================================================

h1(doc, "Approval Signatures")
tbl_appr = doc.add_table(rows=1, cols=4)
tbl_appr.style = "Table Grid"
header_row(tbl_appr, "Role", "Name", "Signature", "Date")
data_row(tbl_appr, "Author",        "Joep Rous", "", "2026-03-18", alt=True)
data_row(tbl_appr, "Reviewer",      "[Name]",    "", "",            alt=False)
data_row(tbl_appr, "Approver (QA)", "[Name]",    "", "",            alt=True)
set_col_widths(tbl_appr, [4, 4, 5, 4])

doc.add_page_break()

# ===========================================================================
# 1. PURPOSE
# ===========================================================================

h1(doc, "1.  Purpose")
para(doc,
    "This document reports the execution and outcome of the Operational Qualification "
    "(OQ) test suite for the JR Validated Environment SPC Module, version 1.0. It "
    "provides objective evidence that all SPC scripts perform as specified in the OQ "
    "Validation Plan JR-VP-SPC-001 v1.0 under the defined test conditions.")
para(doc,
    "This report records the test environment, execution results for all 55 test cases, "
    "and the formal OQ conclusion.")

# ===========================================================================
# 2. SCOPE
# ===========================================================================

h1(doc, "2.  Scope")
h2(doc, "2.1  In Scope")
for item in [
    "All 5 R scripts in repos/spc/R/ (jrc_spc_imr, jrc_spc_xbar_r, jrc_spc_xbar_s, "
    "jrc_spc_p, jrc_spc_c)",
    "Correct computation of control limits and all 8 Western Electric rule violations "
    "for valid inputs",
    "Correct rejection of invalid inputs with informative error messages",
    "Bypass-protection verification (RENV_PATHS_ROOT check)",
    "PNG output to ~/Downloads/",
]:
    bullet(doc, item)

h2(doc, "2.2  Out of Scope")
for item in [
    "Infrastructure scripts in bin/ and admin/ (covered by JR-IQ-001)",
    "Performance qualification (PQ) with real production data",
    "Validation of R interpreter and third-party packages (ggplot2, grid)",
]:
    bullet(doc, item)

# ===========================================================================
# 3. REFERENCE DOCUMENTS
# ===========================================================================

h1(doc, "3.  Reference Documents")
doc.add_paragraph()
tbl_ref = doc.add_table(rows=1, cols=3)
tbl_ref.style = "Table Grid"
header_row(tbl_ref, "Document ID", "Title", "Location")
refs = [
    ("JR-VP-SPC-001 v1.0",
     "Validation Plan — OQ, SPC Module",
     "repos/spc/docs/ignore/spc_validation_plan.docx"),
    ("JR-VP-002 v1.0",
     "Validation Plan — OQ, Community Script Suite v1.1.0",
     "docs/ignore/oq_validation_plan.docx"),
    ("JR-IQ-001",
     "IQ Execution Evidence",
     "docs/IQ_validation_20260311_205146.txt"),
    ("OQ Evidence",
     "SPC OQ Execution Evidence (pytest output + header)",
     "~/.jrscript/MyProject/validation/spc_oq_execution_<timestamp>.txt"),
]
for i, row in enumerate(refs):
    data_row(tbl_ref, *row, alt=(i % 2 == 0))
set_col_widths(tbl_ref, [4, 7, 6])

# ===========================================================================
# 4. TEST ENVIRONMENT
# ===========================================================================

h1(doc, "4.  Test Environment")
doc.add_paragraph()
tbl_env = doc.add_table(rows=1, cols=2)
tbl_env.style = "Table Grid"
header_row(tbl_env, "Item", "Value")
env_rows = [
    ("Execution date / time",       "2026-03-18T19:35:23"),
    ("Completion date / time",      "2026-03-18T19:35:49"),
    ("Host",                        "Joeps-MacBook-Air.local"),
    ("Operating system",            "macOS 15.7.3 (Darwin 24.6.0, aarch64)"),
    ("R version",                   "4.5.0"),
    ("Python version",              "3.11.9"),
    ("pytest version",              "8.3.4"),
    ("OQ virtual environment",      "~/.venvs/MyProject_oq/  (reused from core OQ suite)"),
    ("Script entry point",          "bin/jrrun  (sets RENV_PATHS_ROOT, loads renv library)"),
    ("Test runner",                 "repos/spc/admin_spc_oq"),
    ("Test directory",              "repos/spc/oq/"),
    ("Test data directory",         "repos/spc/oq/data/  (17 committed synthetic CSV files)"),
    ("Evidence file",               "~/.jrscript/MyProject/validation/"
                                    "spc_oq_execution_<timestamp>.txt"),
]
for i, row in enumerate(env_rows):
    data_row(tbl_env, *row, alt=(i % 2 == 0))
set_col_widths(tbl_env, [6, 11])

# ===========================================================================
# 5. TEST DATA
# ===========================================================================

h1(doc, "5.  Test Data")
para(doc,
    "All test data files are synthetic CSV files committed to repos/spc/oq/data/ "
    "and included in the project integrity check (admin/project_integrity.sha256). "
    "No real patient or production data is used.")
doc.add_paragraph()
tbl_data = doc.add_table(rows=1, cols=3)
tbl_data.style = "Table Grid"
header_row(tbl_data, "File", "Content", "Used by")
data_files = [
    ("imr_stable.csv",
     "25 observations, id/value, mean \u2248 10.07, no WE rule fires",
     "TC-SPC-IMR-001, 002, 005, 006"),
    ("imr_ooc.csv",
     "Same as stable but obs 13 = 11.50 (Rule 1 violation, UCL \u2248 11.27)",
     "TC-SPC-IMR-003, 004"),
    ("imr_missing_col.csv",
     "Missing 'value' column",
     "TC-SPC-IMR-009"),
    ("imr_one_obs.csv",
     "Only 1 observation (< 2 required for moving range)",
     "TC-SPC-IMR-010"),
    ("xbar_r_stable.csv",
     "20 subgroups \u00d7 n = 5, long format; no WE rule fires",
     "TC-SPC-XBR-001, 002, 005, 006"),
    ("xbar_r_ooc.csv",
     "Same but sg09 replaced with mean \u2248 53 (Rule 1 violation)",
     "TC-SPC-XBR-003, 004"),
    ("xbar_r_unbalanced.csv",
     "sg01 n = 5, sg02 n = 4, sg03 n = 5 (unbalanced)",
     "TC-SPC-XBR-010, TC-SPC-XBS-010"),
    ("xbar_r_n_too_large.csv",
     "3 subgroups \u00d7 n = 11 (exceeds X-bar/R maximum)",
     "TC-SPC-XBR-011"),
    ("xbar_r_missing_col.csv",
     "Missing 'value' column",
     "TC-SPC-XBR-009"),
    ("xbar_s_stable.csv",
     "20 subgroups \u00d7 n = 8; three spread types to prevent Rule 7; no WE rule fires",
     "TC-SPC-XBS-001, 002, 005, 006"),
    ("xbar_s_ooc.csv",
     "Same but sg09 replaced with mean \u2248 102 (Rule 1 violation)",
     "TC-SPC-XBS-003, 004"),
    ("xbar_s_missing_col.csv",
     "Missing 'value' column",
     "TC-SPC-XBS-009"),
    ("p_stable.csv",
     "25 subgroups, n = 100 each, defectives 1\u20138; p_bar = 0.0416; no WE rule fires",
     "TC-SPC-P-001, 002, 003, 006"),
    ("p_ooc.csv",
     "Same but subgroup 13 = 15 defectives (p = 0.15 > UCL \u2248 0.10)",
     "TC-SPC-P-004, 005"),
    ("p_missing_col.csv",
     "Missing 'defectives' column",
     "TC-SPC-P-009"),
    ("p_invalid.csv",
     "Row with defectives = 101 > n = 100",
     "TC-SPC-P-010"),
    ("c_stable.csv",
     "25 subgroups; defects 1\u20139; c_bar \u2248 4.96; no WE rule fires",
     "TC-SPC-C-001, 002, 003, 006"),
    ("c_ooc.csv",
     "Same but subgroup 12 = 15 defects (> UCL \u2248 12.2)",
     "TC-SPC-C-004, 005"),
    ("c_missing_col.csv",
     "Missing 'defects' column",
     "TC-SPC-C-009"),
]
for i, row in enumerate(data_files):
    data_row(tbl_data, *row, alt=(i % 2 == 0))
set_col_widths(tbl_data, [5, 6.5, 5.5])

# ===========================================================================
# 6. TEST EXECUTION RESULTS
# ===========================================================================

h1(doc, "6.  Test Execution Results")
h2(doc, "6.1  Summary")
doc.add_paragraph()
tbl_sum = doc.add_table(rows=1, cols=2)
tbl_sum.style = "Table Grid"
header_row(tbl_sum, "Item", "Value")
summary = [
    ("Total test cases",    "55"),
    ("Passed",              "55"),
    ("Failed",              "0"),
    ("Errors",              "0"),
    ("Duration",            "25.59 s"),
    ("OQ result",           "PASS"),
]
for i, row in enumerate(summary):
    data_row(tbl_sum, *row, alt=(i % 2 == 0))
set_col_widths(tbl_sum, [6, 11])

doc.add_paragraph()
h2(doc, "6.2  Results by Test File")
doc.add_paragraph()
tbl_byfile = doc.add_table(rows=1, cols=4)
tbl_byfile.style = "Table Grid"
header_row(tbl_byfile, "Test file", "Script(s) covered", "Tests", "Result")
byfile = [
    ("test_spc_imr.py",    "jrc_spc_imr",    "11", "11 / 11  PASS"),
    ("test_spc_xbar_r.py", "jrc_spc_xbar_r", "12", "12 / 12  PASS"),
    ("test_spc_xbar_s.py", "jrc_spc_xbar_s", "11", "11 / 11  PASS"),
    ("test_spc_p.py",      "jrc_spc_p",      "11", "11 / 11  PASS"),
    ("test_spc_c.py",      "jrc_spc_c",      "10", "10 / 10  PASS"),
]
for i, row in enumerate(byfile):
    data_row(tbl_byfile, *row, alt=(i % 2 == 0))
set_col_widths(tbl_byfile, [5.5, 5, 2, 4.5])

doc.add_paragraph()
h2(doc, "6.3  Individual Test Case Results")
para(doc,
    "All 55 test cases passed. The full pytest output, including individual test IDs, "
    "pass/fail status, and timing, is recorded in the evidence file:")
para(doc,
    "    ~/.jrscript/MyProject/validation/spc_oq_execution_<timestamp>.txt",
    bold=True)
para(doc,
    "That file also contains the evidence header (host, R version, Python version, "
    "pytest version, timestamp) and the OQ RESULT: PASS footer. The per-test "
    "results are reproduced in the table below for traceability.")

doc.add_paragraph()
tbl_tc = doc.add_table(rows=1, cols=4)
tbl_tc.style = "Table Grid"
header_row(tbl_tc, "Test Case ID", "Description", "Expected", "Result")
tc_rows = [
    # --- IMR ---
    ("TC-SPC-IMR-001", "Happy path — exit 0, sections present",           "Exit 0, sections present",         "PASS"),
    ("TC-SPC-IMR-002", "Stable data — IN CONTROL verdict",                "IN CONTROL or STABLE",             "PASS"),
    ("TC-SPC-IMR-003", "OOC data — exit 0, OUT OF CONTROL",               "Exit 0, OUT OF CONTROL/SIGNALS",   "PASS"),
    ("TC-SPC-IMR-004", "OOC data — obs 13 flagged as Rule 1",             "'13' and '[1]' in output",         "PASS"),
    ("TC-SPC-IMR-005", "--ucl / --lcl flags accepted",                    "Exit 0, limit value in output",    "PASS"),
    ("TC-SPC-IMR-006", "PNG written to ~/Downloads/",                     "PNG present, recent mtime",        "PASS"),
    ("TC-SPC-IMR-007", "No arguments \u2192 usage message",               "Exit \u2260 0, 'Usage' in output", "PASS"),
    ("TC-SPC-IMR-008", "File not found",                                  "Exit \u2260 0, 'not found'",       "PASS"),
    ("TC-SPC-IMR-009", "Missing 'value' column",                          "Exit \u2260 0, 'value' in output", "PASS"),
    ("TC-SPC-IMR-010", "Too few observations (1 row)",                    "Exit \u2260 0",                    "PASS"),
    ("TC-SPC-IMR-011", "Bypass protection — direct Rscript call",         "Exit \u2260 0, RENV_PATHS_ROOT",   "PASS"),
    # --- XBar-R ---
    ("TC-SPC-XBR-001", "Happy path — exit 0, sections present",           "Exit 0, sections present",         "PASS"),
    ("TC-SPC-XBR-002", "Stable data — IN CONTROL verdict",                "IN CONTROL or STABLE",             "PASS"),
    ("TC-SPC-XBR-003", "OOC data — exit 0, OUT OF CONTROL",               "Exit 0, OUT OF CONTROL/SIGNALS",   "PASS"),
    ("TC-SPC-XBR-004", "OOC data — sg09 flagged",                         "'sg09' in output",                 "PASS"),
    ("TC-SPC-XBR-005", "--ucl / --lcl flags accepted",                    "Exit 0, limit value in output",    "PASS"),
    ("TC-SPC-XBR-006", "PNG written to ~/Downloads/",                     "PNG present, recent mtime",        "PASS"),
    ("TC-SPC-XBR-007", "No arguments \u2192 usage message",               "Exit \u2260 0, 'Usage' in output", "PASS"),
    ("TC-SPC-XBR-008", "File not found",                                  "Exit \u2260 0, 'not found'",       "PASS"),
    ("TC-SPC-XBR-009", "Missing 'value' column",                          "Exit \u2260 0, 'value' in output", "PASS"),
    ("TC-SPC-XBR-010", "Unbalanced subgroups",                            "Exit \u2260 0, 'unbalanced'",      "PASS"),
    ("TC-SPC-XBR-011", "Subgroup size n > 10",                            "Exit \u2260 0, suggests xbar_s",   "PASS"),
    ("TC-SPC-XBR-012", "Bypass protection — direct Rscript call",         "Exit \u2260 0, RENV_PATHS_ROOT",   "PASS"),
    # --- XBar-S ---
    ("TC-SPC-XBS-001", "Happy path — exit 0, sections present",           "Exit 0, sections present",         "PASS"),
    ("TC-SPC-XBS-002", "Stable data — IN CONTROL verdict",                "IN CONTROL or STABLE",             "PASS"),
    ("TC-SPC-XBS-003", "OOC data — exit 0, OUT OF CONTROL",               "Exit 0, OUT OF CONTROL/SIGNALS",   "PASS"),
    ("TC-SPC-XBS-004", "OOC data — sg09 flagged",                         "'sg09' in output",                 "PASS"),
    ("TC-SPC-XBS-005", "--ucl / --lcl flags accepted",                    "Exit 0, limit value in output",    "PASS"),
    ("TC-SPC-XBS-006", "PNG written to ~/Downloads/",                     "PNG present, recent mtime",        "PASS"),
    ("TC-SPC-XBS-007", "No arguments \u2192 usage message",               "Exit \u2260 0, 'Usage' in output", "PASS"),
    ("TC-SPC-XBS-008", "File not found",                                  "Exit \u2260 0, 'not found'",       "PASS"),
    ("TC-SPC-XBS-009", "Missing 'value' column",                          "Exit \u2260 0, 'value' in output", "PASS"),
    ("TC-SPC-XBS-010", "Unbalanced subgroups",                            "Exit \u2260 0, 'unbalanced'",      "PASS"),
    ("TC-SPC-XBS-011", "Bypass protection — direct Rscript call",         "Exit \u2260 0, RENV_PATHS_ROOT",   "PASS"),
    # --- P-chart ---
    ("TC-SPC-P-001",   "Happy path — exit 0, sections present",           "Exit 0, sections present",         "PASS"),
    ("TC-SPC-P-002",   "Stable data — IN CONTROL verdict",                "IN CONTROL or STABLE",             "PASS"),
    ("TC-SPC-P-003",   "Known p-bar \u2248 0.0416",                       "Output matches 0.04[01234]",       "PASS"),
    ("TC-SPC-P-004",   "OOC data — exit 0, OUT OF CONTROL",               "Exit 0, OUT OF CONTROL/SIGNALS",   "PASS"),
    ("TC-SPC-P-005",   "OOC data — subgroup 13 flagged",                  "'13' in output",                   "PASS"),
    ("TC-SPC-P-006",   "PNG written to ~/Downloads/",                     "PNG present, recent mtime",        "PASS"),
    ("TC-SPC-P-007",   "No arguments \u2192 usage message",               "Exit \u2260 0, 'Usage' in output", "PASS"),
    ("TC-SPC-P-008",   "File not found",                                  "Exit \u2260 0, 'not found'",       "PASS"),
    ("TC-SPC-P-009",   "Missing 'defectives' column",                     "Exit \u2260 0, 'defective'",       "PASS"),
    ("TC-SPC-P-010",   "Defectives > n",                                  "Exit \u2260 0, validation error",  "PASS"),
    ("TC-SPC-P-011",   "Bypass protection — direct Rscript call",         "Exit \u2260 0, RENV_PATHS_ROOT",   "PASS"),
    # --- C-chart ---
    ("TC-SPC-C-001",   "Happy path — exit 0, sections present",           "Exit 0, sections present",         "PASS"),
    ("TC-SPC-C-002",   "Stable data — IN CONTROL verdict",                "IN CONTROL or STABLE",             "PASS"),
    ("TC-SPC-C-003",   "Known c-bar \u2248 4.96",                         "Output matches 4.8[0-9]|4.9[0-9]", "PASS"),
    ("TC-SPC-C-004",   "OOC data — exit 0, OUT OF CONTROL",               "Exit 0, OUT OF CONTROL/SIGNALS",   "PASS"),
    ("TC-SPC-C-005",   "OOC data — subgroup 12 flagged",                  "'12' in output",                   "PASS"),
    ("TC-SPC-C-006",   "PNG written to ~/Downloads/",                     "PNG present, recent mtime",        "PASS"),
    ("TC-SPC-C-007",   "No arguments \u2192 usage message",               "Exit \u2260 0, 'Usage' in output", "PASS"),
    ("TC-SPC-C-008",   "File not found",                                  "Exit \u2260 0, 'not found'",       "PASS"),
    ("TC-SPC-C-009",   "Missing 'defects' column",                        "Exit \u2260 0, 'defect'",          "PASS"),
    ("TC-SPC-C-010",   "Bypass protection — direct Rscript call",         "Exit \u2260 0, RENV_PATHS_ROOT",   "PASS"),
]
for i, row in enumerate(tc_rows):
    data_row(tbl_tc, *row, alt=(i % 2 == 0))
set_col_widths(tbl_tc, [3.5, 6, 4, 3])

# ===========================================================================
# 7. DEVIATIONS
# ===========================================================================

h1(doc, "7.  Deviations from OQ Plan (JR-VP-SPC-001 v1.0)")
para(doc, "No deviations from the approved OQ plan were identified during this OQ execution.")
doc.add_paragraph()
tbl_dev = doc.add_table(rows=1, cols=4)
tbl_dev.style = "Table Grid"
header_row(tbl_dev, "ID", "Description", "Resolution", "Status")
data_row(tbl_dev, "\u2014", "None", "\u2014", "N/A", alt=True)
set_col_widths(tbl_dev, [2, 7, 6, 2])

# ===========================================================================
# 8. REQUIREMENTS TRACEABILITY
# ===========================================================================

h1(doc, "8.  Requirements Traceability Matrix")
para(doc,
    "The table below maps each User Requirement from JR-VP-SPC-001 to the "
    "test cases executed and their result.")
doc.add_paragraph()
tbl_rtm = doc.add_table(rows=1, cols=4)
tbl_rtm.style = "Table Grid"
header_row(tbl_rtm, "UR", "Requirement (summary)", "Test Cases", "Result")
rtm = [
    ("UR-SPC-001", "jrc_spc_imr — X-bar, MR-bar, control limits",
     "TC-SPC-IMR-001, 002", "2 / 2  PASS"),
    ("UR-SPC-002", "jrc_spc_imr — WE rules, violation list",
     "TC-SPC-IMR-002, 003, 004", "3 / 3  PASS"),
    ("UR-SPC-003", "jrc_spc_imr — --ucl / --lcl",
     "TC-SPC-IMR-005", "1 / 1  PASS"),
    ("UR-SPC-004", "jrc_spc_imr — verdict",
     "TC-SPC-IMR-001, 002, 003", "3 / 3  PASS"),
    ("UR-SPC-005", "jrc_spc_imr — PNG to ~/Downloads/",
     "TC-SPC-IMR-006", "1 / 1  PASS"),
    ("UR-SPC-006", "jrc_spc_imr — reject invalid inputs",
     "TC-SPC-IMR-007, 008, 009, 010", "4 / 4  PASS"),
    ("UR-SPC-007", "jrc_spc_xbar_r — grand mean, R-bar, limits",
     "TC-SPC-XBR-001, 002", "2 / 2  PASS"),
    ("UR-SPC-008", "jrc_spc_xbar_r — WE rules, violation list",
     "TC-SPC-XBR-002, 003, 004", "3 / 3  PASS"),
    ("UR-SPC-009", "jrc_spc_xbar_r — --ucl / --lcl",
     "TC-SPC-XBR-005", "1 / 1  PASS"),
    ("UR-SPC-010", "jrc_spc_xbar_r — verdict",
     "TC-SPC-XBR-001, 002, 003", "3 / 3  PASS"),
    ("UR-SPC-011", "jrc_spc_xbar_r — PNG to ~/Downloads/",
     "TC-SPC-XBR-006", "1 / 1  PASS"),
    ("UR-SPC-012", "jrc_spc_xbar_r — reject unbalanced / n > 10",
     "TC-SPC-XBR-010, 011", "2 / 2  PASS"),
    ("UR-SPC-013", "jrc_spc_xbar_s — grand mean, S-bar, limits via c4",
     "TC-SPC-XBS-001, 002", "2 / 2  PASS"),
    ("UR-SPC-014", "jrc_spc_xbar_s — WE rules, violation list",
     "TC-SPC-XBS-002, 003, 004", "3 / 3  PASS"),
    ("UR-SPC-015", "jrc_spc_xbar_s — --ucl / --lcl",
     "TC-SPC-XBS-005", "1 / 1  PASS"),
    ("UR-SPC-016", "jrc_spc_xbar_s — verdict",
     "TC-SPC-XBS-001, 002, 003", "3 / 3  PASS"),
    ("UR-SPC-017", "jrc_spc_xbar_s — PNG to ~/Downloads/",
     "TC-SPC-XBS-006", "1 / 1  PASS"),
    ("UR-SPC-018", "jrc_spc_xbar_s — reject unbalanced subgroups",
     "TC-SPC-XBS-010", "1 / 1  PASS"),
    ("UR-SPC-019", "jrc_spc_p — p-bar, variable limits, WE rules",
     "TC-SPC-P-001, 002, 003, 004, 005", "5 / 5  PASS"),
    ("UR-SPC-020", "jrc_spc_p — verdict",
     "TC-SPC-P-001, 002, 004", "3 / 3  PASS"),
    ("UR-SPC-021", "jrc_spc_p — reject defectives > n",
     "TC-SPC-P-010", "1 / 1  PASS"),
    ("UR-SPC-022", "jrc_spc_p — PNG to ~/Downloads/",
     "TC-SPC-P-006", "1 / 1  PASS"),
    ("UR-SPC-023", "jrc_spc_c — c-bar, sigma, limits, WE rules",
     "TC-SPC-C-001, 002, 003, 004, 005", "5 / 5  PASS"),
    ("UR-SPC-024", "jrc_spc_c — verdict; PNG to ~/Downloads/",
     "TC-SPC-C-001, 002, 004, 006", "4 / 4  PASS"),
    ("UR-SPC-025", "All scripts — RENV_PATHS_ROOT bypass protection",
     "TC-SPC-IMR-011, XBR-012, XBS-011, P-011, C-010", "5 / 5  PASS"),
    ("UR-SPC-026", "All scripts — no-arguments usage message",
     "TC-SPC-IMR-007, XBR-007, XBS-007, P-007, C-007", "5 / 5  PASS"),
    ("UR-SPC-027", "All scripts — file not found error",
     "TC-SPC-IMR-008, XBR-008, XBS-008, P-008, C-008", "5 / 5  PASS"),
]
for i, row in enumerate(rtm):
    data_row(tbl_rtm, *row, alt=(i % 2 == 0))
set_col_widths(tbl_rtm, [2.5, 7, 5, 2.5])

# ===========================================================================
# 9. OQ CONCLUSION
# ===========================================================================

h1(doc, "9.  OQ Conclusion")
para(doc,
    "The Operational Qualification of the JR Validated Environment SPC Module "
    "v1.0 is complete. All 55 test cases defined in JR-VP-SPC-001 were executed "
    "on 2026-03-18 and passed with no failures, no errors, and no deviations.")
doc.add_paragraph()
para(doc,
    "All 27 user requirements (UR-SPC-001 through UR-SPC-027) are covered by at "
    "least one passing test case. The OQ acceptance criteria stated in "
    "JR-VP-SPC-001 Section 7 are fully met.")
doc.add_paragraph()
para(doc,
    "The SPC module is hereby declared OPERATIONALLY QUALIFIED and released for "
    "use in process monitoring and control activities in accordance with the JR "
    "Validated Environment quality system.", bold=True)
doc.add_paragraph()
para(doc,
    "OQ Evidence file: "
    "~/.jrscript/MyProject/validation/spc_oq_execution_<timestamp>.txt",
    italic=True)

# ===========================================================================
# Save
# ===========================================================================

doc.save(OUT)
print(f"Saved: {OUT}")
