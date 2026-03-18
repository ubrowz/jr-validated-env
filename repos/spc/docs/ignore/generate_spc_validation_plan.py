"""
Generates spc_validation_plan.docx — FDA-acceptable OQ Validation Plan
for the JR Validated Environment SPC module v1.0.

Run from the repo root:
    python3 repos/spc/docs/ignore/generate_spc_validation_plan.py
Output: repos/spc/docs/ignore/spc_validation_plan.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "spc_validation_plan.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text="", bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def shade_cell(cell, hex_fill, font_color=None):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)
    if font_color:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = font_color


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, "1A1A2E", RGBColor(0xFF, 0xFF, 0xFF))
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            shade_cell(cell, bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_tc_block(doc, tc_id, tests_ur, command, rationale, pass_criterion):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tc_id)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x21, 0x66, 0xAC)

    def labeled_line(label, value):
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Cm(0.8)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        r1 = lp.add_run(label + ": ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = lp.add_run(value)
        r2.font.size = Pt(9)
        if label == "Command":
            r2.font.name = "Courier New"

    labeled_line("Tests", tests_ur)
    labeled_line("Command", command)
    if rationale:
        labeled_line("Rationale", rationale)
    labeled_line("Pass criterion", pass_criterion)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ---------------------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------------------

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("VALIDATION PLAN\nOPERATIONAL QUALIFICATION\nSPC MODULE")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("JR Validated Environment — SPC Module v1.0")
r2.font.size = Pt(13)
r2.bold = True

doc.add_paragraph()

add_table(doc,
    ["Field", "Value"],
    [
        ("Document Number",  "JR-VP-SPC-001"),
        ("Title",            "Validation Plan — Operational Qualification, SPC Module"),
        ("System",           "JR Validated Environment — SPC Module"),
        ("Module Version",   "1.0"),
        ("Document Version", "1.0"),
        ("Status",           "Draft"),
        ("Effective Date",   "2026-03-18"),
        ("Author",           "Joep Rous"),
        ("Reviewer",         "[Name]"),
        ("Approver",         "[Name]"),
    ],
    col_widths=[5, 11]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# VERSION HISTORY
# ---------------------------------------------------------------------------

add_heading(doc, "Version History", level=1)
add_table(doc,
    ["Version", "Date", "Author", "Description"],
    [("1.0", "2026-03-18", "Joep Rous",
      "Initial release. Covers all 5 SPC scripts in module v1.0.")],
    col_widths=[2, 3, 4, 8]
)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# APPROVAL SIGNATURES
# ---------------------------------------------------------------------------

add_heading(doc, "Approval Signatures", level=1)
add_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ("Author",        "", "", ""),
        ("Reviewer",      "", "", ""),
        ("Approver (QA)", "", "", ""),
    ],
    col_widths=[4, 5, 5, 3]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. PURPOSE
# ---------------------------------------------------------------------------

add_heading(doc, "1.  Purpose", level=1)
add_para(doc,
    "This document defines the Validation Plan for the Operational Qualification (OQ) "
    "of the Statistical Process Control (SPC) module included in version 1.0 of the "
    "JR Validated Environment. It specifies the user requirements, test cases, "
    "acceptance criteria, and traceability between requirements and tests needed to "
    "demonstrate that each SPC script performs correctly for its intended use.")
add_para(doc,
    "Execution of the tests defined in this plan, together with satisfactory results "
    "and documented evidence, constitutes the OQ for the SPC module.")

# ---------------------------------------------------------------------------
# 2. SCOPE
# ---------------------------------------------------------------------------

add_heading(doc, "2.  Scope", level=1)
add_heading(doc, "2.1  In Scope", level=2)

in_scope = [
    "All 5 R scripts in the repos/spc/R/ directory: jrc_spc_imr, jrc_spc_xbar_r, "
    "jrc_spc_xbar_s, jrc_spc_p, jrc_spc_c",
    "Correct computation of control limits and all 8 Western Electric rule violations "
    "for valid inputs",
    "Correct rejection of invalid inputs with informative error messages",
    "Correct file I/O behaviour (CSV input, PNG output to ~/Downloads/)",
    "Verification that all R scripts enforce execution within the jrrun validated "
    "environment by checking RENV_PATHS_ROOT at startup",
    "Balanced-design validation for subgrouped charts (X-bar/R, X-bar/S)",
    "Subgroup size constraint enforcement (X-bar/R: n = 2–10; X-bar/S: n ≥ 2)",
    "Defectives-exceed-n validation for P-chart",
]
for item in in_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

add_heading(doc, "2.2  Out of Scope", level=2)
out_scope = [
    "Infrastructure scripts in bin/ and admin/ (covered by IQ; see JR-VP-001)",
    "Performance qualification (PQ) — production process monitoring studies are "
    "conducted separately under individual process control protocols",
    "User acceptance testing (UAT)",
    "Validation of R language interpreter and third-party packages (ggplot2, grid) — "
    "package integrity is confirmed at installation via SHA256 (JR-IQ-001)",
    "Statistical validation of the underlying SPC methods — method references are "
    "cited in Section 3",
]
for item in out_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 3. REGULATORY BASIS
# ---------------------------------------------------------------------------

add_heading(doc, "3.  Regulatory Basis and References", level=1)
add_heading(doc, "3.1  Regulations", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("21 CFR Part 820, §820.70(i)", "Quality System Regulation — Automated Data Processing"),
        ("21 CFR Part 11",              "Electronic Records; Electronic Signatures"),
    ],
    col_widths=[5, 12]
)

doc.add_paragraph()
add_heading(doc, "3.2  FDA Guidance", level=2)
add_table(doc,
    ["Reference", "Title", "Date"],
    [
        ("FDA GPSV",
         "General Principles of Software Validation; Final Guidance for Industry and FDA Staff",
         "January 2002"),
        ("FDA CSA",
         "Computer Software Assurance for Production and Quality System Software; Draft Guidance",
         "September 2022"),
    ],
    col_widths=[3, 11, 3]
)

doc.add_paragraph()
add_heading(doc, "3.3  Standards and Methods", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("AIAG SPC, 2nd Ed.",
         "Statistical Process Control Reference Manual, Automotive Industry Action Group, 2005"),
        ("Montgomery (2020)",
         "Montgomery DC. Introduction to Statistical Quality Control, 8th Ed. Wiley."),
        ("Western Electric (1956)",
         "Statistical Quality Control Handbook. Western Electric Company, Indianapolis."),
        ("ISO 7870-2:2023",
         "Control charts — Part 2: Shewhart control charts"),
        ("ISO 7870-3:2012",
         "Control charts — Part 3: Acceptance control charts"),
        ("ISO 13485:2016",
         "Medical devices — Quality management systems — Requirements for regulatory purposes"),
    ],
    col_widths=[4, 13]
)

doc.add_paragraph()
add_heading(doc, "3.4  Internal Documents", level=2)
add_table(doc,
    ["Document Number", "Title"],
    [
        ("JR-VP-001", "Validation Plan — Installation Qualification, JR Validated Environment v1.0.0"),
        ("JR-IQ-001", "IQ Execution Evidence (docs/IQ_validation_20260311_205146.txt)"),
        ("JR-VP-002", "Validation Plan — OQ, Community Script Suite v1.1.0"),
    ],
    col_widths=[4, 13]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 4. USER REQUIREMENTS
# ---------------------------------------------------------------------------

add_heading(doc, "4.  User Requirements", level=1)
add_para(doc,
    "The following user requirements define the functional behaviour that the SPC "
    "module must satisfy. Each requirement is traced to one or more test cases in "
    "Section 6.")
doc.add_paragraph()

ur_rows = [
    # IMR
    ("UR-SPC-001", "jrc_spc_imr",
     "Compute X-bar, MR-bar, UCL and LCL for both the Individuals chart and the "
     "Moving Range chart from a valid id/value CSV. Report all values in terminal output."),
    ("UR-SPC-002", "jrc_spc_imr",
     "Apply all 8 Western Electric rules to the Individuals chart and Rule 1 only "
     "to the Moving Range chart. List each out-of-control observation by ID, value, "
     "and rule(s) fired."),
    ("UR-SPC-003", "jrc_spc_imr",
     "Accept optional --ucl and --lcl arguments. When provided, use user-specified "
     "limits in place of computed limits and label them as user-specified in output."),
    ("UR-SPC-004", "jrc_spc_imr",
     "Output a clear IN CONTROL or OUT OF CONTROL verdict based on WE rule results."),
    ("UR-SPC-005", "jrc_spc_imr",
     "Save a two-panel PNG (Individuals chart + MR chart) to ~/Downloads/ with a "
     "datetime prefix."),
    ("UR-SPC-006", "jrc_spc_imr",
     "Exit non-zero with an informative error message when: no arguments are "
     "supplied; the input file is not found; a required column is missing; "
     "fewer than 2 observations are present."),
    # XBAR-R
    ("UR-SPC-007", "jrc_spc_xbar_r",
     "Compute grand mean, R-bar, and X-bar/R control limits using tabulated A2, D3, "
     "D4 constants for n = 2–10 from a valid subgroup/value long-format CSV."),
    ("UR-SPC-008", "jrc_spc_xbar_r",
     "Apply all 8 Western Electric rules to the X-bar chart and Rule 1 only to the "
     "R chart. List each out-of-control subgroup by label, value, and rule(s) fired."),
    ("UR-SPC-009", "jrc_spc_xbar_r",
     "Accept optional --ucl and --lcl arguments for user-specified X-bar limits."),
    ("UR-SPC-010", "jrc_spc_xbar_r",
     "Output a clear IN CONTROL or OUT OF CONTROL verdict."),
    ("UR-SPC-011", "jrc_spc_xbar_r",
     "Save a two-panel PNG (X-bar chart + R chart) to ~/Downloads/ with a datetime prefix."),
    ("UR-SPC-012", "jrc_spc_xbar_r",
     "Reject unbalanced subgroups with a non-zero exit and informative message. "
     "Reject subgroup size n > 10 with a non-zero exit and suggestion to use X-bar/S."),
    # XBAR-S
    ("UR-SPC-013", "jrc_spc_xbar_s",
     "Compute grand mean, S-bar, and X-bar/S control limits using the analytical "
     "c4(n) formula (gamma function) for any n ≥ 2."),
    ("UR-SPC-014", "jrc_spc_xbar_s",
     "Apply all 8 Western Electric rules to the X-bar chart and Rule 1 only to the "
     "S chart. List each out-of-control subgroup by label, value, and rule(s) fired."),
    ("UR-SPC-015", "jrc_spc_xbar_s",
     "Accept optional --ucl and --lcl arguments for user-specified X-bar limits."),
    ("UR-SPC-016", "jrc_spc_xbar_s",
     "Output a clear IN CONTROL or OUT OF CONTROL verdict."),
    ("UR-SPC-017", "jrc_spc_xbar_s",
     "Save a two-panel PNG (X-bar chart + S chart) to ~/Downloads/ with a datetime prefix."),
    ("UR-SPC-018", "jrc_spc_xbar_s",
     "Reject unbalanced subgroups with a non-zero exit and informative message."),
    # P-chart
    ("UR-SPC-019", "jrc_spc_p",
     "Compute p-bar and per-subgroup UCL/LCL from a subgroup/n/defectives CSV. "
     "Apply all 8 Western Electric rules to standardised z-values. List violations."),
    ("UR-SPC-020", "jrc_spc_p",
     "Output a clear IN CONTROL or OUT OF CONTROL verdict."),
    ("UR-SPC-021", "jrc_spc_p",
     "Reject any row where defectives > n with a non-zero exit and informative message."),
    ("UR-SPC-022", "jrc_spc_p",
     "Save a PNG (P-chart with step-line variable control limits) to ~/Downloads/."),
    # C-chart
    ("UR-SPC-023", "jrc_spc_c",
     "Compute c-bar, sigma = sqrt(c-bar), UCL = c-bar + 3*sigma, "
     "LCL = max(0, c-bar - 3*sigma) from a subgroup/defects CSV. "
     "Apply all 8 Western Electric rules. List violations."),
    ("UR-SPC-024", "jrc_spc_c",
     "Output a clear IN CONTROL or OUT OF CONTROL verdict and save a PNG with "
     "sigma zone lines to ~/Downloads/."),
    # Cross-cutting
    ("UR-SPC-025", "All scripts",
     "Each script must check for the RENV_PATHS_ROOT environment variable at "
     "startup and exit non-zero with an informative message if it is absent, "
     "preventing execution outside the jrrun validated environment."),
    ("UR-SPC-026", "All scripts",
     "Calling any script with no arguments must exit non-zero and print a usage "
     "message showing the expected syntax."),
    ("UR-SPC-027", "All scripts",
     "Supplying a non-existent CSV path must exit non-zero with a message "
     "indicating that the file was not found."),
]

add_table(doc,
    ["UR", "Script(s)", "Requirement"],
    ur_rows,
    col_widths=[2.5, 3.5, 11]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 5. TEST ENVIRONMENT
# ---------------------------------------------------------------------------

add_heading(doc, "5.  Test Environment Requirements", level=1)
add_para(doc,
    "The following environment must be in place before executing the OQ tests.")
doc.add_paragraph()

add_table(doc,
    ["Item", "Requirement"],
    [
        ("Operating system",       "macOS (Apple Silicon or Intel)"),
        ("R version",              "4.5.x (as specified in admin/r_version.txt)"),
        ("Python version",         "3.11.x (as specified in repos/spc/oq/python_version.txt)"),
        ("pytest version",         "8.3.4 (as specified in repos/spc/oq/requirements.txt)"),
        ("OQ virtual environment", "~/.venvs/<PROJECT_ID>_oq/  (created by admin_msa_oq or admin_spc_oq)"),
        ("jrrun",                  "bin/jrrun present and executable; project integrity check passing"),
        ("RENV_PATHS_ROOT",        "Set by jrrun at execution time; must NOT be set in the test shell"),
        ("Test data",              "repos/spc/oq/data/ — 17 synthetic CSV files committed to the repo"),
        ("Test runner",            "repos/spc/admin_spc_oq"),
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 6. TEST CASES
# ---------------------------------------------------------------------------

add_heading(doc, "6.  Test Cases", level=1)
add_para(doc,
    "Each test case is implemented as a pytest function in repos/spc/oq/. "
    "Tests are executed via repos/spc/admin_spc_oq, which activates the OQ "
    "virtual environment, runs pytest with -v, and captures the output to a "
    "timestamped evidence file.")

# ── 6.1  I-MR ──────────────────────────────────────────────────────────────
add_heading(doc, "6.1  jrc_spc_imr", level=2)

add_tc_block(doc,
    "TC-SPC-IMR-001  —  Stable dataset: exit 0, key sections present",
    "UR-SPC-001, UR-SPC-004",
    "jrc_spc_imr oq/data/imr_stable.csv",
    "Verifies the script completes successfully for a well-behaved dataset.",
    "Exit code 0; output contains control limit section, process stability section, "
    "and a verdict line.")

add_tc_block(doc,
    "TC-SPC-IMR-002  —  Stable dataset: IN CONTROL verdict",
    "UR-SPC-002, UR-SPC-004",
    "jrc_spc_imr oq/data/imr_stable.csv",
    "imr_stable.csv is designed so no WE rule fires.",
    "Exit code 0; 'IN CONTROL' or 'STABLE' present in output.")

add_tc_block(doc,
    "TC-SPC-IMR-003  —  OOC dataset: exit 0, OUT OF CONTROL in output",
    "UR-SPC-002, UR-SPC-004",
    "jrc_spc_imr oq/data/imr_ooc.csv",
    "imr_ooc.csv contains observation 13 = 11.50, clearly beyond 3\u03c3 (UCL \u2248 11.27).",
    "Exit code 0; 'OUT OF CONTROL' or 'SIGNALS' present in output.")

add_tc_block(doc,
    "TC-SPC-IMR-004  —  OOC dataset: observation 13 flagged as Rule 1 violation",
    "UR-SPC-002",
    "jrc_spc_imr oq/data/imr_ooc.csv",
    "Rule 1 (beyond 3\u03c3) must fire for observation 13.",
    "Output contains '13' and one of '[1]', 'Rule 1', or 'beyond 3' (case-insensitive).")

add_tc_block(doc,
    "TC-SPC-IMR-005  —  User-specified limits accepted",
    "UR-SPC-003",
    "jrc_spc_imr oq/data/imr_stable.csv --ucl 11.0 --lcl 9.0",
    "Confirms --ucl / --lcl flags are parsed and reported.",
    "Exit code 0; output contains '11.0' or '11.00' or 'user' or 'specified'.")

add_tc_block(doc,
    "TC-SPC-IMR-006  —  PNG written to ~/Downloads/",
    "UR-SPC-005",
    "jrc_spc_imr oq/data/imr_stable.csv",
    "A PNG file must appear in ~/Downloads/ with mtime \u2265 test start time.",
    "File matching *_jrc_spc_imr.png exists in ~/Downloads/ with recent mtime.")

add_tc_block(doc,
    "TC-SPC-IMR-007  —  No arguments: non-zero exit, usage message",
    "UR-SPC-006",
    "jrc_spc_imr",
    "",
    "Exit code \u2260 0; 'Usage' or 'usage' present in output.")

add_tc_block(doc,
    "TC-SPC-IMR-008  —  File not found: non-zero exit",
    "UR-SPC-006, UR-SPC-027",
    "jrc_spc_imr /tmp/no_such_file_xyz.csv",
    "",
    "Exit code \u2260 0; 'not found' or the filename present in output.")

add_tc_block(doc,
    "TC-SPC-IMR-009  —  Missing column: non-zero exit, column named",
    "UR-SPC-006",
    "jrc_spc_imr oq/data/imr_missing_col.csv",
    "imr_missing_col.csv lacks the 'value' column.",
    "Exit code \u2260 0; 'value' present in output.")

add_tc_block(doc,
    "TC-SPC-IMR-010  —  Too few observations: non-zero exit",
    "UR-SPC-006",
    "jrc_spc_imr oq/data/imr_one_obs.csv",
    "Moving range requires at least 2 observations.",
    "Exit code \u2260 0.")

add_tc_block(doc,
    "TC-SPC-IMR-011  —  Bypass protection: direct Rscript call fails",
    "UR-SPC-025",
    "Rscript repos/spc/R/jrc_spc_imr.R oq/data/imr_stable.csv  (RENV_PATHS_ROOT unset)",
    "Script must detect absent RENV_PATHS_ROOT and refuse to run.",
    "Exit code \u2260 0; 'RENV_PATHS_ROOT' present in stderr/stdout.")

# ── 6.2  X-bar/R ────────────────────────────────────────────────────────────
add_heading(doc, "6.2  jrc_spc_xbar_r", level=2)

add_tc_block(doc,
    "TC-SPC-XBR-001  —  Stable dataset: exit 0, key sections present",
    "UR-SPC-007, UR-SPC-010",
    "jrc_spc_xbar_r oq/data/xbar_r_stable.csv",
    "Verifies the script completes successfully for a well-behaved dataset.",
    "Exit code 0; output contains X-bar/grand mean section, R-bar section, and verdict.")

add_tc_block(doc,
    "TC-SPC-XBR-002  —  Stable dataset: IN CONTROL verdict",
    "UR-SPC-008, UR-SPC-010",
    "jrc_spc_xbar_r oq/data/xbar_r_stable.csv",
    "xbar_r_stable.csv is designed so no WE rule fires.",
    "Exit code 0; 'IN CONTROL' or 'STABLE' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-003  —  OOC dataset: exit 0, OUT OF CONTROL in output",
    "UR-SPC-008, UR-SPC-010",
    "jrc_spc_xbar_r oq/data/xbar_r_ooc.csv",
    "xbar_r_ooc.csv has subgroup sg09 with mean \u2248 53, far above UCL \u2248 50.92.",
    "Exit code 0; 'OUT OF CONTROL' or 'SIGNALS' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-004  —  OOC dataset: subgroup sg09 flagged",
    "UR-SPC-008",
    "jrc_spc_xbar_r oq/data/xbar_r_ooc.csv",
    "",
    "Output contains 'sg09'.")

add_tc_block(doc,
    "TC-SPC-XBR-005  —  User-specified limits accepted",
    "UR-SPC-009",
    "jrc_spc_xbar_r oq/data/xbar_r_stable.csv --ucl 52.0 --lcl 48.0",
    "",
    "Exit code 0; output contains '52.0' or '52.00' or 'user' or 'specified'.")

add_tc_block(doc,
    "TC-SPC-XBR-006  —  PNG written to ~/Downloads/",
    "UR-SPC-011",
    "jrc_spc_xbar_r oq/data/xbar_r_stable.csv",
    "",
    "File matching *_jrc_spc_xbar_r.png exists in ~/Downloads/ with recent mtime.")

add_tc_block(doc,
    "TC-SPC-XBR-007  —  No arguments: non-zero exit, usage message",
    "UR-SPC-026",
    "jrc_spc_xbar_r",
    "",
    "Exit code \u2260 0; 'Usage' or 'usage' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-008  —  File not found: non-zero exit",
    "UR-SPC-027",
    "jrc_spc_xbar_r /tmp/no_such_file_xyz.csv",
    "",
    "Exit code \u2260 0; 'not found' or filename present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-009  —  Missing column: non-zero exit, column named",
    "UR-SPC-007",
    "jrc_spc_xbar_r oq/data/xbar_r_missing_col.csv",
    "",
    "Exit code \u2260 0; 'value' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-010  —  Unbalanced subgroups: non-zero exit",
    "UR-SPC-012",
    "jrc_spc_xbar_r oq/data/xbar_r_unbalanced.csv",
    "xbar_r_unbalanced.csv has subgroups of different sizes.",
    "Exit code \u2260 0; one of 'unbalanced', 'equal', or 'same size' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-011  —  Subgroup size n > 10: non-zero exit",
    "UR-SPC-012",
    "jrc_spc_xbar_r oq/data/xbar_r_n_too_large.csv",
    "xbar_r_n_too_large.csv has n = 11; tabulated constants only cover n = 2\u201310.",
    "Exit code \u2260 0; one of '10', 'xbar_s', or 'large' present in output.")

add_tc_block(doc,
    "TC-SPC-XBR-012  —  Bypass protection: direct Rscript call fails",
    "UR-SPC-025",
    "Rscript repos/spc/R/jrc_spc_xbar_r.R oq/data/xbar_r_stable.csv  (RENV_PATHS_ROOT unset)",
    "",
    "Exit code \u2260 0; 'RENV_PATHS_ROOT' present in output.")

# ── 6.3  X-bar/S ────────────────────────────────────────────────────────────
add_heading(doc, "6.3  jrc_spc_xbar_s", level=2)

add_tc_block(doc,
    "TC-SPC-XBS-001  —  Stable dataset: exit 0, key sections present",
    "UR-SPC-013, UR-SPC-016",
    "jrc_spc_xbar_s oq/data/xbar_s_stable.csv",
    "20 subgroups \u00d7 n = 8; no WE rule should fire.",
    "Exit code 0; output contains X-bar section, S-bar/Std Dev section, and verdict.")

add_tc_block(doc,
    "TC-SPC-XBS-002  —  Stable dataset: IN CONTROL verdict",
    "UR-SPC-014, UR-SPC-016",
    "jrc_spc_xbar_s oq/data/xbar_s_stable.csv",
    "",
    "Exit code 0; 'IN CONTROL' or 'STABLE' present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-003  —  OOC dataset: exit 0, OUT OF CONTROL in output",
    "UR-SPC-014, UR-SPC-016",
    "jrc_spc_xbar_s oq/data/xbar_s_ooc.csv",
    "xbar_s_ooc.csv has subgroup sg09 with mean \u2248 102, far above UCL \u2248 100.55.",
    "Exit code 0; 'OUT OF CONTROL' or 'SIGNALS' present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-004  —  OOC dataset: subgroup sg09 flagged",
    "UR-SPC-014",
    "jrc_spc_xbar_s oq/data/xbar_s_ooc.csv",
    "",
    "Output contains 'sg09'.")

add_tc_block(doc,
    "TC-SPC-XBS-005  —  User-specified limits accepted",
    "UR-SPC-015",
    "jrc_spc_xbar_s oq/data/xbar_s_stable.csv --ucl 102.0 --lcl 98.0",
    "",
    "Exit code 0; output contains '102.0' or '102.00' or 'user' or 'specified'.")

add_tc_block(doc,
    "TC-SPC-XBS-006  —  PNG written to ~/Downloads/",
    "UR-SPC-017",
    "jrc_spc_xbar_s oq/data/xbar_s_stable.csv",
    "",
    "File matching *_jrc_spc_xbar_s.png exists in ~/Downloads/ with recent mtime.")

add_tc_block(doc,
    "TC-SPC-XBS-007  —  No arguments: non-zero exit, usage message",
    "UR-SPC-026",
    "jrc_spc_xbar_s",
    "",
    "Exit code \u2260 0; 'Usage' or 'usage' present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-008  —  File not found: non-zero exit",
    "UR-SPC-027",
    "jrc_spc_xbar_s /tmp/no_such_file_xyz.csv",
    "",
    "Exit code \u2260 0; 'not found' or filename present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-009  —  Missing column: non-zero exit, column named",
    "UR-SPC-013",
    "jrc_spc_xbar_s oq/data/xbar_s_missing_col.csv",
    "",
    "Exit code \u2260 0; 'value' present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-010  —  Unbalanced subgroups: non-zero exit",
    "UR-SPC-018",
    "jrc_spc_xbar_s oq/data/xbar_r_unbalanced.csv",
    "Reuses the xbar_r unbalanced dataset (same column structure).",
    "Exit code \u2260 0; one of 'unbalanced', 'equal', or 'same size' present in output.")

add_tc_block(doc,
    "TC-SPC-XBS-011  —  Bypass protection: direct Rscript call fails",
    "UR-SPC-025",
    "Rscript repos/spc/R/jrc_spc_xbar_s.R oq/data/xbar_s_stable.csv  (RENV_PATHS_ROOT unset)",
    "",
    "Exit code \u2260 0; 'RENV_PATHS_ROOT' present in output.")

# ── 6.4  P-chart ────────────────────────────────────────────────────────────
add_heading(doc, "6.4  jrc_spc_p", level=2)

add_tc_block(doc,
    "TC-SPC-P-001  —  Stable dataset: exit 0, key sections present",
    "UR-SPC-019, UR-SPC-020",
    "jrc_spc_p oq/data/p_stable.csv",
    "25 subgroups, n = 100 each, defectives vary 1\u20138; no WE rule should fire.",
    "Exit code 0; output contains p-bar section, UCL section, and verdict.")

add_tc_block(doc,
    "TC-SPC-P-002  —  Stable dataset: IN CONTROL verdict",
    "UR-SPC-019, UR-SPC-020",
    "jrc_spc_p oq/data/p_stable.csv",
    "",
    "Exit code 0; 'IN CONTROL' or 'STABLE' present in output.")

add_tc_block(doc,
    "TC-SPC-P-003  —  Known p-bar computed correctly",
    "UR-SPC-019",
    "jrc_spc_p oq/data/p_stable.csv",
    "With 104 total defectives across 2500 inspected, p_bar = 0.0416.",
    "Output matches regex 0.04[01234] or 4.[01234]%.")

add_tc_block(doc,
    "TC-SPC-P-004  —  OOC dataset: exit 0, OUT OF CONTROL in output",
    "UR-SPC-019, UR-SPC-020",
    "jrc_spc_p oq/data/p_ooc.csv",
    "p_ooc.csv has subgroup 13 = 15 defectives (p = 0.15 > UCL \u2248 0.10).",
    "Exit code 0; 'OUT OF CONTROL' or 'SIGNALS' present in output.")

add_tc_block(doc,
    "TC-SPC-P-005  —  OOC dataset: subgroup 13 flagged",
    "UR-SPC-019",
    "jrc_spc_p oq/data/p_ooc.csv",
    "",
    "Output contains '13'.")

add_tc_block(doc,
    "TC-SPC-P-006  —  PNG written to ~/Downloads/",
    "UR-SPC-022",
    "jrc_spc_p oq/data/p_stable.csv",
    "",
    "File matching *_jrc_spc_p.png exists in ~/Downloads/ with recent mtime.")

add_tc_block(doc,
    "TC-SPC-P-007  —  No arguments: non-zero exit, usage message",
    "UR-SPC-026",
    "jrc_spc_p",
    "",
    "Exit code \u2260 0; 'Usage' or 'usage' present in output.")

add_tc_block(doc,
    "TC-SPC-P-008  —  File not found: non-zero exit",
    "UR-SPC-027",
    "jrc_spc_p /tmp/no_such_file_xyz.csv",
    "",
    "Exit code \u2260 0; 'not found' or filename present in output.")

add_tc_block(doc,
    "TC-SPC-P-009  —  Missing column: non-zero exit, column named",
    "UR-SPC-019",
    "jrc_spc_p oq/data/p_missing_col.csv",
    "p_missing_col.csv lacks the 'defectives' column.",
    "Exit code \u2260 0; 'defective' present in output.")

add_tc_block(doc,
    "TC-SPC-P-010  —  Defectives exceed n: non-zero exit",
    "UR-SPC-021",
    "jrc_spc_p oq/data/p_invalid.csv",
    "p_invalid.csv contains a row where defectives = 101 > n = 100.",
    "Exit code \u2260 0; one of 'defective', 'exceed', 'greater', 'cannot', or 'invalid' "
    "present in output.")

add_tc_block(doc,
    "TC-SPC-P-011  —  Bypass protection: direct Rscript call fails",
    "UR-SPC-025",
    "Rscript repos/spc/R/jrc_spc_p.R oq/data/p_stable.csv  (RENV_PATHS_ROOT unset)",
    "",
    "Exit code \u2260 0; 'RENV_PATHS_ROOT' present in output.")

# ── 6.5  C-chart ────────────────────────────────────────────────────────────
add_heading(doc, "6.5  jrc_spc_c", level=2)

add_tc_block(doc,
    "TC-SPC-C-001  —  Stable dataset: exit 0, key sections present",
    "UR-SPC-023, UR-SPC-024",
    "jrc_spc_c oq/data/c_stable.csv",
    "25 subgroups; defect counts vary 1\u20139; no WE rule should fire.",
    "Exit code 0; output contains c-bar section, UCL section, and verdict.")

add_tc_block(doc,
    "TC-SPC-C-002  —  Stable dataset: IN CONTROL verdict",
    "UR-SPC-023, UR-SPC-024",
    "jrc_spc_c oq/data/c_stable.csv",
    "",
    "Exit code 0; 'IN CONTROL' or 'STABLE' present in output.")

add_tc_block(doc,
    "TC-SPC-C-003  —  Known c-bar computed correctly",
    "UR-SPC-023",
    "jrc_spc_c oq/data/c_stable.csv",
    "With 124 total defects across 25 subgroups, c_bar = 4.96.",
    "Output matches regex 4.8[0-9] or 4.9[0-9].")

add_tc_block(doc,
    "TC-SPC-C-004  —  OOC dataset: exit 0, OUT OF CONTROL in output",
    "UR-SPC-023, UR-SPC-024",
    "jrc_spc_c oq/data/c_ooc.csv",
    "c_ooc.csv has subgroup 12 = 15 defects, above UCL \u2248 12.2.",
    "Exit code 0; 'OUT OF CONTROL' or 'SIGNALS' present in output.")

add_tc_block(doc,
    "TC-SPC-C-005  —  OOC dataset: subgroup 12 flagged",
    "UR-SPC-023",
    "jrc_spc_c oq/data/c_ooc.csv",
    "",
    "Output contains '12'.")

add_tc_block(doc,
    "TC-SPC-C-006  —  PNG written to ~/Downloads/",
    "UR-SPC-024",
    "jrc_spc_c oq/data/c_stable.csv",
    "",
    "File matching *_jrc_spc_c.png exists in ~/Downloads/ with recent mtime.")

add_tc_block(doc,
    "TC-SPC-C-007  —  No arguments: non-zero exit, usage message",
    "UR-SPC-026",
    "jrc_spc_c",
    "",
    "Exit code \u2260 0; 'Usage' or 'usage' present in output.")

add_tc_block(doc,
    "TC-SPC-C-008  —  File not found: non-zero exit",
    "UR-SPC-027",
    "jrc_spc_c /tmp/no_such_file_xyz.csv",
    "",
    "Exit code \u2260 0; 'not found' or filename present in output.")

add_tc_block(doc,
    "TC-SPC-C-009  —  Missing column: non-zero exit, column named",
    "UR-SPC-023",
    "jrc_spc_c oq/data/c_missing_col.csv",
    "c_missing_col.csv lacks the 'defects' column.",
    "Exit code \u2260 0; 'defect' present in output.")

add_tc_block(doc,
    "TC-SPC-C-010  —  Bypass protection: direct Rscript call fails",
    "UR-SPC-025",
    "Rscript repos/spc/R/jrc_spc_c.R oq/data/c_stable.csv  (RENV_PATHS_ROOT unset)",
    "",
    "Exit code \u2260 0; 'RENV_PATHS_ROOT' present in output.")

doc.add_page_break()

# ---------------------------------------------------------------------------
# 7. ACCEPTANCE CRITERIA
# ---------------------------------------------------------------------------

add_heading(doc, "7.  Acceptance Criteria", level=1)
add_para(doc,
    "The OQ is considered passed when all of the following conditions are met:")
criteria = [
    "All 55 test cases defined in Section 6 pass with exit code as specified.",
    "No test case fails (i.e. pytest reports 0 failures and 0 errors).",
    "The OQ evidence file (timestamped pytest output) is written to "
    "~/.jrscript/<PROJECT_ID>/validation/ by admin_spc_oq.",
    "All 27 user requirements (UR-SPC-001 through UR-SPC-027) are covered by at "
    "least one passing test case.",
    "Any deviation from this plan is documented and assessed for impact before "
    "the OQ is declared passed.",
]
for item in criteria:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 8. TRACEABILITY MATRIX
# ---------------------------------------------------------------------------

add_heading(doc, "8.  Requirements Traceability Matrix", level=1)
add_para(doc,
    "The table below maps each user requirement to the test cases that verify it.")
doc.add_paragraph()

rtm_rows = [
    ("UR-SPC-001", "jrc_spc_imr — X-bar, MR-bar, control limits",
     "TC-SPC-IMR-001, 002"),
    ("UR-SPC-002", "jrc_spc_imr — WE rules, violation list",
     "TC-SPC-IMR-002, 003, 004"),
    ("UR-SPC-003", "jrc_spc_imr — --ucl / --lcl",
     "TC-SPC-IMR-005"),
    ("UR-SPC-004", "jrc_spc_imr — IN CONTROL / OUT OF CONTROL verdict",
     "TC-SPC-IMR-001, 002, 003"),
    ("UR-SPC-005", "jrc_spc_imr — PNG to ~/Downloads/",
     "TC-SPC-IMR-006"),
    ("UR-SPC-006", "jrc_spc_imr — reject invalid inputs",
     "TC-SPC-IMR-007, 008, 009, 010"),
    ("UR-SPC-007", "jrc_spc_xbar_r — grand mean, R-bar, limits",
     "TC-SPC-XBR-001, 002"),
    ("UR-SPC-008", "jrc_spc_xbar_r — WE rules, violation list",
     "TC-SPC-XBR-002, 003, 004"),
    ("UR-SPC-009", "jrc_spc_xbar_r — --ucl / --lcl",
     "TC-SPC-XBR-005"),
    ("UR-SPC-010", "jrc_spc_xbar_r — verdict",
     "TC-SPC-XBR-001, 002, 003"),
    ("UR-SPC-011", "jrc_spc_xbar_r — PNG to ~/Downloads/",
     "TC-SPC-XBR-006"),
    ("UR-SPC-012", "jrc_spc_xbar_r — reject unbalanced / n > 10",
     "TC-SPC-XBR-010, 011"),
    ("UR-SPC-013", "jrc_spc_xbar_s — grand mean, S-bar, limits via c4",
     "TC-SPC-XBS-001, 002"),
    ("UR-SPC-014", "jrc_spc_xbar_s — WE rules, violation list",
     "TC-SPC-XBS-002, 003, 004"),
    ("UR-SPC-015", "jrc_spc_xbar_s — --ucl / --lcl",
     "TC-SPC-XBS-005"),
    ("UR-SPC-016", "jrc_spc_xbar_s — verdict",
     "TC-SPC-XBS-001, 002, 003"),
    ("UR-SPC-017", "jrc_spc_xbar_s — PNG to ~/Downloads/",
     "TC-SPC-XBS-006"),
    ("UR-SPC-018", "jrc_spc_xbar_s — reject unbalanced subgroups",
     "TC-SPC-XBS-010"),
    ("UR-SPC-019", "jrc_spc_p — p-bar, variable limits, WE rules",
     "TC-SPC-P-001, 002, 003, 004, 005"),
    ("UR-SPC-020", "jrc_spc_p — verdict",
     "TC-SPC-P-001, 002, 004"),
    ("UR-SPC-021", "jrc_spc_p — reject defectives > n",
     "TC-SPC-P-010"),
    ("UR-SPC-022", "jrc_spc_p — PNG to ~/Downloads/",
     "TC-SPC-P-006"),
    ("UR-SPC-023", "jrc_spc_c — c-bar, sigma, limits, WE rules",
     "TC-SPC-C-001, 002, 003, 004, 005"),
    ("UR-SPC-024", "jrc_spc_c — verdict; PNG to ~/Downloads/",
     "TC-SPC-C-001, 002, 004, 006"),
    ("UR-SPC-025", "All scripts — RENV_PATHS_ROOT bypass protection",
     "TC-SPC-IMR-011, XBR-012, XBS-011, P-011, C-010"),
    ("UR-SPC-026", "All scripts — no-arguments usage message",
     "TC-SPC-IMR-007, XBR-007, XBS-007, P-007, C-007"),
    ("UR-SPC-027", "All scripts — file not found error",
     "TC-SPC-IMR-008, XBR-008, XBS-008, P-008, C-008"),
]

add_table(doc,
    ["UR", "Requirement (summary)", "Test Cases"],
    rtm_rows,
    col_widths=[2.5, 7, 7.5]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 9. ROLES AND RESPONSIBILITIES
# ---------------------------------------------------------------------------

add_heading(doc, "9.  Roles and Responsibilities", level=1)
doc.add_paragraph()
add_table(doc,
    ["Role", "Responsibility"],
    [
        ("Author",
         "Prepares this validation plan; executes the OQ test suite; "
         "documents results in the validation report."),
        ("Reviewer",
         "Reviews this plan and the corresponding validation report for "
         "technical accuracy and completeness."),
        ("Approver (QA)",
         "Approves this plan prior to OQ execution; approves the validation "
         "report upon successful completion of OQ."),
    ],
    col_widths=[4, 13]
)

# ---------------------------------------------------------------------------
# 10. DEVIATIONS
# ---------------------------------------------------------------------------

add_heading(doc, "10.  Deviation Handling", level=1)
add_para(doc,
    "Any deviation from this plan (e.g. a test case that cannot be executed as "
    "written, or a test that fails due to an environmental issue rather than a "
    "software defect) must be documented in the validation report with:")
for item in [
    "A unique deviation ID",
    "Description of the deviation",
    "Root cause assessment",
    "Impact on the OQ conclusion",
    "Disposition (accepted with justification, or requires remediation)",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

add_para(doc,
    "The OQ may not be declared passed if any test case fails due to a software "
    "defect, unless a formal deviation with justification is approved by the QA "
    "approver.")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT)
print(f"Saved: {OUT}")
