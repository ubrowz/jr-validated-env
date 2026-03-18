"""
generate_spc_user_manual.py

Generates repos/spc/docs/spc_user_manual.docx — the engineer-facing user guide
for the JR Validated Environment SPC Module.

Run with:
    python3 repos/spc/docs/ignore/generate_spc_user_manual.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "spc_user_manual.docx")

NAVY  = "1A1A2E"
GRAY  = "F2F2F2"
WHITE = "FFFFFF"
INFO  = "E8F0F7"
WARN  = "FFF3CD"
GREEN = "E8F5E9"

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def header_row(table, cols, widths=None):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        set_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def data_row(table, row_idx, values, shade=True, bold_first=False):
    row = table.rows[row_idx]
    fill = GRAY if (shade and row_idx % 2 == 1) else WHITE
    for i, text in enumerate(values):
        cell = row.cells[i]
        set_shading(cell, fill)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if bold_first and i == 0:
            run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def para(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = "Courier New"


def box(doc, fill, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    p.clear()
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.name = "Calibri"
    doc.add_paragraph()


def info(doc, label, text):  box(doc, INFO, label, text)
def warn(doc, label, text):  box(doc, WARN, label, text)
def tip(doc, label, text):   box(doc, GREEN, label, text)


def script_section(doc, number, script_name, when, args_table, example_cmds, output_note):
    """Standard subsection layout for each script."""
    h2(doc, f"{number}  {script_name}")
    h3(doc, "When do I use this?")
    para(doc, when)
    h3(doc, "What do I need?")
    tbl = doc.add_table(rows=len(args_table) + 1, cols=2)
    tbl.style = "Table Grid"
    header_row(tbl, ["Argument", "What to provide"], [1.8, 5.1])
    for i, (arg, desc) in enumerate(args_table, start=1):
        data_row(tbl, i, [arg, desc], bold_first=True)
    doc.add_paragraph()
    h3(doc, "Example")
    for cmd in example_cmds:
        code(doc, cmd)
    h3(doc, "What to look for in the output")
    para(doc, output_note)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(0.984)
section.top_margin    = Inches(0.984)
section.bottom_margin = Inches(0.984)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
cover.paragraph_format.space_after  = Pt(6)
r = cover.add_run("JR Validated Environment \u2014 SPC Module")
r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("User Guide for Statistical Process Control")
r.bold = True; r.font.size = Pt(13); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  |  Date: 2026-03-18  |  Audience: Design and Manufacturing Engineers")
r.font.size = Pt(10); r.font.name = "Calibri"
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 1 — How to Use This Guide
# ---------------------------------------------------------------------------
h1(doc, "1. How to Use This Guide")
para(doc,
    "This guide describes the five SPC control chart scripts available in the JR Validated "
    "Environment. Each script produces a different chart type suited to a specific data structure "
    "and monitoring goal. Choose the chart that matches your process data:")

bullet(doc, "One measurement per time point (individual units, batches, shifts): I-MR chart")
bullet(doc, "Small subgroups of measurements per time point, subgroup size 2\u201310: X-bar/R chart")
bullet(doc, "Subgroups of any size, or subgroup size > 10: X-bar/S chart")
bullet(doc, "Count of nonconforming units (pass/fail) per inspection lot: P-chart")
bullet(doc, "Count of defects per unit, constant inspection unit size: C-chart")

para(doc,
    "All scripts apply the 8 Western Electric (WE) rules and output an IN CONTROL / "
    "OUT OF CONTROL verdict. The continuous measurement scripts (I-MR, X-bar/R, X-bar/S) "
    "support optional user-specified historical control limits. See Section 3 for "
    "interpretation guidance and Section 6 for the full WE rule definitions.")

# ---------------------------------------------------------------------------
# Section 2 — Quick Reference Table
# ---------------------------------------------------------------------------
h1(doc, "2. Quick Reference Table")
para(doc, "Find the row that matches your data, then go to the section listed.")
doc.add_paragraph()

qr_rows = [
    ("I collect one measurement per time point (one batch, one shift, one unit per day)",
     "jrc_spc_imr", "4.1"),
    ("I collect small groups of units at each time point, subgroup size 2\u201310",
     "jrc_spc_xbar_r", "4.2"),
    ("I collect subgroups at each time point, subgroup size > 10 or any size",
     "jrc_spc_xbar_s", "4.3"),
    ("I record how many units fail inspection in each batch or lot (pass/fail count)",
     "jrc_spc_p", "4.4"),
    ("I count individual defects per unit (one unit can have multiple defects)",
     "jrc_spc_c", "4.5"),
]

tbl = doc.add_table(rows=len(qr_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["My data looks like\u2026", "Script", "Section"], [4.0, 2.0, 0.9])
for i, row_data in enumerate(qr_rows, start=1):
    data_row(tbl, i, list(row_data))
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 3 — Understanding the Results
# ---------------------------------------------------------------------------
h1(doc, "3. Understanding the Results")

h2(doc, "3.1  Control Limits")
para(doc,
    "Control limits (UCL and LCL) are calculated from the data itself \u2014 they are not "
    "specification limits. They represent the expected range of variation for a stable process "
    "(approximately \u00b13\u03c3 from the centreline). A point outside the control limits, or a "
    "non-random pattern within them, is a signal that something has changed in the process.")

info(doc, "INFO:",
    "IN CONTROL does not mean the process meets specification. A process can be perfectly "
    "stable and predictable while still producing out-of-spec results. Whether values meet "
    "specifications is answered by capability analysis (jrc_capability), not by a control chart.")

h2(doc, "3.2  IN CONTROL / OUT OF CONTROL Verdict")
para(doc,
    "The terminal output ends with one of two verdicts:")
bullet(doc, "IN CONTROL \u2014 no WE rule violations detected. The process is stable and predictable.")
bullet(doc,
    "OUT OF CONTROL \u2014 one or more WE rule violations detected. The script lists the rule "
    "numbers and the affected point IDs. Investigate for assignable (special) causes before "
    "using the process outputs for product release.")

h2(doc, "3.3  Western Electric Rules")
para(doc,
    "All eight WE rules are applied to the primary chart (I chart, X-bar chart, P chart, C chart). "
    "The secondary chart (MR chart, Range chart, S chart) has only Rule 1 applied. "
    "Section 6 defines all eight rules. Key guidance:")
bullet(doc,
    "Rule 1 (point beyond 3\u03c3): highest priority signal \u2014 investigate immediately.")
bullet(doc,
    "Rules 2, 3, 5 (runs and trends): indicate a sustained process shift or drift.")
bullet(doc,
    "Rules 4, 6, 7, 8 (patterns): indicate cycles, over-control, or stratification.")

h2(doc, "3.4  Historical Limits")
para(doc,
    "The I-MR, X-bar/R, and X-bar/S scripts accept optional --ucl and --lcl arguments. "
    "Use historical limits when you want to compare current production against a previously "
    "established baseline rather than computing limits from the current data. Historical limits "
    "are displayed on the chart alongside the data but the verdict and WE rules are still "
    "evaluated against these user-supplied values.")

h2(doc, "3.5  Variable Control Limits (P-chart)")
para(doc,
    "The P-chart computes per-subgroup control limits when subgroup sizes (n) vary between "
    "lots. Larger subgroups produce narrower limits; smaller subgroups produce wider limits. "
    "The chart displays limit lines that step with each subgroup. WE pattern rules are applied "
    "using standardised z-values to avoid false signals from the changing limit width.")

# ---------------------------------------------------------------------------
# Section 4 — Per-Script Reference
# ---------------------------------------------------------------------------
h1(doc, "4. Per-Script Reference")

# 4.1 jrc_spc_imr
script_section(doc,
    "4.1", "jrc_spc_imr \u2014 Individuals and Moving Range (I-MR)",
    "Use when you have one measurement per time point \u2014 one reading per shift, one batch "
    "result per day, one patient measurement per visit. The I chart monitors the process level "
    "and the MR chart monitors process variation (as the absolute difference between consecutive "
    "observations). This is the right chart when subgrouping is not practical.",
    [
        ("data.csv",
         "CSV with columns: id, value. The \u2018id\u2019 column labels each observation (shift number, "
         "date, batch ID). The \u2018value\u2019 column is the numeric measurement. "
         "Minimum: 3 observations (to compute at least 2 moving ranges)."),
        ("--ucl <value>",
         "Optional. User-specified Upper Control Limit for the I chart. Use when applying "
         "historical control limits established from a prior stable period."),
        ("--lcl <value>",
         "Optional. User-specified Lower Control Limit for the I chart."),
    ],
    [
        "jrc_spc_imr batch_results.csv",
        "jrc_spc_imr daily_weight.csv --ucl 10.75 --lcl 9.25",
    ],
    "Check the terminal verdict. If OUT OF CONTROL, the listed rule numbers and point IDs "
    "tell you where to investigate. The two-panel PNG shows the I chart on top (with 1\u03c3 and "
    "2\u03c3 zone lines for WE rules) and the MR chart below. Out-of-control points are shown in "
    "red. A signal on the MR chart indicates a sudden change in variability \u2014 investigate "
    "before interpreting the I chart."
)

# 4.2 jrc_spc_xbar_r
script_section(doc,
    "4.2", "jrc_spc_xbar_r \u2014 X-bar and Range",
    "Use when you measure a small group (subgroup) of units at regular intervals \u2014 for "
    "example, five parts pulled from a line every hour, or three readings at the start of each "
    "shift. The X-bar chart tracks the subgroup mean; the Range chart tracks within-subgroup "
    "spread. Best suited for subgroup sizes of 2\u201310. For larger subgroups, use jrc_spc_xbar_s.",
    [
        ("data.csv",
         "CSV with columns: subgroup, value (long format \u2014 one row per observation). "
         "The \u2018subgroup\u2019 label groups observations collected at the same time point. "
         "All subgroups must have the same number of observations (balanced design). "
         "Minimum: 2 subgroups, subgroup size n \u2265 2."),
        ("--ucl <value>",
         "Optional. User-specified UCL for the X-bar chart."),
        ("--lcl <value>",
         "Optional. User-specified LCL for the X-bar chart."),
    ],
    [
        "jrc_spc_xbar_r hourly_samples.csv",
        "jrc_spc_xbar_r fill_weights.csv --ucl 105.2 --lcl 94.8",
    ],
    "The terminal output shows chart constants (d2, D3, D4, A2) for the detected subgroup "
    "size, grand mean, R-bar, control limits, and any WE rule violations. The two-panel PNG "
    "shows the X-bar chart on top and the Range chart below. Investigate any Range chart signals "
    "first \u2014 unstable variation makes the X-bar limits unreliable."
)

# 4.3 jrc_spc_xbar_s
script_section(doc,
    "4.3", "jrc_spc_xbar_s \u2014 X-bar and Standard Deviation",
    "Use for subgrouped data when subgroup size is larger than 10, or when you want the "
    "statistically more efficient standard deviation in place of the range. The X-bar/S chart "
    "is valid for any subgroup size \u2265 2. Chart constants (c4, A3, B3, B4) are computed "
    "analytically from the gamma function, so no lookup table is required.",
    [
        ("data.csv",
         "CSV with columns: subgroup, value (long format \u2014 one row per observation). "
         "All subgroups must have the same number of observations (balanced design). "
         "Minimum: 2 subgroups, subgroup size n \u2265 2."),
        ("--ucl <value>",
         "Optional. User-specified UCL for the X-bar chart."),
        ("--lcl <value>",
         "Optional. User-specified LCL for the X-bar chart."),
    ],
    [
        "jrc_spc_xbar_s subgroup_data.csv",
        "jrc_spc_xbar_s tensile_strength.csv --ucl 52.0 --lcl 44.0",
    ],
    "The terminal output shows chart constants (c4, A3, B3, B4), grand mean, S-bar, control "
    "limits, and any WE rule violations. The two-panel PNG shows the X-bar chart on top and "
    "the S chart below. Investigate S chart signals before interpreting the X-bar chart."
)

tip(doc, "TIP:",
    "When subgroup size is between 2 and 10, jrc_spc_xbar_r and jrc_spc_xbar_s will give "
    "virtually identical results in practice. Choose R for simplicity; choose S if your "
    "organisation\u2019s SOP specifies the standard deviation chart or if subgroup size exceeds 10.")

# 4.4 jrc_spc_p
script_section(doc,
    "4.4", "jrc_spc_p \u2014 P-chart (Proportion Nonconforming)",
    "Use when each data point represents an inspection lot where you record the number of "
    "units inspected and the number found nonconforming (pass/fail). Each unit is classified "
    "as either conforming or nonconforming. Supports variable lot sizes \u2014 control limits "
    "automatically adjust for each subgroup\u2019s sample size.",
    [
        ("data.csv",
         "CSV with columns: subgroup, n, defectives. "
         "\u2018subgroup\u2019 is a label for each lot or time period. "
         "\u2018n\u2019 is the number of units inspected. "
         "\u2018defectives\u2019 is the count of nonconforming units (must be \u2264 n). "
         "Minimum: 2 subgroups. All n \u2265 1. All defectives \u2265 0."),
    ],
    [
        "jrc_spc_p weekly_inspection.csv",
        "jrc_spc_p final_test_results.csv",
    ],
    "The terminal shows p-bar (weighted overall proportion defective), per-subgroup limits "
    "if n varies, and any WE rule violations. The PNG shows each subgroup\u2019s fraction defective "
    "with stepped per-point limit lines. OUT OF CONTROL signals indicate the defect rate has "
    "shifted. IN CONTROL at a high defect rate still requires process improvement \u2014 "
    "the chart confirms stability, not acceptability."
)

# 4.5 jrc_spc_c
script_section(doc,
    "4.5", "jrc_spc_c \u2014 C-chart (Count of Defects per Unit)",
    "Use when you count the number of individual defects on each inspection unit \u2014 for "
    "example, weld spatter points per panel, cosmetic blemishes per device, solder defects per "
    "circuit board, or documentation errors per batch record. A single unit can have multiple "
    "defects. The inspection unit must be of constant size or opportunity for defects. "
    "The statistical basis is the Poisson distribution.",
    [
        ("data.csv",
         "CSV with columns: (subgroup or id), defects. The first column may be named "
         "\u2018subgroup\u2019 or \u2018id\u2019 \u2014 both are accepted. The \u2018defects\u2019 column must contain "
         "non-negative integers. Minimum: 2 subgroups."),
    ],
    [
        "jrc_spc_c cosmetic_inspection.csv",
        "jrc_spc_c weld_defects.csv",
    ],
    "The terminal shows c-bar (mean defects per unit), sigma (square root of c-bar), UCL, "
    "LCL, and any WE rule violations. The PNG shows defect counts per unit with UCL/LCL and "
    "1\u03c3/2\u03c3 zone lines. A Rule 1 signal (sudden high count) requires immediate investigation. "
    "Pattern rules (runs, trends) suggest a gradual drift in the defect-generating process."
)

warn(doc, "NOTE:",
    "The C-chart requires constant inspection opportunity. If your inspection unit varies in "
    "size (different lengths of weld, different numbers of solder joints per board), the "
    "control limits will be incorrect. In that case, calculate defects per unit of opportunity "
    "and use an appropriate U-chart \u2014 not currently available as a JR script.")

# ---------------------------------------------------------------------------
# Section 5 — Preparing Your Data
# ---------------------------------------------------------------------------
h1(doc, "5. Preparing Your Data")
para(doc,
    "Each script expects a CSV file with specific column names. The table below summarises "
    "the required columns for each script.")
doc.add_paragraph()

csv_rows = [
    ("jrc_spc_imr",     "id, value",                          "none"),
    ("jrc_spc_xbar_r",  "subgroup, value (long format)",      "none"),
    ("jrc_spc_xbar_s",  "subgroup, value (long format)",      "none"),
    ("jrc_spc_p",       "subgroup, n, defectives",            "none"),
    ("jrc_spc_c",       "subgroup (or id), defects",          "none"),
]

tbl = doc.add_table(rows=len(csv_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["Script", "Required columns", "Optional columns"], [2.0, 3.7, 1.2])
for i, row_data in enumerate(csv_rows, start=1):
    data_row(tbl, i, list(row_data), bold_first=True)
doc.add_paragraph()

para(doc,
    "Column names are case-insensitive. Extra columns are ignored. The \u2018id\u2019 and \u2018subgroup\u2019 "
    "columns may contain text labels or numbers. Numeric value columns must be numeric. "
    "Minimum row counts are enforced; the script will exit with a clear error if the data "
    "does not meet the minimum requirements.")

info(doc, "INFO:",
    "The X-bar/R and X-bar/S scripts use long format: one row per observation, not one row "
    "per subgroup. If your data is in wide format (one row per subgroup with multiple value "
    "columns), you must reshape it before running the script. Use jrc_convert_csv or standard "
    "spreadsheet operations to pivot the data.")

# ---------------------------------------------------------------------------
# Section 6 — Western Electric Rules Reference
# ---------------------------------------------------------------------------
h1(doc, "6. Western Electric Rules Reference")
para(doc,
    "All scripts apply all 8 Western Electric rules to the primary chart. The secondary "
    "variability chart (MR, R, or S) has only Rule 1 applied. The P-chart applies rules "
    "using standardised z-values to account for variable control limit widths.")
doc.add_paragraph()

we_rows = [
    ("1", "Any single point beyond 3\u03c3 (outside control limits)",
          "Immediate special cause \u2014 highest priority signal"),
    ("2", "9 consecutive points on the same side of the centreline",
          "Sustained process shift \u2014 investigate for a step change"),
    ("3", "6 consecutive points steadily increasing or decreasing",
          "Process drift or trend \u2014 investigate for gradual change (tool wear, temperature)"),
    ("4", "14 consecutive points alternating up and down",
          "Over-control or systematic oscillation \u2014 investigate for two alternating streams"),
    ("5", "2 of 3 consecutive points beyond 2\u03c3 on the same side",
          "Incipient shift \u2014 early warning of a process change"),
    ("6", "4 of 5 consecutive points beyond 1\u03c3 on the same side",
          "Sustained off-centre \u2014 process mean has shifted slightly"),
    ("7", "15 consecutive points within 1\u03c3 of the centreline (stratification)",
          "Two or more streams mixed in the data, or over-tight sampling"),
    ("8", "8 consecutive points beyond 1\u03c3 on either side (mixture)",
          "Two distinct process streams contributing to the data"),
]

tbl = doc.add_table(rows=len(we_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["Rule", "Signal definition", "Likely cause"], [0.5, 3.2, 3.2])
for i, row_data in enumerate(we_rows, start=1):
    data_row(tbl, i, list(row_data), bold_first=True)
doc.add_paragraph()

para(doc,
    "When a rule fires, the terminal output lists the rule number and the IDs of the points "
    "involved. Use this information to identify the time window to investigate. Begin with "
    "Rule 1 signals; pattern rules (2\u20138) indicate sustained process changes that may not "
    "produce individual extreme values.")

warn(doc, "IMPORTANT:",
    "Out-of-control signals require investigation and documented root cause analysis before "
    "accepting the process output for product release in a regulated environment. "
    "Do not re-run the chart with different data to make signals disappear \u2014 "
    "address the assignable cause.")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
