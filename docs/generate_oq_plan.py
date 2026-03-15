"""
Generates oq_validation_plan.docx — FDA-acceptable OQ Validation Plan
for the JR Validated Environment community script suite v1.1.0.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "oq_validation_plan.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text="", bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1A1A2E")
        shading.set(qn("w:color"), "FFFFFF")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), bg)
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_tc_block(doc, tc_id, tests_ur, command, rationale, pass_criterion):
    """Add a formatted test case block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tc_id)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x21, 0x66, 0xAC)

    def labeled_line(label, value):
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Cm(0.8)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        r1 = lp.add_run(label + ": ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = lp.add_run(value)
        r2.font.size = Pt(9)
        if label == "Command":
            r2.font.name = "Courier New"

    labeled_line("Tests", tests_ur)
    labeled_line("Command", command)
    if rationale:
        labeled_line("Rationale", rationale)
    labeled_line("Pass criterion", pass_criterion)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ---------------------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------------------

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("VALIDATION PLAN\nOPERATIONAL QUALIFICATION\nCOMMUNITY SCRIPT SUITE")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("JR Validated Environment — Version 1.1.0")
r2.font.size = Pt(13)
r2.bold = True

doc.add_paragraph()

add_table(doc,
    ["Field", "Value"],
    [
        ("Document Number",  "JR-VP-002"),
        ("Title",            "Validation Plan — Operational Qualification, Community Script Suite"),
        ("System",           "JR Validated Environment"),
        ("System Version",   "1.1.0"),
        ("Document Version", "1.0"),
        ("Status",           "Draft"),
        ("Effective Date",   "2026-03-15"),
        ("Author",           "Joep Rous"),
        ("Reviewer",         "[Name]"),
        ("Approver",         "[Name]"),
    ],
    col_widths=[5, 11]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# VERSION HISTORY
# ---------------------------------------------------------------------------

add_heading(doc, "Version History", level=1)
add_table(doc,
    ["Version", "Date", "Author", "Description"],
    [("1.0", "2026-03-15", "Joep Rous",
      "Initial release. Covers all 22 R and 2 Python community scripts in v1.1.0.")],
    col_widths=[2, 3, 4, 8]
)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# APPROVAL SIGNATURES
# ---------------------------------------------------------------------------

add_heading(doc, "Approval Signatures", level=1)
add_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ("Author",       "", "", ""),
        ("Reviewer",     "", "", ""),
        ("Approver (QA)","", "", ""),
    ],
    col_widths=[4, 5, 5, 3]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. PURPOSE
# ---------------------------------------------------------------------------

add_heading(doc, "1.  Purpose", level=1)
add_para(doc,
    "This document defines the Validation Plan for the Operational Qualification (OQ) "
    "of the community script suite included in version 1.1.0 of the JR Validated "
    "Environment. It specifies the user requirements, test cases, acceptance criteria, "
    "and traceability between requirements and tests needed to demonstrate that each "
    "script performs correctly for its intended use.")
add_para(doc,
    "Execution of the tests defined in this plan, together with satisfactory results "
    "and documented evidence, constitutes the OQ for the community script suite.")

# ---------------------------------------------------------------------------
# 2. SCOPE
# ---------------------------------------------------------------------------

add_heading(doc, "2.  Scope", level=1)
add_heading(doc, "2.1  In Scope", level=2)

in_scope = [
    "All 22 R scripts in the R/ directory (excluding jrhello.R)",
    "All 2 Python scripts in the Python/ directory (excluding jrhello.py)",
    "Correct computation of outputs for valid inputs",
    "Correct rejection of invalid inputs with informative error messages",
    "Correct file I/O behaviour where applicable (CSV input, CSV/PNG output)",
    "Verification that all R scripts enforce execution within the jrrun validated "
    "environment by checking RENV_PATHS_ROOT at startup",
]
for item in in_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

add_heading(doc, "2.2  Out of Scope", level=2)
out_scope = [
    "Infrastructure scripts in bin/ and admin/ (covered by IQ; see JR-VP-001)",
    "Demo scripts jrhello.R and jrhello.py",
    "Performance qualification (PQ) — end-to-end verification studies using "
    "production data are conducted separately under the design verification protocol",
    "User acceptance testing (UAT)",
    "Validation of R and Python language interpreters (assumed qualified by their "
    "respective maintainers)",
    "Validation of third-party packages (tolerance, MASS, e1071, ggplot2, survival, "
    "nortest, outliers) — package integrity is confirmed at installation via SHA256 "
    "and is documented in the IQ evidence report (JR-IQ-001)",
]
for item in out_scope:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 3. REGULATORY BASIS
# ---------------------------------------------------------------------------

add_heading(doc, "3.  Regulatory Basis and References", level=1)
add_heading(doc, "3.1  Regulations", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("21 CFR Part 820, §820.70(i)", "Quality System Regulation — Automated Data Processing"),
        ("21 CFR Part 11",              "Electronic Records; Electronic Signatures"),
    ],
    col_widths=[5, 12]
)

doc.add_paragraph()
add_heading(doc, "3.2  FDA Guidance", level=2)
add_table(doc,
    ["Reference", "Title", "Date"],
    [
        ("FDA GPSV",
         "General Principles of Software Validation; Final Guidance for Industry and FDA Staff",
         "January 2002"),
        ("FDA CSA",
         "Computer Software Assurance for Production and Quality System Software; Draft Guidance",
         "September 2022"),
    ],
    col_widths=[3, 11, 3]
)

doc.add_paragraph()
add_heading(doc, "3.3  Standards", level=2)
add_table(doc,
    ["Reference", "Title"],
    [
        ("ASTM F3172-15",
         "Standard Guide for Design Verification Device Size and Sample Size Selection "
         "for Endovascular Devices"),
        ("ISO 13485:2016",
         "Medical devices — Quality management systems"),
    ],
    col_widths=[4, 13]
)

doc.add_paragraph()
add_heading(doc, "3.4  Internal Documents", level=2)
add_table(doc,
    ["Document Number", "Title"],
    [
        ("JR-VP-001", "Validation Plan — Installation Qualification, JR Validated Environment v1.0.0"),
        ("JR-IQ-001", "IQ Execution Evidence (file: docs/IQ_validation_20260311_205146.txt)"),
    ],
    col_widths=[4, 13]
)

# ---------------------------------------------------------------------------
# 4. DEFINITIONS
# ---------------------------------------------------------------------------

add_heading(doc, "4.  Definitions and Abbreviations", level=1)
add_table(doc,
    ["Term", "Definition"],
    [
        ("IQ",    "Installation Qualification — verifies that the system is installed correctly "
                  "and that the environment is complete and consistent"),
        ("OQ",    "Operational Qualification — verifies that the system operates correctly "
                  "within specified limits across the range of anticipated inputs"),
        ("PQ",    "Performance Qualification — verifies that the system performs consistently "
                  "under real conditions of use"),
        ("UR",    "User Requirement — a statement of what the system must do from the user's perspective"),
        ("TC",    "Test Case — a specific set of inputs, execution conditions, and expected results "
                  "used to verify a user requirement"),
        ("jrrun", "The validated zsh entry-point script that sets the controlled environment "
                  "before executing any community script"),
        ("renv",  "R package management system; used to isolate the validated R library"),
        ("venv",  "Python virtual environment; used to isolate the validated Python library"),
        ("CSV",   "Comma-separated values file; two-column format (id, value) as defined by the jrc suite"),
        ("TOST",  "Two One-Sided Tests — statistical method for equivalence testing"),
        ("MLE",   "Maximum Likelihood Estimation"),
        ("LoA",   "Limits of Agreement (Bland-Altman method)"),
        ("TI",    "Statistical Tolerance Interval"),
        ("LSL",   "Lower Specification Limit"),
        ("USL",   "Upper Specification Limit"),
    ],
    col_widths=[3, 14]
)

# ---------------------------------------------------------------------------
# 5. SYSTEM DESCRIPTION
# ---------------------------------------------------------------------------

add_heading(doc, "5.  System Description", level=1)
add_heading(doc, "5.1  Software Description", level=2)
add_para(doc,
    "The JR Validated Environment is a controlled R and Python execution environment "
    "for medical device design verification and statistical analysis. It consists of "
    "a controlled local R package repository (miniCRAN via Dropbox), a pinned isolated "
    "R library managed by renv, a pinned isolated Python virtual environment, the jrrun "
    "entry-point script which enforces the validated environment and logs every execution, "
    "and a suite of community scripts (jrc_*) for sample size calculation, diagnostic "
    "analysis, statistical analysis, design verification, data generation, and data "
    "conversion. This document covers the community script suite only. The environment "
    "infrastructure was qualified in JR-IQ-001.")

add_heading(doc, "5.2  Intended Use", level=2)
add_para(doc,
    "The community scripts are intended for use by engineers and scientists in the "
    "design verification phase of medical device development. Intended uses include:")
uses = [
    "Calculating minimum sample sizes for discrete and continuous design verification",
    "Assessing the statistical properties of pilot and verification datasets",
    "Computing and reporting process capability, outliers, and normality",
    "Performing method comparison studies (Bland-Altman)",
    "Performing reliability and lifetime analysis (Weibull)",
    "Verifying measurement datasets against specification limits using statistical tolerance intervals",
    "Generating synthetic datasets for validation and testing purposes",
    "Converting raw instrument data files into the standard jrc CSV format",
]
for u in uses:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(u).font.size = Pt(10)
add_para(doc,
    "Scripts produce console output and, where applicable, PNG plots and CSV files. "
    "All outputs are intended to support engineering reports and design history files "
    "(DHF) for FDA submissions.")

add_heading(doc, "5.3  User Profile", level=2)
add_para(doc,
    "Primary users are mechanical, biomedical, or quality engineers with working "
    "knowledge of design verification testing but who are not expected to have deep "
    "statistical programming expertise. Scripts are invoked from the macOS Terminal "
    "via wrapper commands (e.g., jrc_ss_discrete 0.99 0.95).")

add_heading(doc, "5.4  Hardware and Software Environment", level=2)
add_table(doc,
    ["Component", "Specification"],
    [
        ("Operating system", "macOS (arm64 / aarch64-apple-darwin)"),
        ("R version",        "4.5 (as specified in admin/r_version.txt)"),
        ("Python version",   "3.11.9 (as specified in admin/python_version.txt)"),
        ("R packages",       "As listed in admin/R_requirements.txt and admin/renv.lock"),
        ("Python packages",  "As listed in admin/python_requirements.txt"),
        ("Entry point",      "bin/jrrun"),
    ],
    col_widths=[5, 12]
)

# ---------------------------------------------------------------------------
# 6. RISK ASSESSMENT
# ---------------------------------------------------------------------------

add_heading(doc, "6.  Risk Assessment and Validation Level", level=1)
add_heading(doc, "6.1  Risk Classification", level=2)
add_para(doc,
    "The community scripts are used to compute sample sizes, process statistics, and "
    "tolerance intervals that directly inform design verification conclusions reported "
    "in FDA submissions. Incorrect results could lead to underpowered verification "
    "studies that fail to detect non-conforming processes, incorrect pass/fail verdicts "
    "in design verification reports, or erroneous statistical claims in regulatory "
    "submissions. This represents a high-consequence failure mode affecting patient "
    "safety indirectly via the quality of design verification evidence. Per the FDA "
    "General Principles of Software Validation (GPSV 2002), the level of validation "
    "effort is commensurate with this risk.")

add_heading(doc, "6.2  Validation Strategy", level=2)
add_table(doc,
    ["Risk Factor", "Assessment", "Consequence"],
    [
        ("Computational errors in sample size scripts",
         "Medium probability",
         "High: incorrect N in DHF"),
        ("Incorrect file I/O (wrong column, wrong delimiter)",
         "Medium probability",
         "High: analysis performed on wrong data"),
        ("Silent failure (non-zero exit, no error message)",
         "Low probability",
         "Medium: user unaware of failure"),
        ("Incorrect pass/fail verdict in jrc_verify_attr",
         "Low probability",
         "High: incorrect verification conclusion"),
    ],
    col_widths=[6, 4, 7]
)
doc.add_paragraph()
add_para(doc,
    "Strategy: Full OQ testing of all scripts, covering both primary (happy-path) "
    "execution paths and boundary/error-handling paths. Numerical outputs of pure-computation "
    "scripts are verified against independently calculated reference values where feasible.")

# ---------------------------------------------------------------------------
# 7. ROLES AND RESPONSIBILITIES
# ---------------------------------------------------------------------------

add_heading(doc, "7.  Roles and Responsibilities", level=1)
add_table(doc,
    ["Role", "Responsibility"],
    [
        ("Author",
         "Writes and maintains this Validation Plan"),
        ("Reviewer",
         "Reviews test specifications for completeness, accuracy, and regulatory adequacy"),
        ("QA Approver",
         "Approves this plan prior to OQ execution"),
        ("Test Executor",
         "Executes the OQ test suite (admin_oq) and documents results"),
        ("QA Witness",
         "Reviews execution evidence and signs the OQ execution report"),
    ],
    col_widths=[4, 13]
)
doc.add_paragraph()
add_para(doc,
    "The Test Executor and Author may be the same person. The QA Approver must be "
    "independent of the Author for this document.", italic=True)

# ---------------------------------------------------------------------------
# 8. VALIDATION APPROACH
# ---------------------------------------------------------------------------

add_heading(doc, "8.  Validation Approach", level=1)
add_heading(doc, "8.1  Overall Strategy", level=2)
add_para(doc,
    "Validation of the JR Validated Environment follows a three-phase IQ/OQ/PQ model:")
add_table(doc,
    ["Phase", "Description", "Status"],
    [
        ("IQ",
         "Verifies correct installation: R version, Python version, all required packages "
         "installed at pinned versions, SHA256 integrity of all scripts",
         "Completed — JR-IQ-001"),
        ("OQ",
         "Verifies correct operation of each community script across the range of anticipated "
         "inputs (this document)",
         "Planned"),
        ("PQ",
         "Verifies end-to-end performance under real conditions; executed under individual "
         "design verification protocols",
         "Out of scope here"),
    ],
    col_widths=[2, 12, 3]
)

doc.add_paragraph()
add_heading(doc, "8.2  OQ Execution", level=2)
add_para(doc,
    "OQ tests are executed by the automated test runner admin_oq, which:")
steps = [
    "Activates the OQ Python virtual environment (~/.venvs/MyProject_oq/)",
    "Invokes pytest against the test suite in oq/",
    "For each test case, calls the community script as a subprocess via jrrun",
    "Captures exit code, stdout, and stderr",
    "Evaluates pass/fail against the criteria defined in Section 13",
    "Writes a timestamped qualification report to ~/.jrscript/MyProject/validation/",
]
for s in steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)
add_para(doc,
    "All test cases must pass for the OQ to be considered successful. A single failing "
    "test case constitutes an OQ failure and must be resolved before the system is "
    "released to users.")

add_heading(doc, "8.3  Sequence of Execution", level=2)
seq = [
    "Generate and commit test data files to oq/data/ (see Section 11)",
    "Obtain QA approval of this plan",
    "Execute admin_oq on the target machine",
    "Review the qualification report",
    "Document any deviations (see Section 15)",
    "Obtain QA sign-off on the execution evidence",
]
for s in seq:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)

# ---------------------------------------------------------------------------
# 9. USER REQUIREMENTS
# ---------------------------------------------------------------------------

add_heading(doc, "9.  User Requirements", level=1)
add_para(doc,
    "The following user requirements define what the community script suite must do "
    "from the user's perspective. Each requirement is uniquely identified and traceable "
    "to one or more OQ test cases in the Requirements Traceability Matrix (Section 14).")

add_heading(doc, "9.1  Sample Size Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-001",
         "The system shall compute the minimum sample size for discrete (pass/fail) design "
         "verification for a given proportion and confidence level, using the exact binomial "
         "(chi-squared) method, for 0 to 10 allowed failures."),
        ("UR-002",
         "The system shall compute the achieved proportion for a discrete test result given "
         "confidence level, sample size, and number of failures."),
        ("UR-003",
         "The system shall compute the minimum sample size for continuous attribute design "
         "verification using statistical tolerance intervals, for 1-sided and 2-sided "
         "specifications, applying Box-Cox transformation when data are non-normal."),
        ("UR-004",
         "The system shall verify whether a user-specified planned sample size meets the "
         "tolerance interval requirement for a given dataset and specification."),
        ("UR-005",
         "The system shall compute the achieved proportion coverage for a continuous attribute "
         "verification dataset using tolerance intervals."),
        ("UR-006",
         "The system shall compute the minimum pilot sample size required to obtain a reliable "
         "estimate of the process standard deviation, for 1-sided and 2-sided specification cases."),
        ("UR-007",
         "The system shall compute the minimum sample size for a paired comparison study for a "
         "given detectable difference, standard deviation of differences, and test direction "
         "(1- or 2-sided)."),
        ("UR-008",
         "The system shall compute the minimum sample size for equivalence testing (TOST) for "
         "a given equivalence margin, standard deviation, and test direction."),
        ("UR-009",
         "The system shall compute the minimum sample size for fatigue and lifetime testing "
         "using a Weibull reliability model, for a given reliability, confidence, shape parameter, "
         "and acceleration factor, for 0 to 5 allowed failures."),
        ("UR-010",
         "The system shall provide Gauge R&R study design guidance (operators x replicates "
         "combinations) for a given target %GRR and reference type (process variation or tolerance)."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.2  Diagnostic Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-011",
         "The system shall test a dataset for normality using skewness, Shapiro-Wilk, and "
         "Anderson-Darling methods, and shall attempt a Box-Cox transformation if the data are "
         "non-normal, reporting the result and recommended approach."),
        ("UR-012",
         "The system shall detect outliers in a dataset using the Grubbs test (iterative) and "
         "the IQR method, reporting the row identifiers of flagged observations."),
        ("UR-013",
         "The system shall compute process capability indices (Cp, Cpk, Pp, Ppk) with 95% "
         "confidence intervals for a dataset and one or two specification limits."),
        ("UR-014",
         "The system shall compute a descriptive statistics summary for a dataset column, "
         "including: N, mean, median, SD, variance, CV, min, max, range, 5th/25th/75th/95th "
         "percentiles, IQR, skewness, excess kurtosis, and 95% CI on the mean."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.3  Statistical Analysis Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-015",
         "The system shall perform a Bland-Altman method comparison analysis for two paired "
         "datasets, reporting bias, SD of differences, limits of agreement with 95% CIs, and "
         "proportional bias test result, and shall save a Bland-Altman plot as a PNG file."),
        ("UR-016",
         "The system shall fit a 2-parameter Weibull distribution to lifetime data including "
         "right-censored observations using MLE, reporting shape and scale parameters with 95% "
         "CIs and B1, B10, and B50 life estimates with 95% CIs, and shall save a Weibull "
         "probability plot as a PNG file."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.4  Verification Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-017",
         "The system shall compute a statistical tolerance interval for a verification dataset "
         "for a given proportion and confidence, compare the tolerance limits against one or two "
         "specification limits, report a pass or fail verdict for each limit, apply Box-Cox "
         "transformation when data are non-normal, and save a histogram PNG file."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.5  Data Generation Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-018",
         "The system shall generate reproducible synthetic datasets from the following "
         "distributions: normal, log-normal, chi-squared scaled (sqrt), Weibull (boxcox), "
         "and uniform. Generated CSV files shall have two columns (id, value). When a seed "
         "is provided, repeated calls with identical parameters shall produce identical output."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.6  Data Conversion Requirements", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-019",
         "The system shall convert a multi-column delimited file to the standard jrc two-column "
         "CSV format (id, value) by extracting a user-specified column (by name or 1-based number) "
         "with configurable header line skipping and delimiter selection (tab, space, comma, or "
         "auto-detect). Non-numeric values shall be skipped with a warning."),
        ("UR-020",
         "The system shall convert a single-column plain text file (one numeric value per line) "
         "to the standard jrc two-column CSV format (id, value), with optional line range "
         "selection. Non-numeric and empty lines shall be skipped with a warning."),
    ],
    col_widths=[2.5, 14.5]
)

doc.add_paragraph()
add_heading(doc, "9.7  General Requirements (Applicable to All Scripts)", level=2)
add_table(doc,
    ["ID", "Requirement"],
    [
        ("UR-021",
         "The system shall validate all input arguments before computation. On receipt of "
         "invalid input (wrong type, out-of-range value, missing argument, file not found, "
         "column not found), the system shall exit with a non-zero exit code and print an "
         "informative error message that identifies the invalid argument or missing resource."),
        ("UR-022",
         "All R scripts shall verify at startup that the RENV_PATHS_ROOT environment variable "
         "is set, and shall halt with an informative error if it is not, preventing execution "
         "outside the jrrun validated environment."),
    ],
    col_widths=[2.5, 14.5]
)

# ---------------------------------------------------------------------------
# 10. TEST ENVIRONMENT
# ---------------------------------------------------------------------------

add_heading(doc, "10.  Test Environment", level=1)
add_table(doc,
    ["Item", "Specification"],
    [
        ("Test runner",             "pytest (version per oq/requirements.txt)"),
        ("Test execution method",   "Subprocess calls via jrrun; exit codes and stdout/stderr captured"),
        ("OQ venv",                 "~/.venvs/MyProject_oq/ — separate from user venv"),
        ("OQ venv requirements",    "oq/requirements.txt (frozen; not modified after OQ execution)"),
        ("R environment",           "As qualified in JR-IQ-001"),
        ("Python environment",      "As qualified in JR-IQ-001"),
        ("Test data location",      "oq/data/ (committed to the repository)"),
        ("Evidence output",         "~/.jrscript/MyProject/validation/oq_<datetime>.txt"),
        ("Machine",                 "The machine on which IQ was executed, or a machine where IQ has been completed"),
    ],
    col_widths=[5, 12]
)

# ---------------------------------------------------------------------------
# 11. TEST DATA
# ---------------------------------------------------------------------------

add_heading(doc, "11.  Test Data", level=1)
add_para(doc,
    "The following files shall be committed to oq/data/ before OQ execution. All CSV "
    "files use the standard two-column format (id, value) unless otherwise noted.")
add_table(doc,
    ["File", "Content", "Generation Method"],
    [
        ("normal_n30_mean10_sd1_seed42.csv",
         "30 values, N(mean=10, sd=1), seed 42",
         "jrc_gen_normal 30 10 1 oq/data 42"),
        ("skewed_n30_lognormal_seed42.csv",
         "30 values, log-normal (meanlog=2, sdlog=0.5), seed 42",
         "jrc_gen_lognormal 30 2.0 0.5 oq/data 42"),
        ("outlier_n30_seed42.csv",
         "As normal_n30_mean10_sd1_seed42.csv with row 15 replaced by 15.0 (≥5σ outlier)",
         "Generated then manually edited"),
        ("bland_altman_method1_seed42.csv",
         "25 values, N(mean=10, sd=1), seed 42",
         "jrc_gen_normal 25 10 1 oq/data 42 (rename)"),
        ("bland_altman_method2_seed42.csv",
         "Same 25 method1 values plus N(mean=0.3, sd=0.2) noise, seed 99",
         "oq/data/generate_ba_method2.R"),
        ("weibull_n20_seed42.csv",
         "20 rows; columns id, cycles, status; 15 failures, 5 censored; Weibull(shape=2, scale=1000), seed 42",
         "oq/data/generate_weibull.R"),
        ("method1_short.csv",
         "Rows 1-10 of bland_altman_method1_seed42.csv",
         "Subset of method1 file"),
        ("all_censored.csv",
         "10 rows; columns id, cycles, status; all status=0",
         "Prepared manually"),
        ("neg_times.csv",
         "10 rows; columns id, cycles, status; row 1 cycles = -100",
         "Prepared manually"),
        ("bad_status.csv",
         "10 rows; columns id, cycles, status; status values include 2",
         "Prepared manually"),
        ("convert_multicolumn.txt",
         "Tab-delimited; 3 metadata header lines; columns SampleID, ForceN, Temperature; 20 data rows",
         "Prepared manually"),
        ("convert_singlecolumn.txt",
         "200 lines, one numeric value per line, values from N(5, 0.5)",
         "Prepared manually"),
    ],
    col_widths=[5.5, 6.5, 5]
)
doc.add_paragraph()
add_para(doc,
    "Data integrity: SHA256 checksums of all files in oq/data/ shall be recorded in "
    "oq/data/checksums.sha256 and verified by admin_oq before test execution begins.",
    italic=True)

# ---------------------------------------------------------------------------
# 12. OQ TEST SPECIFICATIONS
# ---------------------------------------------------------------------------

add_heading(doc, "12.  OQ Test Specifications", level=1)

add_heading(doc, "Conventions", level=2)
add_para(doc,
    "Command — the command as typed by the user (via the jrrun wrapper). Full paths to "
    "data files are replaced by filename only for readability; admin_oq uses absolute paths.")
add_para(doc,
    "Expected output — text that must appear on stdout or stderr. R scripts use message() "
    "for all output (captured on stderr). Python scripts use print() (stdout).")
add_para(doc,
    "Pass criterion — the automated assertion made by the test runner.")
add_para(doc,
    "All test case IDs follow the pattern TC-<SCRIPT>-<NNN>.")

# ---- 12.1 jrc_ss_discrete ----
add_heading(doc, "12.1  jrc_ss_discrete    [Tests UR-001, UR-021]", level=2)
tcs = [
    ("TC-DISC-001", "UR-001", "jrc_ss_discrete 0.99 0.95",
     "For P=0.99, C=0.95, f=0: N = ceil(log(0.05)/log(0.99)) = 299 (exact binomial formula).",
     "Exit code 0; '299' present in output on the f=0 row."),
    ("TC-DISC-002", "UR-001", "jrc_ss_discrete 0.99 0.80",
     "Lower confidence reduces required N.",
     "Exit code 0; f=0 N is a positive integer less than 299."),
    ("TC-DISC-003", "UR-021", "jrc_ss_discrete 1.5 0.95", "",
     "Exit code non-zero; output contains 'proportion'."),
    ("TC-DISC-004", "UR-021", "jrc_ss_discrete 0.99 0",   "",
     "Exit code non-zero; output contains 'confidence'."),
    ("TC-DISC-005", "UR-021", "jrc_ss_discrete 0.99",     "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.2 jrc_ss_discrete_ci ----
add_heading(doc, "12.2  jrc_ss_discrete_ci    [Tests UR-002, UR-021]", level=2)
tcs = [
    ("TC-DISCICI-001", "UR-002", "jrc_ss_discrete_ci 0.95 299 0",
     "At N=299, f=0, C=0.95 the achieved proportion must be ≥ 0.99.",
     "Exit code 0; reported achieved proportion ≥ 0.99."),
    ("TC-DISCICI-002", "UR-002", "jrc_ss_discrete_ci 0.95 299 1",
     "One failure reduces achieved proportion.",
     "Exit code 0; achieved proportion < that reported in TC-DISCICI-001."),
    ("TC-DISCICI-003", "UR-021", "jrc_ss_discrete_ci 0.95 10 15", "",
     "Exit code non-zero; output references 'f' or 'n'."),
    ("TC-DISCICI-004", "UR-021", "jrc_ss_discrete_ci 0.95 299",   "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.3 jrc_ss_attr ----
add_heading(doc, "12.3  jrc_ss_attr    [Tests UR-003, UR-021]", level=2)
tcs = [
    ("TC-ATTR-001", "UR-003",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 -", "",
     "Exit code 0; output contains a positive integer minimum N ≥ 10; output contains '✅'."),
    ("TC-ATTR-002", "UR-003",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value - 11.0", "",
     "Exit code 0; output contains a positive integer minimum N ≥ 10."),
    ("TC-ATTR-003", "UR-003",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 11.0",
     "2-sided tolerance interval requires more samples than 1-sided.",
     "Exit code 0; 2-sided minimum N ≥ 1-sided N from TC-ATTR-001."),
    ("TC-ATTR-004", "UR-021",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 11.0 9.0", "",
     "Exit code non-zero; output contains 'spec2'."),
    ("TC-ATTR-005", "UR-021",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value - -", "",
     "Exit code non-zero."),
    ("TC-ATTR-006", "UR-021",
     "jrc_ss_attr 0.95 0.95 nonexistent.csv value 9.0 -", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-ATTR-007", "UR-021",
     "jrc_ss_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv badcol 9.0 -", "",
     "Exit code non-zero; output contains 'not found' or 'Available'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.4 jrc_ss_attr_check ----
add_heading(doc, "12.4  jrc_ss_attr_check    [Tests UR-004, UR-021]", level=2)
tcs = [
    ("TC-ATTRCK-001", "UR-004",
     "jrc_ss_attr_check 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 - 50", "",
     "Exit code 0; output contains a pass indicator ('✅' or 'PASS')."),
    ("TC-ATTRCK-002", "UR-004",
     "jrc_ss_attr_check 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 - 5", "",
     "Exit code 0; output contains a fail indicator ('❌' or 'FAIL')."),
    ("TC-ATTRCK-003", "UR-021",
     "jrc_ss_attr_check 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 -", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.5 jrc_ss_attr_ci ----
add_heading(doc, "12.5  jrc_ss_attr_ci    [Tests UR-005, UR-021]", level=2)
tcs = [
    ("TC-ATTRCI-001", "UR-005",
     "jrc_ss_attr_ci 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 -", "",
     "Exit code 0; output contains a proportion value between 0 and 1 (exclusive)."),
    ("TC-ATTRCI-002", "UR-005",
     "jrc_ss_attr_ci 0.95 normal_n30_mean10_sd1_seed42.csv value 9.0 11.0",
     "2-sided interval is more restrictive than 1-sided.",
     "Exit code 0; 2-sided proportion ≤ proportion from TC-ATTRCI-001."),
    ("TC-ATTRCI-003", "UR-021",
     "jrc_ss_attr_ci 0.95 nonexistent.csv value 9.0 -", "",
     "Exit code non-zero; output contains 'not found'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.6 jrc_ss_sigma ----
add_heading(doc, "12.6  jrc_ss_sigma    [Tests UR-006, UR-021]", level=2)
tcs = [
    ("TC-SIGMA-001", "UR-006", "jrc_ss_sigma 1.5 9.0 -", "",
     "Exit code 0; output contains a table of N values for multiple power/confidence combinations."),
    ("TC-SIGMA-002", "UR-006", "jrc_ss_sigma 1.5 9.0 11.0",
     "2-sided hypothesis test requires larger N than 1-sided.",
     "Exit code 0; N values ≥ corresponding values from TC-SIGMA-001."),
    ("TC-SIGMA-003", "UR-021", "jrc_ss_sigma -1.0 9.0 -", "",
     "Exit code non-zero."),
    ("TC-SIGMA-004", "UR-021", "jrc_ss_sigma 1.5", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.7 jrc_ss_paired ----
add_heading(doc, "12.7  jrc_ss_paired    [Tests UR-007, UR-021]", level=2)
tcs = [
    ("TC-PAIRED-001", "UR-007", "jrc_ss_paired 0.5 1.0 2", "",
     "Exit code 0; output contains a table of N values; all N ≥ 10."),
    ("TC-PAIRED-002", "UR-007", "jrc_ss_paired 0.5 1.0 1",
     "1-sided test requires fewer samples than 2-sided.",
     "Exit code 0; N values strictly less than corresponding values from TC-PAIRED-001."),
    ("TC-PAIRED-003", "UR-021", "jrc_ss_paired 0.5 1.0 3", "",
     "Exit code non-zero."),
    ("TC-PAIRED-004", "UR-021", "jrc_ss_paired 0.5 0 2", "",
     "Exit code non-zero."),
    ("TC-PAIRED-005", "UR-021", "jrc_ss_paired 0.5", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.8 jrc_ss_equivalence ----
add_heading(doc, "12.8  jrc_ss_equivalence    [Tests UR-008, UR-021]", level=2)
tcs = [
    ("TC-EQUIV-001", "UR-008", "jrc_ss_equivalence 0.5 1.0 2", "",
     "Exit code 0; output contains 'TOST' or 'equivalence' (case-insensitive); all N ≥ 10."),
    ("TC-EQUIV-002", "UR-008", "jrc_ss_equivalence 0.5 1.0 1",
     "1-sided non-inferiority requires fewer samples.",
     "Exit code 0; N values < corresponding values from TC-EQUIV-001."),
    ("TC-EQUIV-003", "UR-021", "jrc_ss_equivalence 0.5 1.0 0", "",
     "Exit code non-zero."),
    ("TC-EQUIV-004", "UR-021", "jrc_ss_equivalence 0.5", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.9 jrc_ss_fatigue ----
add_heading(doc, "12.9  jrc_ss_fatigue    [Tests UR-009, UR-021]", level=2)
tcs = [
    ("TC-FAT-001", "UR-009", "jrc_ss_fatigue 0.90 0.95 2.0 1.0", "",
     "Exit code 0; output contains a table for f=0..5; f=0 row has the largest N."),
    ("TC-FAT-002", "UR-009", "jrc_ss_fatigue 0.90 0.95 2.0 2.0",
     "Higher acceleration factor reduces required sample size.",
     "Exit code 0; f=0 N strictly less than f=0 N from TC-FAT-001."),
    ("TC-FAT-003", "UR-021", "jrc_ss_fatigue 1.0 0.95 2.0 1.0", "",
     "Exit code non-zero."),
    ("TC-FAT-004", "UR-021", "jrc_ss_fatigue 0.90 0.95 2.0 0.5", "",
     "Exit code non-zero."),
    ("TC-FAT-005", "UR-021", "jrc_ss_fatigue 0.90 0.95 2.0", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.10 jrc_ss_gauge_rr ----
add_heading(doc, "12.10  jrc_ss_gauge_rr    [Tests UR-010, UR-021]", level=2)
tcs = [
    ("TC-GRR-001", "UR-010", "jrc_ss_gauge_rr 10 process 1.0", "",
     "Exit code 0; output contains a table of operators x replicates; output contains '%GRR' and 'ndc'."),
    ("TC-GRR-002", "UR-010", "jrc_ss_gauge_rr 10 tolerance 5.0", "",
     "Exit code 0; output contains a table."),
    ("TC-GRR-003", "UR-021", "jrc_ss_gauge_rr 10 badtype 1.0", "",
     "Exit code non-zero."),
    ("TC-GRR-004", "UR-021", "jrc_ss_gauge_rr 10", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.11 jrc_normality ----
add_heading(doc, "12.11  jrc_normality    [Tests UR-011, UR-021]", level=2)
tcs = [
    ("TC-NORM-001", "UR-011",
     "jrc_normality normal_n30_mean10_sd1_seed42.csv value",
     "Data is N(10,1); normality should be confirmed.",
     "Exit code 0; output contains 'normal' (case-insensitive) and '✅'."),
    ("TC-NORM-002", "UR-011",
     "jrc_normality skewed_n30_lognormal_seed42.csv value",
     "Log-normal data is non-normal; Box-Cox path should be triggered.",
     "Exit code 0; output contains 'Box-Cox' or 'not normal'."),
    ("TC-NORM-003", "UR-021", "jrc_normality nonexistent.csv value", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-NORM-004", "UR-021",
     "jrc_normality normal_n30_mean10_sd1_seed42.csv badcol", "",
     "Exit code non-zero; output contains 'not found' or 'Available'."),
    ("TC-NORM-005", "UR-021",
     "jrc_normality normal_n30_mean10_sd1_seed42.csv", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.12 jrc_outliers ----
add_heading(doc, "12.12  jrc_outliers    [Tests UR-012, UR-021]", level=2)
tcs = [
    ("TC-OUT-001", "UR-012",
     "jrc_outliers normal_n30_mean10_sd1_seed42.csv value",
     "Data has no injected outliers; Grubbs and IQR should find none.",
     "Exit code 0; output contains 'no outlier' (case-insensitive) or a count of 0 flagged observations."),
    ("TC-OUT-002", "UR-012",
     "jrc_outliers outlier_n30_seed42.csv value",
     "Row 15 = 15.0 is ≥ 5σ above the mean and must be detected.",
     "Exit code 0; output references row '15' or reports ≥ 1 flagged observation."),
    ("TC-OUT-003", "UR-021", "jrc_outliers nonexistent.csv value", "",
     "Exit code non-zero."),
    ("TC-OUT-004", "UR-021",
     "jrc_outliers normal_n30_mean10_sd1_seed42.csv", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.13 jrc_capability ----
add_heading(doc, "12.13  jrc_capability    [Tests UR-013, UR-021]", level=2)
tcs = [
    ("TC-CAP-001", "UR-013",
     "jrc_capability normal_n30_mean10_sd1_seed42.csv value 7.0 13.0",
     "Data is N(10,1); spec is ±3σ. Cp ≈ 1.0; Cpk should be positive.",
     "Exit code 0; output contains 'Cp', 'Cpk', 'Pp', 'Ppk'; reported Cpk > 0."),
    ("TC-CAP-002", "UR-013",
     "jrc_capability normal_n30_mean10_sd1_seed42.csv value - 13.0", "",
     "Exit code 0; output contains 'Cpk' or 'Ppk'; no computation error."),
    ("TC-CAP-003", "UR-021",
     "jrc_capability normal_n30_mean10_sd1_seed42.csv value - -", "",
     "Exit code non-zero."),
    ("TC-CAP-004", "UR-021",
     "jrc_capability nonexistent.csv value 7.0 13.0", "",
     "Exit code non-zero."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.14 jrc_descriptive ----
add_heading(doc, "12.14  jrc_descriptive    [Tests UR-014, UR-021]", level=2)
tcs = [
    ("TC-DESC-001", "UR-014",
     "jrc_descriptive normal_n30_mean10_sd1_seed42.csv value", "",
     "Exit code 0; output contains 'mean', 'median', 'SD', 'min', 'max', 'skewness'; "
     "reported mean is within ±0.5 of 10.0."),
    ("TC-DESC-002", "UR-021", "jrc_descriptive nonexistent.csv value", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-DESC-003", "UR-021",
     "jrc_descriptive normal_n30_mean10_sd1_seed42.csv badcol", "",
     "Exit code non-zero."),
    ("TC-DESC-004", "UR-021",
     "jrc_descriptive normal_n30_mean10_sd1_seed42.csv", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.15 jrc_bland_altman ----
add_heading(doc, "12.15  jrc_bland_altman    [Tests UR-015, UR-021]", level=2)
tcs = [
    ("TC-BA-001", "UR-015",
     "jrc_bland_altman bland_altman_method1_seed42.csv value bland_altman_method2_seed42.csv value",
     "Method 2 was generated with a mean bias of +0.3; reported bias should be approximately +0.3.",
     "Exit code 0; output contains 'Bias'; output contains 'Limits of Agreement' or 'LoA'; "
     "reported bias is within ±0.2 of 0.3; PNG file created in directory of method1 file."),
    ("TC-BA-002", "UR-015",
     "jrc_bland_altman bland_altman_method1_seed42.csv value bland_altman_method2_seed42.csv value",
     "",
     "Exit code 0; output contains 'proportional bias' (case-insensitive)."),
    ("TC-BA-003", "UR-021",
     "jrc_bland_altman method1_short.csv value bland_altman_method2_seed42.csv value",
     "method1_short.csv has 10 rows; method2 has 25 rows — mismatch must be detected.",
     "Exit code non-zero; output contains 'different numbers' or 'observations'."),
    ("TC-BA-004", "UR-021",
     "jrc_bland_altman nonexistent.csv value bland_altman_method2_seed42.csv value", "",
     "Exit code non-zero."),
    ("TC-BA-005", "UR-021",
     "jrc_bland_altman bland_altman_method1_seed42.csv value bland_altman_method2_seed42.csv", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.16 jrc_weibull ----
add_heading(doc, "12.16  jrc_weibull    [Tests UR-016, UR-021]", level=2)
tcs = [
    ("TC-WEIB-001", "UR-016",
     "jrc_weibull weibull_n20_seed42.csv cycles status",
     "Data generated from Weibull(shape=2, scale=1000); MLE estimates should be within ±50% "
     "of true values given N=20.",
     "Exit code 0; output contains 'beta' and 'eta'; reported beta is in range (1.0, 3.0); "
     "output contains 'B1', 'B10', 'B50'; PNG file created."),
    ("TC-WEIB-002", "UR-021",
     "jrc_weibull all_censored.csv cycles status",
     "All units censored — fewer than 2 failures.",
     "Exit code non-zero; output contains 'failure' or 'At least 2'."),
    ("TC-WEIB-003", "UR-021",
     "jrc_weibull neg_times.csv cycles status",
     "Row 1 has cycles = -100.",
     "Exit code non-zero; output contains 'positive'."),
    ("TC-WEIB-004", "UR-021",
     "jrc_weibull bad_status.csv cycles status",
     "Status column contains value 2.",
     "Exit code non-zero; output contains '0' and '1'."),
    ("TC-WEIB-005", "UR-021", "jrc_weibull nonexistent.csv cycles status", "",
     "Exit code non-zero."),
    ("TC-WEIB-006", "UR-021", "jrc_weibull weibull_n20_seed42.csv cycles", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.17 jrc_verify_attr ----
add_heading(doc, "12.17  jrc_verify_attr    [Tests UR-017, UR-021]", level=2)
tcs = [
    ("TC-VER-001", "UR-017",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 7.0 -",
     "Data is N(10,1); LSL=7.0 is 3σ below mean; TI lower limit should exceed 7.0.",
     "Exit code 0; output contains '✅' and 'Lower Tolerance Limit greater than Lower Spec Limit'."),
    ("TC-VER-002", "UR-017",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 9.8 -",
     "LSL=9.8 is close to the mean; TI lower limit should be below 9.8.",
     "Exit code 0; output contains '❌' and 'Lower Tolerance Limit less than Lower Spec Limit'."),
    ("TC-VER-003", "UR-017",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 7.0 13.0", "",
     "Exit code 0; output contains '✅' and 'inside Spec'."),
    ("TC-VER-004", "UR-017",
     "jrc_verify_attr 0.95 0.95 skewed_n30_lognormal_seed42.csv value 1.0 -",
     "Log-normal data triggers Box-Cox path; tolerance limit must be in original units.",
     "Exit code 0; output contains 'boxcox' or 'Box-Cox'; no computation error."),
    ("TC-VER-005", "UR-021",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 11.0 9.0", "",
     "Exit code non-zero; output contains 'spec2'."),
    ("TC-VER-006", "UR-017",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 7.0 -",
     "PNG output is required per UR-017.",
     "A file matching '*_tolerance.png' is present in the directory of the input CSV after execution."),
    ("TC-VER-007", "UR-021",
     "jrc_verify_attr 0.95 0.95 nonexistent.csv value 7.0 -", "",
     "Exit code non-zero."),
    ("TC-VER-008", "UR-021",
     "jrc_verify_attr 0.95 0.95 normal_n30_mean10_sd1_seed42.csv value 7.0", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.18–12.22 gen scripts ----
gen_specs = [
    ("12.18", "jrc_gen_normal",   "UR-018, UR-021",
     "normal",    "n50_mean10_sd1_seed42",   "jrc_gen_normal 50 10.0 1.0 <output_dir> 42",
     "sd <= 0",   "jrc_gen_normal 50 10.0 0 <output_dir> 42",
     "all values",  "all values in 'value' column present; no constraint on range"),
    ("12.19", "jrc_gen_lognormal","UR-018, UR-021",
     "lognormal", "n50_meanlog2_sdlog0.5_seed42", "jrc_gen_lognormal 50 2.0 0.5 <output_dir> 42",
     "sdlog <= 0", "jrc_gen_lognormal 50 2.0 0 <output_dir> 42",
     "all values > 0", "all values in 'value' column are strictly positive"),
    ("12.20", "jrc_gen_sqrt",     "UR-018, UR-021",
     "sqrt",      "n50_df3_scale1_seed42",   "jrc_gen_sqrt 50 3 1.0 <output_dir> 42",
     "df <= 0",   "jrc_gen_sqrt 50 0 1.0 <output_dir> 42",
     "all values >= 0", "all values in 'value' column are non-negative"),
    ("12.21", "jrc_gen_boxcox",   "UR-018, UR-021",
     "boxcox",    "n50_shape2_scale1000_seed42", "jrc_gen_boxcox 50 2.0 1000.0 <output_dir> 42",
     "shape <= 0","jrc_gen_boxcox 50 0 1000.0 <output_dir> 42",
     "all values > 0", "all values in 'value' column are strictly positive"),
    ("12.22", "jrc_gen_uniform",  "UR-018, UR-021",
     "uniform",   "n50_min0_max10_seed42",   "jrc_gen_uniform 50 0.0 10.0 <output_dir> 42",
     "max <= min","jrc_gen_uniform 50 10.0 5.0 <output_dir> 42",
     "all values in [0.0, 10.0]", "all values satisfy 0.0 <= value <= 10.0"),
]
for sec, script, tests_ur, dist, stem, cmd, err_desc, err_cmd, bounds_label, bounds_check in gen_specs:
    add_heading(doc, f"{sec}  {script}    [Tests {tests_ur}]", level=2)
    tcs_gen = [
        (f"TC-GEN-{dist.upper()[:2]}-001", tests_ur.split(",")[0].strip(),
         cmd,
         f"With seed 42 the output is deterministic.",
         f"Exit code 0; CSV created; header is 'id,value'; exactly 50 data rows; {bounds_check}."),
        (f"TC-GEN-{dist.upper()[:2]}-002", tests_ur.split(",")[0].strip(),
         f"Run TC-GEN-{dist.upper()[:2]}-001 twice; compare SHA256 of both outputs.",
         "Reproducibility requirement per UR-018.",
         "SHA256 checksums of both output files are identical."),
        (f"TC-GEN-{dist.upper()[:2]}-003", "UR-021",
         err_cmd,
         f"Invalid argument: {err_desc}.",
         "Exit code non-zero."),
        (f"TC-GEN-{dist.upper()[:2]}-004", "UR-021",
         f"{script} 50", "",
         "Exit code non-zero; output contains 'Usage'."),
    ]
    for tc in tcs_gen:
        add_tc_block(doc, *tc)

# ---- 12.23 jrc_convert_csv ----
add_heading(doc, "12.23  jrc_convert_csv    [Tests UR-019, UR-021]", level=2)
tcs = [
    ("TC-CCSV-001", "UR-019",
     "jrc_convert_csv convert_multicolumn.txt ForceN 3",
     "Column by name; auto-detect delimiter; skip 3 header lines.",
     "Exit code 0; output CSV created; header is 'id,value'; row count equals data rows in source; output contains '✅'."),
    ("TC-CCSV-002", "UR-019",
     "jrc_convert_csv convert_multicolumn.txt 2 3",
     "Column 2 (1-based) = ForceN; result must match TC-CCSV-001.",
     "Exit code 0; output CSV values identical to TC-CCSV-001."),
    ("TC-CCSV-003", "UR-019",
     "jrc_convert_csv convert_multicolumn.txt ForceN 3 tab",
     "Explicit tab delimiter; result must match TC-CCSV-001.",
     "Exit code 0; result identical to TC-CCSV-001."),
    ("TC-CCSV-004", "UR-021",
     "jrc_convert_csv convert_multicolumn.txt ForceN 999", "",
     "Exit code non-zero; output contains 'skip_lines'."),
    ("TC-CCSV-005", "UR-021",
     "jrc_convert_csv convert_multicolumn.txt NonExistent 3", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-CCSV-006", "UR-021",
     "jrc_convert_csv nonexistent.txt ForceN 0", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-CCSV-007", "UR-021",
     "jrc_convert_csv convert_multicolumn.txt ForceN 3 pipe", "",
     "Exit code non-zero; output contains 'delimiter'."),
    ("TC-CCSV-008", "UR-021",
     "jrc_convert_csv convert_multicolumn.txt ForceN", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---- 12.24 jrc_convert_txt ----
add_heading(doc, "12.24  jrc_convert_txt    [Tests UR-020, UR-021]", level=2)
tcs = [
    ("TC-CTXT-001", "UR-020",
     "jrc_convert_txt convert_singlecolumn.txt", "",
     "Exit code 0; output CSV created; header is 'id,value'; exactly 200 data rows; output contains '✅'."),
    ("TC-CTXT-002", "UR-020",
     "jrc_convert_txt convert_singlecolumn.txt 50 100", "",
     "Exit code 0; output CSV has exactly 51 data rows; filename contains 'lines50to100'."),
    ("TC-CTXT-003", "UR-020",
     "jrc_convert_txt convert_singlecolumn.txt 150",
     "No end_line; defaults to end of file (line 200).",
     "Exit code 0; output CSV has 51 data rows; filename contains 'lines150to200'."),
    ("TC-CTXT-004", "UR-021",
     "jrc_convert_txt convert_singlecolumn.txt 500", "",
     "Exit code non-zero; output contains 'start_line' or 'exceeds'."),
    ("TC-CTXT-005", "UR-021",
     "jrc_convert_txt convert_singlecolumn.txt 100 50", "",
     "Exit code non-zero; output contains 'end_line'."),
    ("TC-CTXT-006", "UR-021",
     "jrc_convert_txt nonexistent.txt", "",
     "Exit code non-zero; output contains 'not found'."),
    ("TC-CTXT-007", "UR-021",
     "jrc_convert_txt", "",
     "Exit code non-zero; output contains 'Usage'."),
]
for tc in tcs:
    add_tc_block(doc, *tc)

# ---------------------------------------------------------------------------
# 13. ACCEPTANCE CRITERIA
# ---------------------------------------------------------------------------

add_heading(doc, "13.  Acceptance Criteria", level=1)
add_heading(doc, "13.1  Individual Test Case Criterion", level=2)
add_para(doc, "A test case passes when all of the following are true:")
items = [
    "The subprocess exit code matches the expected value (0 for success cases, non-zero for error cases).",
    "All required strings are present in the combined stdout + stderr output.",
    "Where a file output is asserted, the file exists at the expected path.",
    "Where a numerical bound is asserted (e.g., N ≥ 10, reported mean within ±0.5 of 10.0), "
    "the extracted value satisfies the bound.",
]
for item in items:
    p = doc.add_paragraph(style="List Number")
    p.add_run(item).font.size = Pt(10)

add_heading(doc, "13.2  OQ Acceptance Criterion", level=2)
add_para(doc,
    "The OQ is passed when all 122 test cases pass with no deviations. A single failing "
    "test case constitutes an OQ failure and must be resolved before the system is "
    "released for use.")

add_heading(doc, "13.3  Exclusions", level=2)
add_para(doc,
    "No test cases may be excluded without documented justification reviewed and approved "
    "by QA. An exclusion is not the same as a deviation — exclusions must be agreed before "
    "OQ execution; deviations are handled after (Section 15).")

# ---------------------------------------------------------------------------
# 14. TRACEABILITY MATRIX
# ---------------------------------------------------------------------------

add_heading(doc, "14.  Requirements Traceability Matrix", level=1)
add_table(doc,
    ["User Requirement", "Test Cases"],
    [
        ("UR-001 — Discrete sample size",
         "TC-DISC-001, TC-DISC-002, TC-DISC-003, TC-DISC-004, TC-DISC-005"),
        ("UR-002 — Discrete achieved proportion",
         "TC-DISCICI-001, TC-DISCICI-002, TC-DISCICI-003, TC-DISCICI-004"),
        ("UR-003 — Attribute sample size",
         "TC-ATTR-001, TC-ATTR-002, TC-ATTR-003, TC-ATTR-004, TC-ATTR-005, TC-ATTR-006, TC-ATTR-007"),
        ("UR-004 — Attribute check",
         "TC-ATTRCK-001, TC-ATTRCK-002, TC-ATTRCK-003"),
        ("UR-005 — Attribute achieved proportion",
         "TC-ATTRCI-001, TC-ATTRCI-002, TC-ATTRCI-003"),
        ("UR-006 — Sigma pilot size",
         "TC-SIGMA-001, TC-SIGMA-002, TC-SIGMA-003, TC-SIGMA-004"),
        ("UR-007 — Paired comparison size",
         "TC-PAIRED-001, TC-PAIRED-002, TC-PAIRED-003, TC-PAIRED-004, TC-PAIRED-005"),
        ("UR-008 — Equivalence size",
         "TC-EQUIV-001, TC-EQUIV-002, TC-EQUIV-003, TC-EQUIV-004"),
        ("UR-009 — Fatigue sample size",
         "TC-FAT-001, TC-FAT-002, TC-FAT-003, TC-FAT-004, TC-FAT-005"),
        ("UR-010 — Gauge R&R design",
         "TC-GRR-001, TC-GRR-002, TC-GRR-003, TC-GRR-004"),
        ("UR-011 — Normality testing",
         "TC-NORM-001, TC-NORM-002, TC-NORM-003, TC-NORM-004, TC-NORM-005"),
        ("UR-012 — Outlier detection",
         "TC-OUT-001, TC-OUT-002, TC-OUT-003, TC-OUT-004"),
        ("UR-013 — Process capability",
         "TC-CAP-001, TC-CAP-002, TC-CAP-003, TC-CAP-004"),
        ("UR-014 — Descriptive statistics",
         "TC-DESC-001, TC-DESC-002, TC-DESC-003, TC-DESC-004"),
        ("UR-015 — Bland-Altman analysis",
         "TC-BA-001, TC-BA-002, TC-BA-003, TC-BA-004, TC-BA-005"),
        ("UR-016 — Weibull analysis",
         "TC-WEIB-001, TC-WEIB-002, TC-WEIB-003, TC-WEIB-004, TC-WEIB-005, TC-WEIB-006"),
        ("UR-017 — Tolerance interval verification",
         "TC-VER-001, TC-VER-002, TC-VER-003, TC-VER-004, TC-VER-005, TC-VER-006, TC-VER-007, TC-VER-008"),
        ("UR-018 — Synthetic data generation",
         "TC-GEN-NO-001..004, TC-GEN-LO-001..004, TC-GEN-SQ-001..004, "
         "TC-GEN-BO-001..004, TC-GEN-UN-001..004"),
        ("UR-019 — CSV conversion",
         "TC-CCSV-001, TC-CCSV-002, TC-CCSV-003, TC-CCSV-004, TC-CCSV-005, "
         "TC-CCSV-006, TC-CCSV-007, TC-CCSV-008"),
        ("UR-020 — Text file conversion",
         "TC-CTXT-001, TC-CTXT-002, TC-CTXT-003, TC-CTXT-004, TC-CTXT-005, "
         "TC-CTXT-006, TC-CTXT-007"),
        ("UR-021 — Input validation and error handling",
         "At least one error-path TC per script (throughout Section 12)"),
        ("UR-022 — Environment enforcement",
         "TC-GEN-NO-001 (RENV_PATHS_ROOT verified at startup via jrrun)"),
    ],
    col_widths=[6, 11]
)

# ---------------------------------------------------------------------------
# 15. DEVIATION HANDLING
# ---------------------------------------------------------------------------

add_heading(doc, "15.  Deviation Handling", level=1)
add_para(doc,
    "A deviation is any departure from the test procedure or any test case result "
    "that does not meet the pass criterion defined in Section 13.")

add_heading(doc, "15.1  During Execution", level=2)
add_para(doc, "If a test case fails during OQ execution:")
steps = [
    "The test executor records the deviation: test case ID, observed result, expected result, "
    "and the full output captured by admin_oq.",
    "Execution of the remaining test cases continues unless the failure prevents subsequent "
    "tests from running.",
    "The test executor does not modify test data, test scripts, or community scripts during "
    "execution without prior QA approval.",
]
for s in steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(10)

add_heading(doc, "15.2  Deviation Resolution", level=2)
add_para(doc,
    "Each deviation must be resolved before the OQ is closed. Resolution options are:")
add_table(doc,
    ["Resolution", "Condition"],
    [
        ("Correction + retest",
         "A defect in the community script is identified and corrected; affected test cases "
         "are re-executed after the correction is committed and the integrity hash is regenerated."),
        ("Test case correction",
         "The test case itself contains an error (e.g., wrong expected value); the correction "
         "is documented, reviewed, and approved before retesting."),
        ("Exclusion with justification",
         "In exceptional circumstances, QA may accept a documented justification for excluding "
         "a test case; this requires written approval and must be recorded in the OQ Execution Report."),
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()
add_para(doc,
    "All deviations and their resolutions shall be documented in the OQ Execution Report.")

# ---------------------------------------------------------------------------
# 16. EVIDENCE AND DOCUMENTATION
# ---------------------------------------------------------------------------

add_heading(doc, "16.  Evidence and Documentation", level=1)
add_para(doc,
    "Successful completion of the OQ produces the following evidence:")
add_table(doc,
    ["Evidence Item", "Location", "Retained By"],
    [
        ("This Validation Plan (approved, signed)",
         "docs/oq_validation_plan.docx + signed PDF", "QA"),
        ("Test data files with SHA256 checksums",
         "oq/data/ (committed to repository)", "Repository"),
        ("OQ test suite source code",
         "oq/ (committed to repository)", "Repository"),
        ("OQ execution report (automated, timestamped)",
         "~/.jrscript/MyProject/validation/oq_<datetime>.txt", "QA"),
        ("OQ execution report (copy)",
         "docs/", "Repository"),
        ("Deviation log (if any)",
         "Appended to OQ execution report", "QA"),
    ],
    col_widths=[6, 7, 4]
)
doc.add_paragraph()
add_para(doc,
    "All evidence shall be retained for the lifetime of the device plus the period required "
    "by applicable regulations (minimum 2 years post-market per 21 CFR 820.180, or per the "
    "applicable device-specific regulation).")

# ---------------------------------------------------------------------------
# 17. RELATED DOCUMENTS
# ---------------------------------------------------------------------------

add_heading(doc, "17.  Related Documents", level=1)
add_table(doc,
    ["Document", "Location"],
    [
        ("JR-VP-001 — IQ Validation Plan",          "docs/validation_plan.docx"),
        ("IQ Execution Evidence (JR-IQ-001)",        "docs/IQ_validation_20260311_205146.txt"),
        ("Validation Report template",               "docs/templates/jr_validation_report_template.docx"),
        ("CHANGELOG",                                "CHANGELOG.md"),
        ("Infrastructure Backlog",                   "INFRA_BACKLOG.md"),
        ("CONTRIBUTING guide",                       "CONTRIBUTING.md"),
    ],
    col_widths=[8, 9]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("End of Document — JR-VP-002 v1.0")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT)
print(f"✅  Saved: {OUT}")
